from fastapi import FastAPI, Depends, UploadFile, File, Form, BackgroundTasks, Request
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from sqlalchemy.orm import Session
from sqlalchemy import text, func, or_
from typing import List, Dict, Any
import os
from pathlib import Path
import uuid
from dotenv import load_dotenv

from database import get_db, engine
from models import Base, Feedback, UploadedFile, MandatoryFile, ProjectKnowledgeBaseFile, ChatMessage, Conversation, Project
from schemas import (
    LoginRequest, LoginResponse, 
    SprintStartRequest, SprintStartResponse,
    ChatRequest, ChatResponse,
    LLMChatRequest, LLMChatResponse,
    RiskAssessmentStartRequest, RiskAssessmentStartResponse,
    RiskAssessmentChatRequest, RiskAssessmentChatResponse,
    RiskAssessmentFinishRequest, RiskAssessmentFinishResponse,
    GenerateRiskAssessmentRequest,
    FeedbackRequest, FeedbackResponse
)
from services import (
    auth_service, sprint_service, 
    llm_service, gemini_service, db_service
)
from services.risk_service import risk_service
from services.docx_service import docx_service
from services.risk_docx_service import risk_docx_service
from services.pdf_service import pdf_service

# Load environment variables
load_dotenv()

# Run automatic migrations first (may drop/recreate tables)
try:
    from db_migrations import run_migrations
    run_migrations()
    print("[OK] Database migrations completed")
except Exception as e:
    print(f"[WARNING] Database migrations failed: {str(e)}")
    print("[INFO] Continuing startup - some features may not work correctly")

# Create database tables (creates new tables if they don't exist)
# This runs after migrations to recreate any dropped tables
Base.metadata.create_all(bind=engine)

app = FastAPI(
    title="Sprint Planning Demo API",
    description="A demo API for sprint planning with LLM integration",
    version="1.0.0"
)

def _get_structured_html_system_prompt() -> str:
    """System prompt that enforces structured HTML responses for the chatbot."""
    return (
        "You are a structured project assistant chatbot.\n\n"
        "Your goal is to provide clean, organized, and readable answers based on the uploaded project file content.\n"
        "Make sure your responses follow these formatting rules:\n\n"
        "1. Use HTML structure for responses.\n"
        "2. Headings → use <h3> or <h4> tags.\n"
        "3. Lists → use <ul><li> for bullet points and <ol><li> for numbered lists.\n"
        "4. For any links (e.g., URLs found in text), make them clickable using: <a href=\"URL\" target=\"_blank\">View Document</a>\n"
        "5. Avoid showing asterisks (*) or Markdown formatting.\n"
        "6. Keep responses concise, professional, and well-indented.\n"
        "7. If a section includes multiple deliverables or items, list each one separately for readability.\n\n"
        "Example formatting:\n"
        "---------------------\n"
        "<h3>Define Phase Deliverables</h3>\n"
        "<ul>\n"
        "  <li><strong>Project Plan:</strong> Identifies all phases, activities, deliverables, and milestones.</li>\n"
        "  <li><strong>Contact List:</strong> Contains team member contact information.</li>\n"
        "  <li><strong>Project Control Procedures:</strong> Describes reporting mechanisms and change control process.</li>\n"
        "</ul>\n"
        "<p>For more details, refer to <a href=\"https://example.com\" target=\"_blank\">Project Control Document</a>.</p>\n\n"
        "Now, based on the uploaded document, answer the user’s question clearly and in this structured HTML format."
    )

# Global variable to store prompt data in main flow
GLOBAL_PROMPT_DATA = None

def get_global_prompt_data():
    """Get the global prompt data stored in main flow"""
    global GLOBAL_PROMPT_DATA
    print(f"🔍 [GLOBAL PROMPT] Variable Name: GLOBAL_PROMPT_DATA")
    print(f"🔍 [GLOBAL PROMPT] Variable Value: {GLOBAL_PROMPT_DATA[:100] if GLOBAL_PROMPT_DATA else 'None'}...")
    print(f"🔍 [GLOBAL PROMPT] Variable Type: {type(GLOBAL_PROMPT_DATA)}")
    return GLOBAL_PROMPT_DATA

# CORS middleware
def _get_cors_origins() -> List[str]:
    """Return the list of allowed CORS origins pulled strictly from environment."""
    raw_origins = os.getenv("CORS_ORIGINS", "")
    return [origin.strip() for origin in raw_origins.split(",") if origin.strip()]


app.add_middleware(
    CORSMiddleware,
    allow_origins=_get_cors_origins(),
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.get("/")
async def root():
    """Health check endpoint"""
    return {"message": "Sprint Planning Demo API", "status": "running"}

# Static file serving endpoint for mandatory files (by ID)
@app.get("/api/mandatory-files/{file_id}/download")
async def download_mandatory_file(file_id: int, db: Session = Depends(get_db)):
    """Download a mandatory file by ID from database"""
    from fastapi.responses import Response
    
    try:
        mandatory_file = db.query(MandatoryFile).filter(
            MandatoryFile.id == file_id,
            MandatoryFile.is_active == True
        ).first()
        
        if not mandatory_file:
            return {"success": False, "error": "File not found"}
        
        # Check if file content exists in database
        if mandatory_file.file_content:
            # Serve file from database
            file_size = len(mandatory_file.file_content) if mandatory_file.file_content else 0
            print(f"📥 [MANDATORY-DOWNLOAD] Serving file '{mandatory_file.file_name}' (ID: {file_id}) from DATABASE - Size: {file_size} bytes")
            return Response(
                content=mandatory_file.file_content,
                media_type='application/octet-stream',
                headers={
                    "Content-Disposition": f'attachment; filename="{mandatory_file.file_name}"'
                }
            )
        elif mandatory_file.file_path:
            # Fallback: Try to serve from file system (for legacy files)
            file_path = Path(mandatory_file.file_path)
            if file_path.exists():
                print(f"📥 [MANDATORY-DOWNLOAD] Serving file '{mandatory_file.file_name}' (ID: {file_id}) from FILE SYSTEM (legacy) - Path: {file_path}")
                return FileResponse(
                    path=str(file_path),
                    filename=mandatory_file.file_name,
                    media_type='application/octet-stream'
                )
            else:
                print(f"⚠️ [MANDATORY-DOWNLOAD] File '{mandatory_file.file_name}' (ID: {file_id}) not found in database or file system")
                return {"success": False, "error": "File content not found in database or file system"}
        else:
            print(f"⚠️ [MANDATORY-DOWNLOAD] File '{mandatory_file.file_name}' (ID: {file_id}) has no file_content or file_path")
            return {"success": False, "error": "File content not found"}
            
    except Exception as e:
        print(f"❌ [MANDATORY-DOWNLOAD] Error serving file ID {file_id}: {str(e)}")
        return {"success": False, "error": f"Error serving file: {str(e)}"}

# Legacy endpoint for backward compatibility
@app.get("/mandatory/{filename:path}")
async def get_mandatory_file(filename: str):
    """Serve mandatory files from backend/mandatory directory (legacy)"""
    try:
        # Get the backend directory path
        backend_dir = Path(__file__).parent
        mandatory_dir = backend_dir / "mandatory"
        file_path = mandatory_dir / filename
        
        # Security: Ensure file is within mandatory directory (prevent path traversal)
        try:
            file_path.resolve().relative_to(mandatory_dir.resolve())
        except ValueError:
            return {"success": False, "error": "Invalid file path"}
        
        # Check if file exists
        if not file_path.exists():
            return {"success": False, "error": f"File not found: {filename}"}
        
        # Return file with appropriate media type
        return FileResponse(
            path=str(file_path),
            filename=filename,
            media_type='application/octet-stream'
        )
    except Exception as e:
        return {"success": False, "error": f"Error serving file: {str(e)}"}

# Mandatory Files Management Endpoints
@app.get("/api/mandatory-files")
async def get_mandatory_files(db: Session = Depends(get_db), include_content: bool = False):
    """Get all active mandatory files"""
    try:
        files = db.query(MandatoryFile).filter(MandatoryFile.is_active == True).order_by(MandatoryFile.uploaded_at.desc()).all()
        return {
            "success": True,
            "files": [
                {
                    "id": f.id,
                    "file_name": f.file_name,
                    "file_type": f.file_type,
                    "file_size": f.file_size,
                    "uploaded_by": f.uploaded_by,
                    "uploaded_at": f.uploaded_at.isoformat() if f.uploaded_at else None,
                    "description": f.description,
                    "extracted_text": f.extracted_text if include_content else None,  # Only include if requested
                    "has_content": bool(f.extracted_text)  # Indicate if content exists
                }
                for f in files
            ]
        }
    except Exception as e:
        return {"success": False, "error": f"Error fetching mandatory files: {str(e)}"}

@app.post("/api/mandatory-files/upload")
async def upload_mandatory_file(
    file: UploadFile = File(...),
    uploaded_by: str = Form(None),
    description: str = Form(None),
    background_tasks: BackgroundTasks = None,
    db: Session = Depends(get_db)
):
    """Upload a new mandatory file"""
    import uuid
    from datetime import datetime
    
    try:
        # Validate file type
        file_extension = file.filename.lower().split('.')[-1] if '.' in file.filename else ''
        
        if file_extension not in ['pdf', 'docx', 'txt', 'doc', 'xlsx', 'xls', 'pptx', 'ppt']:
            return {
                "success": False,
                "error": f"Invalid file type. Supported formats: PDF, DOCX, TXT, XLSX, PPTX. Got: {file_extension}"
            }
        
        # Read file content into memory (stored in database, not file system)
        file_content = await file.read()
        file_size = len(file_content)
        
        # Extract text content based on file type
        extracted_text = ""
        try:
            if file_extension == 'pdf':
                result = pdf_service.extract_text_from_pdf(file_content)
                if result['success']:
                    extracted_text = result['text']
            elif file_extension in ['docx', 'doc']:
                # Use enhanced docx extraction to preserve hyperlinks
                try:
                    from services.docx_extraction_helper import extract_text_with_hyperlinks_from_docx
                    extracted_text = extract_text_with_hyperlinks_from_docx(file_content)
                except Exception as e:
                    # Fallback to simple extraction if enhanced extraction fails
                    print(f"⚠️ [MANDATORY-UPLOAD] Enhanced DOCX extraction failed, using fallback: {str(e)}")
                    from docx import Document
                    import io
                    doc = Document(io.BytesIO(file_content))
                    text_parts = []
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            text_parts.append(paragraph.text.strip())
                    extracted_text = '\n'.join(text_parts)
            elif file_extension == 'txt':
                extracted_text = file_content.decode('utf-8', errors='ignore')
            elif file_extension in ['xlsx', 'xls']:
                import io
                from openpyxl import load_workbook
                workbook = load_workbook(filename=io.BytesIO(file_content), data_only=True)
                lines = []
                for sheet in workbook.worksheets:
                    lines.append(f"Sheet: {sheet.title}")
                    for row in sheet.iter_rows(values_only=True):
                        cells = [str(cell) for cell in row if cell is not None]
                        if cells:
                            lines.append("\t".join(cells))
                extracted_text = "\n".join(lines)
            elif file_extension in ['pptx', 'ppt']:
                # PowerPoint files - basic extraction (can be enhanced later)
                extracted_text = f"[PowerPoint file: {file.filename}]"
        except Exception as e:
            print(f"Warning: Could not extract text from file: {e}")
            extracted_text = None  # Store None if extraction fails
        
        # Save to database (file content stored in DB, not file system)
        user_email = uploaded_by or "anonymous"
        mandatory_file = MandatoryFile(
            file_name=file.filename,
            file_type=file_extension,
            file_path=None,  # No longer using file system storage
            file_content=file_content,  # Store file content in database
            file_size=file_size,
            uploaded_by=user_email,
            description=description,
            is_active=True,
            extracted_text=extracted_text  # Store extracted text content
        )
        
        db.add(mandatory_file)
        db.commit()
        db.refresh(mandatory_file)
        
        print(f"💾 [MANDATORY-UPLOAD] Saved file '{file.filename}' (ID: {mandatory_file.id}) to DATABASE - Size: {file_size} bytes, User: {user_email}")
        
        return {
            "success": True,
            "file_id": mandatory_file.id,
            "file_name": mandatory_file.file_name,
            "message": f"File '{file.filename}' uploaded successfully"
        }
        
    except Exception as e:
        db.rollback()
        return {
            "success": False,
            "error": f"Error uploading file: {str(e)}"
        }

@app.delete("/api/mandatory-files/{file_id}")
async def delete_mandatory_file(file_id: int, db: Session = Depends(get_db)):
    """Delete a mandatory file permanently from database"""
    try:
        mandatory_file = db.query(MandatoryFile).filter(MandatoryFile.id == file_id).first()
        
        if not mandatory_file:
            return {"success": False, "error": "File not found"}
        
        file_name = mandatory_file.file_name
        
        # Remove from all users' project knowledge bases first
        db.query(ProjectKnowledgeBaseFile).filter(
            ProjectKnowledgeBaseFile.mandatory_file_id == file_id
        ).delete()
        
        # Hard delete - permanently remove from database (file content is in DB, no file system cleanup needed)
        db.delete(mandatory_file)
        db.commit()
        
        return {
            "success": True,
            "message": f"File '{file_name}' deleted successfully"
        }
        
    except Exception as e:
        db.rollback()
        return {
            "success": False,
            "error": f"Error deleting file: {str(e)}"
        }

# Project Knowledge Base Endpoints
@app.post("/api/project-knowledge-base/add")
async def add_project_knowledge_base_file(
    file_id: int = Form(...),
    user_email: str = Form(...),
    db: Session = Depends(get_db)
):
    """Add a file to user's project knowledge base"""
    try:
        # Verify file exists and is active
        mandatory_file = db.query(MandatoryFile).filter(
            MandatoryFile.id == file_id,
            MandatoryFile.is_active == True
        ).first()
        
        if not mandatory_file:
            return {
                "success": False,
                "error": "File not found or inactive"
            }
        
        # Check if already exists
        existing = db.query(ProjectKnowledgeBaseFile).filter(
            ProjectKnowledgeBaseFile.user_email == user_email,
            ProjectKnowledgeBaseFile.mandatory_file_id == file_id
        ).first()
        
        if existing:
            # File already in knowledge base, but still ensure it's indexed in Pinecone
            print(f"📝 [KNOWLEDGE-BASE] File {file_id} already in knowledge base, ensuring Pinecone index exists...")
            should_index = True  # Still index/reindex to ensure Pinecone is up to date
        else:
            # Add to knowledge base
            knowledge_base_file = ProjectKnowledgeBaseFile(
                user_email=user_email,
                mandatory_file_id=file_id
            )
            
            db.add(knowledge_base_file)
            db.commit()
            db.refresh(knowledge_base_file)
            should_index = True
        
        # Index file to Pinecone (separate index per file)
        # This runs whether file is new or existing to ensure Pinecone is synced
        if should_index:
            try:
                from services.pinecone_service import pinecone_service
                from services.chunking_service import chunking_service
                from services.embedding_service import embedding_service
                
                print(f"🌲 [PINECONE] Indexing mandatory file {file_id} ({mandatory_file.file_name}) to Pinecone...")
                
                # Check if file has extracted text
                if not mandatory_file.extracted_text:
                    print(f"⚠️ [PINECONE] File {file_id} has no extracted_text, skipping Pinecone indexing")
                else:
                    # Create index for this file
                    index_result = pinecone_service.create_index_for_file(
                        file_id=file_id,
                        file_name=mandatory_file.file_name
                    )
                    
                    if not index_result.get("success"):
                        print(f"⚠️ [PINECONE] Failed to create index: {index_result.get('error')}")
                    else:
                        # Chunk text (400 chars, 100 overlap)
                        chunks = chunking_service.chunk_text_by_characters(
                            text=mandatory_file.extracted_text,
                            chunk_size=400,
                            chunk_overlap=100,
                            metadata={
                                "file_id": file_id,
                                "file_name": mandatory_file.file_name,
                                "file_type": mandatory_file.file_type or "unknown"
                            }
                        )
                        
                        if chunks:
                            # Generate embeddings
                            chunk_texts = [chunk["text"] for chunk in chunks]
                            embeddings = embedding_service.embed(chunk_texts)
                            
                            # Index to Pinecone
                            index_chunks_result = pinecone_service.index_file_chunks(
                                file_id=file_id,
                                file_name=mandatory_file.file_name,
                                chunks=chunks,
                                embeddings=embeddings
                            )
                            
                            if index_chunks_result.get("success"):
                                print(f"✅ [PINECONE] Successfully indexed {index_chunks_result.get('chunks_indexed', 0)} chunks for file {file_id}")
                            else:
                                print(f"⚠️ [PINECONE] Failed to index chunks: {index_chunks_result.get('error')}")
                        else:
                            print(f"⚠️ [PINECONE] No chunks created for file {file_id}")
            except Exception as e:
                print(f"⚠️ [PINECONE] Error during Pinecone indexing: {str(e)}")
                import traceback
                print(traceback.format_exc())
                # Don't fail the request if Pinecone indexing fails
        
        return {
            "success": True,
            "message": f"File '{mandatory_file.file_name}' added to knowledge base" + (" (reindexed)" if existing else ""),
            "file_id": file_id
        }
        
    except Exception as e:
        db.rollback()
        return {
            "success": False,
            "error": f"Error adding file to knowledge base: {str(e)}"
        }

@app.delete("/api/project-knowledge-base/remove")
async def remove_project_knowledge_base_file(
    file_id: int,
    user_email: str,
    db: Session = Depends(get_db)
):
    """Remove a file from user's project knowledge base"""
    try:
        knowledge_base_file = db.query(ProjectKnowledgeBaseFile).filter(
            ProjectKnowledgeBaseFile.user_email == user_email,
            ProjectKnowledgeBaseFile.mandatory_file_id == file_id
        ).first()
        
        if not knowledge_base_file:
            return {
                "success": False,
                "error": "File not found in knowledge base"
            }
        
        # Get file info before deletion
        mandatory_file = db.query(MandatoryFile).filter(
            MandatoryFile.id == file_id
        ).first()
        
        db.delete(knowledge_base_file)
        db.commit()
        
        # Delete Pinecone index for this file
        if mandatory_file:
            try:
                from services.pinecone_service import pinecone_service
                delete_result = pinecone_service.delete_index(
                    file_id=file_id,
                    file_name=mandatory_file.file_name
                )
                if delete_result.get("success"):
                    print(f"✅ [PINECONE] Deleted index for file {file_id} ({mandatory_file.file_name})")
                else:
                    print(f"⚠️ [PINECONE] Failed to delete index: {delete_result.get('error')}")
            except Exception as e:
                print(f"⚠️ [PINECONE] Error deleting index: {str(e)}")
        
        return {
            "success": True,
            "message": "File removed from knowledge base",
            "file_id": file_id
        }
        
    except Exception as e:
        db.rollback()
        return {
            "success": False,
            "error": f"Error removing file from knowledge base: {str(e)}"
        }

@app.get("/api/project-knowledge-base")
async def get_project_knowledge_base(user_email: str, db: Session = Depends(get_db)):
    """Get all files in user's project knowledge base"""
    try:
        knowledge_base_files = db.query(ProjectKnowledgeBaseFile).filter(
            ProjectKnowledgeBaseFile.user_email == user_email
        ).all()
        
        file_ids = [kb_file.mandatory_file_id for kb_file in knowledge_base_files]
        
        # Get file details
        files = db.query(MandatoryFile).filter(
            MandatoryFile.id.in_(file_ids),
            MandatoryFile.is_active == True
        ).all()
        
        return {
            "success": True,
            "file_ids": file_ids,
            "files": [
                {
                    "id": f.id,
                    "file_name": f.file_name,
                    "file_type": f.file_type,
                    "file_size": f.file_size,
                    "uploaded_by": f.uploaded_by,
                    "uploaded_at": f.uploaded_at.isoformat() if f.uploaded_at else None
                }
                for f in files
            ]
        }
        
    except Exception as e:
        return {
            "success": False,
            "error": f"Error fetching knowledge base: {str(e)}"
        }

@app.post("/api/project-knowledge-base/reindex-all")
async def reindex_all_knowledge_base_files(db: Session = Depends(get_db)):
    """
    Reindex all files in the knowledge base to Pinecone.
    Useful for indexing files that were added before Pinecone integration.
    """
    try:
        from services.pinecone_service import pinecone_service
        from services.chunking_service import chunking_service
        from services.embedding_service import embedding_service
        
        # Get all knowledge base files
        knowledge_base_files = db.query(ProjectKnowledgeBaseFile).all()
        
        if not knowledge_base_files:
            return {
                "success": False,
                "error": "No files in knowledge base to reindex"
            }
        
        # Get unique file IDs
        file_ids = list(set([kb_file.mandatory_file_id for kb_file in knowledge_base_files]))
        
        # Get mandatory files
        mandatory_files = db.query(MandatoryFile).filter(
            MandatoryFile.id.in_(file_ids),
            MandatoryFile.is_active == True,
            MandatoryFile.extracted_text.isnot(None),
            MandatoryFile.extracted_text != ""
        ).all()
        
        results = []
        success_count = 0
        error_count = 0
        
        for mandatory_file in mandatory_files:
            try:
                print(f"🌲 [REINDEX] Indexing file {mandatory_file.id} ({mandatory_file.file_name})...")
                
                # Create index
                index_result = pinecone_service.create_index_for_file(
                    file_id=mandatory_file.id,
                    file_name=mandatory_file.file_name
                )
                
                if not index_result.get("success"):
                    error_count += 1
                    results.append({
                        "file_id": mandatory_file.id,
                        "file_name": mandatory_file.file_name,
                        "success": False,
                        "error": index_result.get("error", "Failed to create index")
                    })
                    continue
                
                # Chunk text
                chunks = chunking_service.chunk_text_by_characters(
                    text=mandatory_file.extracted_text,
                    chunk_size=400,
                    chunk_overlap=100,
                    metadata={
                        "file_id": mandatory_file.id,
                        "file_name": mandatory_file.file_name,
                        "file_type": mandatory_file.file_type or "unknown"
                    }
                )
                
                if not chunks:
                    error_count += 1
                    results.append({
                        "file_id": mandatory_file.id,
                        "file_name": mandatory_file.file_name,
                        "success": False,
                        "error": "No chunks created"
                    })
                    continue
                
                # Generate embeddings
                chunk_texts = [chunk["text"] for chunk in chunks]
                embeddings = embedding_service.embed(chunk_texts)
                
                # Index to Pinecone
                index_chunks_result = pinecone_service.index_file_chunks(
                    file_id=mandatory_file.id,
                    file_name=mandatory_file.file_name,
                    chunks=chunks,
                    embeddings=embeddings
                )
                
                if index_chunks_result.get("success"):
                    success_count += 1
                    results.append({
                        "file_id": mandatory_file.id,
                        "file_name": mandatory_file.file_name,
                        "success": True,
                        "chunks_indexed": index_chunks_result.get("chunks_indexed", 0),
                        "index_name": index_result.get("index_name")
                    })
                    print(f"✅ [REINDEX] Successfully indexed {index_chunks_result.get('chunks_indexed', 0)} chunks for file {mandatory_file.id}")
                else:
                    error_count += 1
                    results.append({
                        "file_id": mandatory_file.id,
                        "file_name": mandatory_file.file_name,
                        "success": False,
                        "error": index_chunks_result.get("error", "Failed to index chunks")
                    })
                    
            except Exception as e:
                error_count += 1
                results.append({
                    "file_id": mandatory_file.id,
                    "file_name": mandatory_file.file_name,
                    "success": False,
                    "error": str(e)
                })
                print(f"⚠️ [REINDEX] Error indexing file {mandatory_file.id}: {str(e)}")
        
        return {
            "success": True,
            "total_files": len(mandatory_files),
            "success_count": success_count,
            "error_count": error_count,
            "results": results
        }
        
    except Exception as e:
        import traceback
        print(f"❌ [REINDEX] Error: {str(e)}")
        print(traceback.format_exc())
        return {
            "success": False,
            "error": f"Error reindexing files: {str(e)}"
        }

# Authentication endpoints
@app.get("/api/auth/google/url")
async def get_google_auth_url(prompt: str = None):
    """Get Google OAuth URL"""
    try:
        print(f"🔐 [AUTH] Generating Google OAuth URL...")
        print(f"🔐 [AUTH] GOOGLE_CLIENT_ID: {'Set' if os.getenv('GOOGLE_CLIENT_ID') else 'NOT SET'}")
        print(f"🔐 [AUTH] GOOGLE_REDIRECT_URI: {os.getenv('GOOGLE_REDIRECT_URI', 'NOT SET')}")
        auth_url = auth_service.get_google_auth_url(prompt=prompt)
        print(f"🔐 [AUTH] Generated auth URL: {auth_url[:100]}...")
        return {"auth_url": auth_url}
    except ValueError as e:
        print(f"❌ [AUTH] ValueError: {str(e)}")
        return {"error": str(e), "auth_url": None}
    except Exception as e:
        print(f"❌ [AUTH] Exception: {str(e)}")
        import traceback
        print(traceback.format_exc())
        return {"error": f"Failed to generate Google OAuth URL: {str(e)}", "auth_url": None}

from pydantic import BaseModel

class GoogleCallbackRequest(BaseModel):
    code: str

class EditSprintPlanRequest(BaseModel):
    sprint_overview: dict
    team_capacity: dict
    product_backlog: dict
    definition_of_done: dict
    risks_and_impediments: dict
    additional_comments: dict
    edit_comments: str
    edited_by: str

def generate_word_document_content(generated_plan: str, user_inputs: dict) -> str:
    """Generate Word document content that is a replica of HTML rendered output"""
    try:
        # Remove markdown code blocks but keep HTML formatting
        html_content = generated_plan or ''
        import re
        html_content = re.sub(r'```html\s*', '', html_content, flags=re.IGNORECASE)
        html_content = re.sub(r'```\s*$', '', html_content, flags=re.IGNORECASE)
        html_content = html_content.strip()
        
        if not html_content:
            return '<p>No content available for Word document.</p>'
        
        # Create Word-compatible HTML with all formatting preserved
        word_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>Sprint Plan - {user_inputs.get('sprint_overview', {}).get('SprintNumber', 'N/A')}</title>
            <style>
                body {{
                    font-family: 'Calibri', 'Arial', sans-serif;
                    margin: 20px;
                    padding: 0;
                    background: white;
                    color: #2c3e50;
                    line-height: 1.6;
                    font-size: 11pt;
                }}
                h1, h2, h3, h4 {{
                    color: #2d3748;
                    margin: 20px 0 15px 0;
                    font-weight: 600;
                }}
                h1 {{ font-size: 18pt; }}
                h2 {{ font-size: 16pt; }}
                h3 {{ font-size: 14pt; }}
                h4 {{ font-size: 12pt; }}
                p {{ margin: 10px 0; }}
                table {{
                    width: 100%;
                    border-collapse: collapse;
                    margin: 20px 0;
                    border: 1px solid #ddd;
                }}
                th, td {{
                    border: 1px solid #ddd;
                    padding: 8px 12px;
                    text-align: left;
                    vertical-align: top;
                }}
                th {{
                    background-color: #f8f9fa;
                    font-weight: bold;
                    color: #2d3748;
                }}
                ul, ol {{ 
                    margin: 15px 0; 
                    padding-left: 30px; 
                }}
                li {{ 
                    margin: 8px 0; 
                    line-height: 1.6;
                }}
                strong {{
                    color: #2d3748;
                    font-weight: 600;
                }}
                em {{
                    color: #718096;
                    font-style: italic;
                }}
                .sprint-header {{
                    background: #f7fafc;
                    border: 2px solid #e2e8f0;
                    border-radius: 8px;
                    padding: 20px;
                    margin-bottom: 20px;
                    text-align: center;
                }}
                .sprint-header h1 {{
                    margin: 0 0 10px 0;
                    color: #1a202c;
                }}
                .sprint-header p {{
                    margin: 0;
                    color: #4a5568;
                    font-size: 12pt;
                }}
            </style>
        </head>
        <body>
            <div class="sprint-header">
                <h1>Sprint Plan - Sprint {user_inputs.get('sprint_overview', {}).get('SprintNumber', 'N/A')}</h1>
                <p>Generated on {__import__('datetime').datetime.now().strftime('%Y-%m-%d')}</p>
            </div>
            
            {html_content}
            
            <div style="margin-top: 40px; padding-top: 20px; border-top: 1px solid #e2e8f0; text-align: center; color: #718096; font-size: 10pt;">
                <p>This sprint plan was automatically generated by the Sprint Planning System</p>
                <p>Document ID: SP-{user_inputs.get('sprint_overview', {}).get('SprintNumber', 'N/A')}-{__import__('datetime').datetime.now().strftime('%Y-%m-%d')}</p>
            </div>
        </body>
        </html>
        """
        
        return word_html.strip()
        
    except Exception as e:
        print(f"❌ [WORD GENERATION] Error generating Word document: {str(e)}")
        return f'<p>Error generating Word document: {str(e)}</p>'

class DownloadVersionRequest(BaseModel):
    version_number: int

@app.post("/api/auth/google/callback")
async def google_auth_callback(request: GoogleCallbackRequest, db: Session = Depends(get_db)):
    """Handle Google OAuth callback"""
    try:
        print(f"🔐 [AUTH] Received Google callback with code: {request.code[:20] if request.code else 'None'}...")
        result = auth_service.authenticate_user(request.code, db)
        if result.success:
            print(f"✅ [AUTH] Authentication successful for user: {result.user.get('email', 'Unknown') if result.user else 'None'}")
        else:
            print(f"❌ [AUTH] Authentication failed: {result.message}")
        return result
    except Exception as e:
        print(f"❌ [AUTH] Callback error: {str(e)}")
        import traceback
        print(traceback.format_exc())
        from schemas import LoginResponse
        return LoginResponse(
            success=False,
            session_id="",
            user=None,
            message=f"Authentication error: {str(e)}"
        )

@app.post("/api/auth/login", response_model=LoginResponse)
async def login(request: LoginRequest):
    """Simulate Google OAuth login (for demo)"""
    return auth_service.simulate_login(request)

@app.post("/api/auth/logout")
async def logout():
    """Simulate logout"""
    return auth_service.simulate_logout()

@app.post("/api/upload/docx")
async def upload_docx_file(file: UploadFile = File(...), feature_type: str = "sprint"):
    """Upload and parse DOCX file for data extraction"""
    try:
        # Validate file type
        if not file.filename.lower().endswith('.docx'):
            return {
                "success": False,
                "error": "Invalid file type. Please upload a .docx file."
            }
        
        # Read file content
        file_content = await file.read()
        
        # Use appropriate service based on feature type
        if feature_type == "risk-assessment":
            result = risk_docx_service.parse_docx_file(file_content)
        else:
            result = docx_service.parse_docx_file(file_content)
        
        if result['success']:
            print(f"🔍 [MAIN] Sending data to frontend: {result['data']}")
            return {
                "success": True,
                "data": result['data'],
                "message": result['message']
            }
        else:
            return {
                "success": False,
                "error": result['message']
            }
            
    except Exception as e:
        return {
            "success": False,
            "error": f"Error processing file: {str(e)}"
        }

# SOW upload endpoint: supports both DOCX and PDF files; prints content to console and returns raw text
@app.post("/api/upload/sow")
async def upload_sow_document(file: UploadFile = File(...), user_email: str = Form(None)):
    """Upload SOW document (DOCX or PDF), print raw text to server console, return raw text back."""
    try:
        file_extension = file.filename.lower().split('.')[-1]
        
        if file_extension not in ['docx', 'pdf']:
            return {"success": False, "error": "Invalid file type. Please upload a .docx or .pdf file."}

        file_content = await file.read()
        raw_text = ""

        if file_extension == 'docx':
            # Convert DOCX to HTML preserving formatting
            from docx import Document
            import io
            import re

            doc = Document(io.BytesIO(file_content))
            html_parts = []
            
            # Process paragraphs
            for p in doc.paragraphs:
                if p.text.strip():
                    # Check if paragraph has heading style
                    if p.style.name.startswith('Heading'):
                        level = p.style.name.replace('Heading ', '')
                        html_parts.append(f'<h{level}>{p.text.strip()}</h{level}>')
                    else:
                        # Check for bold/italic runs
                        para_html = ""
                        for run in p.runs:
                            text = run.text
                            if run.bold:
                                text = f"<strong>{text}</strong>"
                            if run.italic:
                                text = f"<em>{text}</em>"
                            para_html += text
                        
                        if para_html.strip():
                            html_parts.append(f'<p>{para_html}</p>')
            
            # Process tables
            for table in doc.tables:
                html_parts.append('<table border="1" style="border-collapse: collapse; width: 100%; margin: 10px 0;">')
                for row_idx, row in enumerate(table.rows):
                    html_parts.append('<tr>')
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            # Check if it's a header row (usually first row)
                            tag = 'th' if row_idx == 0 else 'td'
                            html_parts.append(f'<{tag} style="padding: 8px; border: 1px solid #ddd;">{cell_text}</{tag}>')
                        else:
                            tag = 'th' if row_idx == 0 else 'td'
                            html_parts.append(f'<{tag} style="padding: 8px; border: 1px solid #ddd;">&nbsp;</{tag}>')
                    html_parts.append('</tr>')
                html_parts.append('</table>')
            
            # Join HTML parts
            html_content = '\n'.join(html_parts)
            
            # Also create raw text version for backward compatibility
            raw_text = "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
            
        elif file_extension == 'pdf':
            # Extract raw text from PDF using our PDF service
            pdf_result = pdf_service.extract_text_from_pdf(file_content)
            
            if not pdf_result['success']:
                return {"success": False, "error": f"Error extracting text from PDF: {pdf_result['error']}"}
            
            raw_text = pdf_result['text'].strip()
            print(f"📄 [PDF EXTRACTION] Used method: {pdf_result.get('method', 'Unknown')}")

        if not raw_text:
            return {"success": False, "error": "No text content found in the document."}

        # Print to terminal/console
        print("=" * 80)
        print(f"📄 [SOW UPLOAD] File: {file.filename} ({file_extension.upper()})")
        print("— Raw Document Text —")
        print(raw_text if raw_text else "<No text content>")
        print("=" * 80)

        # Return both HTML and raw text
        response_data = {
            "rawText": raw_text, 
            "fileName": file.filename, 
            "fileType": file_extension
        }
        
        # Add HTML content for DOCX files
        if file_extension == 'docx' and 'html_content' in locals():
            response_data["htmlContent"] = html_content
        
        return {"success": True, "data": response_data}
    except Exception as e:
        return {"success": False, "error": f"Error processing SOW file: {str(e)}"}

# Sprint planning endpoints
@app.post("/api/sprint/start", response_model=SprintStartResponse)
async def start_sprint_planning(
    request: SprintStartRequest,
    db: Session = Depends(get_db)
):
    """Start a new sprint planning session"""
    return sprint_service.start_sprint_planning(request, db)

@app.post("/api/sprint/chat", response_model=ChatResponse)
async def chat_with_llm(
    request: ChatRequest,
    db: Session = Depends(get_db)
):
    """Send message to LLM and get response"""
    return sprint_service.chat_with_llm(request, db, get_global_prompt_data())

@app.get("/api/sprint/prompt")
async def get_sprint_prompt(db: Session = Depends(get_db)):
    """Fetch prompt data from documents table for sprint planning"""
    global GLOBAL_PROMPT_DATA
    try:
        # Print the SQL query being executed
        print("🔍 [PROMPT FETCH] SQL Query: SELECT * FROM documents WHERE feature = 'SprintPlan' LIMIT 1")
        
        # Use raw SQL instead of ORM to avoid importing Document
        result = db.execute(text("SELECT id, prompt, feature FROM documents WHERE feature = 'SprintPlan' LIMIT 1"))
        row = result.fetchone()
        
        if row:
            # Store the prompt data in the global variable in main flow
            GLOBAL_PROMPT_DATA = row.prompt
            
            # Print variable details
            print(f"📦 [PROMPT FETCH] Variable Name: GLOBAL_PROMPT_DATA")
            print(f"📦 [PROMPT FETCH] Variable Type: {type(GLOBAL_PROMPT_DATA)}")
            print(f"📦 [PROMPT FETCH] Variable Length: {len(GLOBAL_PROMPT_DATA) if GLOBAL_PROMPT_DATA else 0} characters")
            print(f"📦 [PROMPT FETCH] Variable Content (first 200 chars): {GLOBAL_PROMPT_DATA[:200] if GLOBAL_PROMPT_DATA else 'None'}...")
            print(f"📦 [PROMPT FETCH] Document ID: {row.id}")
            print(f"📦 [PROMPT FETCH] Document Feature: {row.feature}")
            
            return {
                "success": True,
                "prompt": row.prompt,
                "feature": row.feature
            }
        else:
            print("❌ [PROMPT FETCH] No document found with feature = 'SprintPlan'")
            return {
                "success": False,
                "prompt": "",
                "feature": "",
                "message": "No prompt data found for SprintPlan"
            }
    except Exception as e:
        print(f"❌ [PROMPT FETCH] Error: {str(e)}")
        return {
            "success": False,
            "prompt": "",
            "feature": "",
            "message": f"Error fetching prompt data: {str(e)}"
        }

@app.get("/api/feature/prompt/{feature_name}")
async def get_feature_prompt(feature_name: str, db: Session = Depends(get_db)):
    """Fetch prompt data from documents table for specific features"""
    try:
        # Map feature names to database feature values
        feature_mapping = {
            "weekly-status-report": "WeeklyStatusReport",
            "risk-assessment": "RiskAssessment"
        }
        
        db_feature_name = feature_mapping.get(feature_name, feature_name)
        
        # Print the SQL query being executed
        print(f"🔍 [FEATURE PROMPT FETCH] SQL Query: SELECT * FROM documents WHERE feature = '{db_feature_name}' LIMIT 1")
        
        # Use raw SQL instead of ORM to avoid importing Document
        result = db.execute(text(f"SELECT id, prompt, feature FROM documents WHERE feature = '{db_feature_name}' LIMIT 1"))
        row = result.fetchone()
        
        if row:
            # Print variable details
            print(f"📦 [FEATURE PROMPT FETCH] Successfully fetched prompt for feature: {feature_name}")
            print(f"📦 [FEATURE PROMPT FETCH] Document ID: {row.id}")
            print(f"📦 [FEATURE PROMPT FETCH] Document Feature: {row.feature}")
            print(f"📦 [FEATURE PROMPT FETCH] Prompt Length: {len(row.prompt) if row.prompt else 0} characters")
            print(f"📦 [FEATURE PROMPT FETCH] Prompt Content (first 200 chars): {row.prompt[:200] if row.prompt else 'None'}...")
            
            return {
                "success": True,
                "prompt": row.prompt,
                "feature": row.feature
            }
        else:
            print(f"❌ [FEATURE PROMPT FETCH] No document found with feature = '{db_feature_name}'")
            return {
                "success": False,
                "prompt": "",
                "feature": "",
                "message": f"No prompt data found for {feature_name}"
            }
    except Exception as e:
        print(f"❌ [FEATURE PROMPT FETCH] Error: {str(e)}")
        return {
            "success": False,
            "prompt": "",
            "feature": "",
            "message": f"Error fetching prompt data for {feature_name}: {str(e)}"
        }

# Finish sprint implementation
class QAItem(BaseModel):
    question: str
    answer: str

class CustomSprintFinishRequest(BaseModel):
    qa_list: List[QAItem]

@app.post("/api/sprint/finish")
async def finish_sprint_custom(request: CustomSprintFinishRequest, db: Session = Depends(get_db)):
    """Generate sprint plan summary based on custom Q&A list using Gemini and save to CSV"""
    try:
        from services.gemini_service import gemini_service
        from services.csv_service import csv_service
        from services.db_service import db_service
        
        # Use the stored prompt data from global variable in main flow
        print("🔍 [FINISH ENDPOINT] Getting stored prompt from global variable...")
        stored_prompt = get_global_prompt_data()
        
        if not stored_prompt:
            # Fallback to querying table if no stored prompt
            print("🔍 [FINISH ENDPOINT] No stored prompt found, querying database...")
            from sqlalchemy import text
            fallback_query = "SELECT prompt FROM documents WHERE feature = 'SprintPlan'"
            print(f"🔍 [FINISH ENDPOINT] Fallback SQL Query: {fallback_query}")
            result = db.execute(text(fallback_query))
            row = result.fetchone()
            if not row:
                print("❌ [FINISH ENDPOINT] No prompt found in database")
                return {"success": False, "message": "Reference prompt not found in database"}
            stored_prompt = row[0]
            print(f"📦 [FINISH ENDPOINT] Retrieved prompt from database: {stored_prompt[:100]}...")
            # Store it in global variable for future use
            GLOBAL_PROMPT_DATA = stored_prompt
        else:
            print(f"📦 [FINISH ENDPOINT] Using stored prompt from global variable: {stored_prompt[:100]}...")
        
        # Add the Q&A data to the stored prompt structure
        qa_data = ""
        for item in request.qa_list:
            qa_data += f"{item.question}: {item.answer}\n"
        
        # Combine the stored prompt with the Q&A data
        full_prompt = f"{stored_prompt}\n\nSprint Planning Information:\n{qa_data}"
        
        # Use Gemini to generate the summary
        messages = [
            {"role": "system", "content": stored_prompt},  # Use the global prompt data
            {"role": "user", "content": full_prompt}
        ]
        
        result = gemini_service.chat(messages, max_tokens=3000)
        
        if result["success"]:
            # Convert Q&A list to user_inputs format for CSV and database
            user_inputs = {}
            
            # Define the mapping from form field names to CSV column names
            field_mapping = {
                'sprint number': 'sprint_number',
                'sprint dates': 'sprint_dates', 
                'sprint duration': 'sprint_duration',
                'team name': 'team_name',
                'sprint goal': 'sprint_goal',
                'working hour per person': 'working_hours_per_person',
                'team availability': 'team_member_availability',
                'story points history': 'historical_story_points',
                'pbi summary': 'pbi_summary',
                'pbi criteria': 'pbi_criteria',
                'pbi priority': 'pbi_priority',
                'pbi effort': 'pbi_effort',
                'definition of done': 'definition_of_done',
                'risk dependencies': 'risks_and_dependencies'
            }
            
            for item in request.qa_list:
                # Map the question names to CSV column names
                question_lower = item.question.lower()
                csv_column = field_mapping.get(question_lower, question_lower.replace(" ", "_"))
                user_inputs[csv_column] = item.answer
            
            # Save to CSV
            csv_result = csv_service.save_sprint_plan(user_inputs)
            
            # Save to database
            db_result = db_service.save_sprint_plan(db, user_inputs)
            
            return {
                "success": True, 
                "summary": result["response"],
                "csv_saved": csv_result["success"],
                "csv_message": csv_result.get("message", ""),
                "db_saved": db_result["success"],
                "db_message": db_result.get("message", ""),
                "plan_id": db_result.get("plan_id")
            }
        else:
            return {"success": False, "message": result["response"]}

    except Exception as e:
        return {"success": False, "message": f"Failed to generate summary: {str(e)}"}

# New endpoint for generating plans with the new structure
class GeneratePlanRequest(BaseModel):
    sprint_overview: dict
    team_capacity: dict
    product_backlog: dict
    definition_of_done: dict
    risks_and_impediments: dict
    additional_comments: dict
    sow_content: str | None = None
    current_plan_content: str | None = None  # Include current plan for regeneration
    regenerate_with_current_plan: bool = False  # Flag to indicate regeneration
    user_email: str  # Add user email to track who created the plan
    workspace_id: int | None = None  # Add workspace_id to track workspace

# Update SOW content for existing sprint plan
@app.post("/api/sprint/validate-plan")
async def validate_sprint_plan(request: dict, db: Session = Depends(get_db)):
    """Validate sprint plan against SOW using Planvalidation prompt"""
    try:
        sprint_plan_content = request.get('sprint_plan_content', '')
        sow_content = request.get('sow_content', '')
        
        if not sprint_plan_content or not sow_content:
            return {"success": False, "error": "Missing required fields: sprint_plan_content and sow_content"}
        
        print("=" * 80)
        print("✅ [VALIDATE PLAN] REQUEST RECEIVED")
        print("=" * 80)
        print(f"📋 [VALIDATE PLAN] Sprint Plan Content Length: {len(sprint_plan_content)}")
        print(f"📋 [VALIDATE PLAN] SOW Content Length: {len(sow_content)}")
        print("=" * 80)
        
        from services.gemini_service import gemini_service
        
        # Get Planvalidation prompt from database
        result = db.execute(text("SELECT prompt FROM documents WHERE feature = 'Planvalidation' LIMIT 1"))
        row = result.fetchone()
        
        if not row:
            print("❌ [VALIDATE PLAN] No Planvalidation prompt found in database!")
            return {"success": False, "error": "Planvalidation prompt not found in database"}
        
        validation_prompt = row[0]
        print(f"🔍 [VALIDATE PLAN] Using validation prompt: {len(validation_prompt)} characters")
        print(f"🔍 [VALIDATE PLAN] Prompt preview: {validation_prompt[:200]}...")
        
        # Prepare the validation request
        validation_data = f"""
SPRINT PLAN TO VALIDATE:
{sprint_plan_content[:15000]}  # Limit to avoid token limits

STATEMENT OF WORK (SOW) FOR COMPARISON:
{sow_content[:15000]}  # Limit to avoid token limits

Please validate the sprint plan against the SOW and provide:
1. Percentage alignment (e.g., "85% aligned")
2. Areas that are out of scope
3. Technical boundary compliance
4. Resource allocation alignment
"""
        
        print(f"🔍 [VALIDATE PLAN] Validation data length: {len(validation_data)}")
        
        # Call Gemini service with validation prompt
        messages = [
            {"role": "system", "content": validation_prompt},
            {"role": "user", "content": validation_data}
        ]
        
        print("🔍 [GEMINI CALL] Calling Gemini service for validation...")
        gemini_response = gemini_service.chat(messages, max_tokens=2000)
        
        if not gemini_response or not gemini_response.get('success', False):
            error_msg = gemini_response.get('response', 'Unknown error from Gemini service') if gemini_response else 'No response from Gemini'
            print(f"❌ [VALIDATE PLAN] Gemini service failed: {error_msg}")
            return {"success": False, "error": f"Gemini validation failed: {error_msg}"}
        
        validation_result = gemini_response.get('response', '')
        print(f"✅ [VALIDATE PLAN] Validation completed successfully!")
        print(f"📊 [VALIDATE PLAN] Result length: {len(validation_result)} characters")
        print(f"📊 [VALIDATE PLAN] Result preview: {validation_result[:200]}...")
        
        return {
            "success": True,
            "response": validation_result,
            "message": "Sprint plan validation completed successfully"
        }
        
    except Exception as e:
        print(f"❌ [VALIDATE PLAN] Error: {str(e)}")
        return {"success": False, "error": f"Error validating sprint plan: {str(e)}"}

@app.post("/api/risk/validate-assessment")
async def validate_risk_assessment(request: dict, db: Session = Depends(get_db)):
    """Validate risk assessment against SOW using Riskvalidation prompt"""
    try:
        risk_assessment_content = request.get('risk_assessment_content', '')
        sow_content = request.get('sow_content', '')
        
        if not risk_assessment_content or not sow_content:
            return {"success": False, "error": "Missing required fields: risk_assessment_content and sow_content"}
        
        print("=" * 80)
        print("✅ [VALIDATE RISK ASSESSMENT] REQUEST RECEIVED")
        print("=" * 80)
        print(f"📋 [VALIDATE RISK ASSESSMENT] Risk Assessment Content Length: {len(risk_assessment_content)}")
        print(f"📋 [VALIDATE RISK ASSESSMENT] SOW Content Length: {len(sow_content)}")
        print("=" * 80)
        
        from services.gemini_service import gemini_service
        
        # Get Riskvalidation prompt from database
        result = db.execute(text("SELECT prompt FROM documents WHERE feature = 'Riskvalidation' LIMIT 1"))
        row = result.fetchone()
        
        if not row:
            print("❌ [VALIDATE RISK ASSESSMENT] No Riskvalidation prompt found in database!")
            return {"success": False, "error": "Riskvalidation prompt not found in database"}
        
        validation_prompt = row[0]
        print(f"🔍 [VALIDATE RISK ASSESSMENT] Using validation prompt: {len(validation_prompt)} characters")
        print(f"🔍 [VALIDATE RISK ASSESSMENT] Prompt preview: {validation_prompt[:200]}...")
        
        # Prepare the validation request
        validation_data = f"""
RISK ASSESSMENT TO VALIDATE:
{risk_assessment_content[:15000]}  # Limit to avoid token limits

STATEMENT OF WORK (SOW) FOR COMPARISON:
{sow_content[:15000]}  # Limit to avoid token limits

Please validate the risk assessment against the SOW and provide:
1. Percentage alignment (e.g., "85% aligned")
2. Areas that are out of scope
3. Risk coverage compliance
4. Regulatory alignment
"""
        
        print(f"🔍 [VALIDATE RISK ASSESSMENT] Validation data length: {len(validation_data)}")
        
        # Call Gemini service with validation prompt
        messages = [
            {"role": "system", "content": validation_prompt},
            {"role": "user", "content": validation_data}
        ]
        
        gemini_response = gemini_service.chat(messages, max_tokens=2000)
        
        if not gemini_response or not gemini_response.get('success', False):
            error_msg = gemini_response.get('response', 'Unknown error from Gemini service') if gemini_response else 'No response from Gemini'
            print(f"❌ [VALIDATE RISK ASSESSMENT] Gemini service failed: {error_msg}")
            return {"success": False, "error": f"Gemini validation failed: {error_msg}"}
        
        validation_result = gemini_response.get('response', '')
        print(f"✅ [VALIDATE RISK ASSESSMENT] Validation completed successfully!")
        print(f"📊 [VALIDATE RISK ASSESSMENT] Result length: {len(validation_result)} characters")
        print(f"📊 [VALIDATE RISK ASSESSMENT] Result preview: {validation_result[:200]}...")
        
        return {
            "success": True,
            "response": validation_result,
            "message": "Risk assessment validation completed successfully"
        }
        
    except Exception as e:
        print(f"❌ [VALIDATE RISK ASSESSMENT] Error: {str(e)}")
        return {"success": False, "error": f"Error validating risk assessment: {str(e)}"}

@app.post("/api/sprint/update-sow")
async def update_sow_content(request: dict, db: Session = Depends(get_db)):
    """Update SOW content for an existing sprint plan"""
    try:
        sprint_plan_id = request.get('sprint_plan_id')
        sow_content = request.get('sow_content')
        sow_file_name = request.get('sow_file_name', 'Updated SOW')
        
        if not sprint_plan_id or not sow_content:
            return {"success": False, "error": "Missing required fields: sprint_plan_id and sow_content"}
        
        # Find the sprint plan
        from models import SprintPlan
        sprint_plan = db.query(SprintPlan).filter(SprintPlan.id == sprint_plan_id).first()
        
        if not sprint_plan:
            return {"success": False, "error": "Sprint plan not found"}
        
        # Update the SOW content
        sprint_plan.sow_content = sow_content
        db.commit()
        
        print(f"✅ [SOW UPDATE] Updated SOW content for sprint plan {sprint_plan_id}")
        print(f"📄 [SOW UPDATE] File: {sow_file_name}")
        print(f"📄 [SOW UPDATE] Content length: {len(sow_content)}")
        
        return {"success": True, "message": "SOW content updated successfully"}
        
    except Exception as e:
        print(f"❌ [SOW UPDATE] Error: {str(e)}")
        return {"success": False, "error": f"Error updating SOW content: {str(e)}"}

@app.post("/api/sprint/generate-plan")
async def generate_sprint_plan(request: GeneratePlanRequest, db: Session = Depends(get_db)):
    """Generate sprint plan with new structure"""
    try:
        import json
        
        # Log the complete JSON request data
        print("=" * 80)
        print("🚀 [GENERATE SPRINT PLAN] REQUEST RECEIVED")
        print("=" * 80)
        print("📋 [GENERATE SPRINT PLAN] Complete JSON Request Data:")
        print("=" * 80)
        
        # Convert request to dict for logging
        request_dict = {
            "sprint_overview": request.sprint_overview,
            "team_capacity": request.team_capacity,
            "product_backlog": request.product_backlog,
            "definition_of_done": request.definition_of_done,
            "risks_and_impediments": request.risks_and_impediments,
            "additional_comments": request.additional_comments
        }
        
        # Pretty print the JSON
        print(json.dumps(request_dict, indent=2, ensure_ascii=False))
        print("=" * 80)
        print("📊 [GENERATE SPRINT PLAN] Request Summary:")
        print(f"   - Sprint Overview Keys: {list(request.sprint_overview.keys())}")
        print(f"   - Team Capacity Keys: {list(request.team_capacity.keys())}")
        print(f"   - Product Backlog Keys: {list(request.product_backlog.keys())}")
        print(f"   - Definition of Done Keys: {list(request.definition_of_done.keys())}")
        print(f"   - Risks & Impediments Keys: {list(request.risks_and_impediments.keys())}")
        print(f"   - Additional Comments Keys: {list(request.additional_comments.keys())}")
        
        # Log specific data points
        print("📝 [GENERATE SPRINT PLAN] Key Data Points:")
        print(f"   - Sprint Number: {request.sprint_overview.get('SprintNumber', 'N/A')}")
        print(f"   - Team Name: {request.sprint_overview.get('TeamName', 'N/A')}")
        print(f"   - Number of Team Members: {request.team_capacity.get('NumberOfMembers', 'N/A')}")
        print(f"   - Number of Backlog Items: {len(request.product_backlog.get('BacklogItems', []))}")
        print(f"   - DoD Content Length: {len(request.definition_of_done.get('DoDContent', ''))}")
        print(f"   - Risks Content Length: {len(request.risks_and_impediments.get('RisksContent', ''))}")
        print(f"   - Additional Comments Length: {len(request.additional_comments.get('CommentsContent', ''))}")
        print("=" * 80)
        
        from services.gemini_service import gemini_service
        from services.db_service import db_service
        
        # Get stored prompt from global variable or fetch from DB
        stored_prompt = get_global_prompt_data()
        if not stored_prompt:
            # Fallback: fetch from database
            print("🔍 [PROMPT LOADING] Global prompt is None, fetching from database...")
            result = db.execute(text("SELECT id, prompt, feature FROM documents WHERE feature = 'SprintPlan' LIMIT 1"))
            row = result.fetchone()
            if row:
                stored_prompt = row[1]
                # Update global variable for future use
                global GLOBAL_PROMPT_DATA
                GLOBAL_PROMPT_DATA = stored_prompt
                print(f"🔍 [PROMPT LOADING] Loaded prompt from database: {len(stored_prompt)} characters")
            else:
                print("❌ [PROMPT LOADING] No prompt found in database!")
                return {"success": False, "message": "No prompt found in database"}
        else:
            print(f"🔍 [PROMPT LOADING] Using global prompt: {len(stored_prompt)} characters")
        
        print(f"🔍 [GENERATE PLAN] Final prompt length: {len(stored_prompt)} characters")
        print(f"🔍 [GENERATE PLAN] Prompt preview: {stored_prompt[:200]}...")
        
        # Validate that we have a proper prompt
        if not stored_prompt or len(stored_prompt.strip()) < 50:
            print("❌ [PROMPT VALIDATION] Prompt is too short or empty!")
            return {"success": False, "message": "Invalid prompt data. Please ensure prompt is properly loaded."}
        
        # Format user inputs into prompt template
        user_inputs = {
            "sprint_overview": request.sprint_overview,
            "team_capacity": request.team_capacity,
            "product_backlog": request.product_backlog,
            "definition_of_done": request.definition_of_done,
            "risks_and_impediments": request.risks_and_impediments,
            "additional_comments": request.additional_comments
        }
        
        # Format user inputs properly
        team_members_text = ""
        if user_inputs['team_capacity'].get('TeamMembers'):
            team_members_text = "\n".join([
                f"- {member.get('roleName', 'N/A')}: {member.get('workingHours', 'N/A')}"
                for member in user_inputs['team_capacity'].get('TeamMembers', [])
            ])

        backlog_items_text = ""
        if user_inputs['product_backlog'].get('BacklogItems'):
            backlog_items_text = "\n".join([
                f"- {item.get('userStorySummary', 'N/A')} (Priority: {item.get('priority', 'N/A')}, Effort: {item.get('effortEstimate', 0)} hours)"
                for item in user_inputs['product_backlog'].get('BacklogItems', [])
            ])

        # Include optional SOW content (trim to avoid excessive token usage)
        sow_text = (request.sow_content or "").strip()
        sow_text_trimmed = sow_text[:8000]  # cap to ~8k chars to be safe
        
        # Include current plan content for regeneration (trim to avoid excessive token usage)
        current_plan_text = (request.current_plan_content or "").strip()
        current_plan_trimmed = current_plan_text[:12000]  # cap to ~12k chars for current plan
        
        # Check if this is a regeneration request
        is_regeneration = request.regenerate_with_current_plan and current_plan_text

        # Create user inputs section
        if is_regeneration:
            print("🔄 [REGENERATION] This is a regeneration request with current plan content")
            user_inputs_text = f"""
I. Sprint Overview & Proposed Goal:
- Sprint Number: {user_inputs['sprint_overview'].get('SprintNumber', 'N/A')}
- Sprint Dates: {user_inputs['sprint_overview'].get('SprintDates', 'N/A')}
- Sprint Duration: {user_inputs['sprint_overview'].get('SprintDuration', 'N/A')}
- Team Name: {user_inputs['sprint_overview'].get('TeamName', 'N/A')}
- Sprint Goal: {user_inputs['sprint_overview'].get('SprintGoal', 'N/A')}

II. Team Capacity & Availability:
- Total Hours per Person: {user_inputs['team_capacity'].get('TotalHoursPerPerson', 'N/A')}
- Number of Members: {user_inputs['team_capacity'].get('NumberOfMembers', 'N/A')}
- Team Members: {team_members_text}
- Historical Story Points: {user_inputs['team_capacity'].get('HistoricalStoryPoints', 'N/A')}

III. Prioritized Product Backlog Items:
{backlog_items_text}

IV. Definition of Done (DoD):
{user_inputs['definition_of_done'].get('DoDContent', 'N/A')}

V. Known Impediments, Dependencies & Risks:
{user_inputs['risks_and_impediments'].get('RisksContent', 'N/A')}

VI. Statement of Work (SOW) Context (verbatim text, truncated if long):
{sow_text_trimmed if sow_text_trimmed else 'N/A'}

VII. CURRENT SPRINT PLAN (for regeneration reference):
{current_plan_trimmed if current_plan_trimmed else 'N/A'}

IMPORTANT REGENERATION INSTRUCTIONS:
- This is a REGENERATION request. Please analyze the current sprint plan above and regenerate it based on the updated SOW content.
- Maintain the same structure and format as the current plan but incorporate any changes needed based on the new SOW.
- Focus on aligning the sprint plan with the updated SOW requirements and regulations.
- Keep the same level of detail and professional formatting as the current plan.
- DO NOT include the SOW content itself in your response - use it only as context for generating the sprint plan.
"""
        else:
            print("🆕 [NEW PLAN] This is a new plan generation request")
            user_inputs_text = f"""
I. Sprint Overview & Proposed Goal:
- Sprint Number: {user_inputs['sprint_overview'].get('SprintNumber', 'N/A')}
- Sprint Dates: {user_inputs['sprint_overview'].get('SprintDates', 'N/A')}
- Sprint Duration: {user_inputs['sprint_overview'].get('SprintDuration', 'N/A')}
- Team Name: {user_inputs['sprint_overview'].get('TeamName', 'N/A')}
- Sprint Goal: {user_inputs['sprint_overview'].get('SprintGoal', 'N/A')}

II. Team Capacity & Availability:
- Total Hours per Person: {user_inputs['team_capacity'].get('TotalHoursPerPerson', 'N/A')}
- Number of Members: {user_inputs['team_capacity'].get('NumberOfMembers', 'N/A')}
- Team Members: {team_members_text}
- Historical Story Points: {user_inputs['team_capacity'].get('HistoricalStoryPoints', 'N/A')}

III. Prioritized Product Backlog Items:
{backlog_items_text}

IV. Definition of Done (DoD):
{user_inputs['definition_of_done'].get('DoDContent', 'N/A')}

V. Known Impediments, Dependencies & Risks:
{user_inputs['risks_and_impediments'].get('RisksContent', 'N/A')}

VI. Statement of Work (SOW) Context (verbatim text, truncated if long):
{sow_text_trimmed if sow_text_trimmed else 'N/A'}

IMPORTANT INSTRUCTIONS:
- Use the SOW content above as context for generating the sprint plan.
- DO NOT include the SOW content itself in your response - use it only as reference for compliance and requirements.
- Generate a clean sprint plan without repeating the SOW regulations in the output.
"""

        print(f"🔍 [GENERATE PLAN] User inputs: {user_inputs_text[:200]}...")
        
        # Call Gemini service with proper message format and higher token limit
        messages = [
            {"role": "system", "content": stored_prompt},
            {"role": "user", "content": user_inputs_text}
        ]
        
        print("🔍 [GEMINI CALL] Calling Gemini service with messages:")
        print(f"   - System message length: {len(stored_prompt)}")
        print(f"   - User message length: {len(user_inputs_text)}")
        if sow_text:
            print(f"   - SOW provided (len={len(sow_text)}), trimmed to {len(sow_text_trimmed)} chars")
        if is_regeneration:
            print(f"   - REGENERATION MODE: Current plan provided (len={len(current_plan_text)}), trimmed to {len(current_plan_trimmed)} chars")
            print(f"   - Regeneration flag: {request.regenerate_with_current_plan}")
        else:
            print(f"   - NEW PLAN MODE: No current plan content provided")

        def _truncate(text, n=800):
            try:
                return text if len(text) <= n else text[:n] + "... [truncated]"
            except Exception:
                return str(text)[:n]

        payload_preview = {
            "messages": [
                {"role": "system", "content": _truncate(stored_prompt, 600)},
                {"role": "user", "content": _truncate(user_inputs_text, 1200)}
            ],
            "max_tokens": 4000
        }
        print("===== LLM PAYLOAD (PREVIEW) =====")
        try:
            import json
            print(json.dumps(payload_preview, ensure_ascii=False, indent=2))
        except Exception:
            print(str(payload_preview))
        print("===== END LLM PAYLOAD (PREVIEW) =====")
        
        gemini_response = gemini_service.chat(messages, max_tokens=4000)
        
        print("🔍 [GEMINI CALL] Gemini service response received:")
        print(f"   - Response object: {gemini_response}")
        print(f"   - Response type: {type(gemini_response)}")
        print(f"   - Response keys: {list(gemini_response.keys()) if isinstance(gemini_response, dict) else 'Not a dict'}")
        
        # Check if Gemini service actually succeeded
        if not gemini_response:
            print("❌ [GEMINI CALL] Gemini service returned no response")
            return {"success": False, "message": "Failed to generate sprint plan from Gemini"}
        
        if not gemini_response.get('success', False):
            error_msg = gemini_response.get('response', 'Unknown error from Gemini service')
            print(f"❌ [GEMINI CALL] Gemini service failed: {error_msg}")
            return {"success": False, "message": f"Gemini service failed: {error_msg}"}
        
        if not gemini_response.get('response'):
            print("❌ [GEMINI CALL] Gemini service returned empty response")
            return {"success": False, "message": "Failed to generate sprint plan from Gemini"}
        
        if gemini_response.get("success"):
            print("🔍 [STEP 1] Initial sprint plan generation completed successfully")
            
            # Count expected Product Backlog Items from user input
            expected_pb_count = len(user_inputs['product_backlog'].get('BacklogItems', []))
            print(f"🔍 [PB COUNT] Expected Product Backlog Items: {expected_pb_count}")
            
            # STEP 2: VALIDATION AND FINE-TUNING FOR PB COMPLETENESS
            print("🔍 [STEP 2] Starting PB completeness validation and fine-tuning...")
            
            # Pre-validation check for missing sections
            original_plan = gemini_response.get("response", "")
            missing_sections = []
            
            if "Detailed Task Breakdown" not in original_plan:
                missing_sections.append("Detailed Task Breakdown")
            if "Committed Sprint Backlog" not in original_plan:
                missing_sections.append("Committed Sprint Backlog")
            if "Sprint Overview" not in original_plan:
                missing_sections.append("Sprint Overview")
            if "Team Capacity" not in original_plan:
                missing_sections.append("Team Capacity")
            
            if missing_sections:
                print(f"❌ [PRE-VALIDATION] Missing critical sections: {missing_sections}")
                print("🔄 [PRE-VALIDATION] Plan will be regenerated to include missing sections")
            else:
                print("✅ [PRE-VALIDATION] All critical sections present")
            
            validation_result = gemini_service.validate_and_finetune_sprint_plan(
                original_plan=original_plan,
                user_inputs=user_inputs_text,
                stored_prompt=stored_prompt,
                expected_pb_count=expected_pb_count
            )
            
            if validation_result.get("success"):
                if validation_result.get("improved"):
                    print(f"🔄 [VALIDATION] Plan was regenerated to include all {expected_pb_count} PBs")
                    # Use the improved plan
                    final_plan = validation_result.get("response", "")
                    gemini_response["response"] = final_plan
                else:
                    print(f"✅ [VALIDATION] Plan passed validation with all {expected_pb_count} PBs")
                    # Use the original plan
                    final_plan = gemini_response.get("response", "")
                
                print(f"🔍 [VALIDATION] Final plan length: {len(final_plan)} characters")
            else:
                print(f"⚠️ [VALIDATION] Validation failed, using original plan: {validation_result.get('validation_error', 'Unknown error')}")
                # Use original plan if validation fails
                final_plan = gemini_response.get("response", "")
            
            print("🔍 [STEP 2] PB completeness validation and fine-tuning completed")
            
            # PDF generation is now handled by frontend html2pdf.js
            
            # Transform team members to match PDF service expectations
            team_members_for_pdf = []
            if user_inputs['team_capacity'].get('TeamMembers'):
                for member in user_inputs['team_capacity'].get('TeamMembers', []):
                    team_members_for_pdf.append({
                        'name': member.get('roleName', 'N/A'),
                        'role': member.get('roleName', 'N/A'),
                        'availability': f"{member.get('workingHours', 'N/A')} hours"
                    })
            
            # Transform backlog items to match PDF service expectations
            backlog_items_for_pdf = []
            if user_inputs['product_backlog'].get('BacklogItems'):
                for item in user_inputs['product_backlog'].get('BacklogItems', []):
                    backlog_items_for_pdf.append({
                        'summary': item.get('userStorySummary', 'N/A'),
                        'priority': item.get('priority', 'N/A'),
                        'effort': f"{item.get('effortEstimate', 'N/A')} hours"
                    })
            
            # CRITICAL: Store generated plan content in a backup variable
            generated_plan_backup = gemini_response.get("response", "")
            print(f"🔍 [BACKUP] Generated plan backup created: {len(generated_plan_backup)} characters")
            print(f"🔍 [BACKUP] Generated plan backup preview: {generated_plan_backup[:200]}...")
            
            # Prepare plan data for PDF generation
            plan_data = {
                "sprint_number": user_inputs['sprint_overview'].get('SprintNumber', ''),
                "sprint_dates": user_inputs['sprint_overview'].get('SprintDates', ''),
                "sprint_duration": user_inputs['sprint_overview'].get('SprintDuration', ''),
                "team_name": user_inputs['sprint_overview'].get('TeamName', ''),
                "sprint_goal": user_inputs['sprint_overview'].get('SprintGoal', ''),
                "total_hours_per_person": user_inputs['team_capacity'].get('TotalHoursPerPerson', ''),
                "number_of_members": user_inputs['team_capacity'].get('NumberOfMembers', ''),
                "team_members": team_members_for_pdf,
                "historical_story_points": user_inputs['team_capacity'].get('HistoricalStoryPoints', ''),
                "backlog_items": backlog_items_for_pdf,
                "definition_of_done": user_inputs['definition_of_done'].get('DoDContent', ''),
                "risks_and_impediments": user_inputs['risks_and_impediments'].get('RisksContent', ''),
                "generated_plan": generated_plan_backup
            }
            
            print("🔍 [PDF GENERATION] Plan data being sent to PDF service:")
            print(f"   - Sprint Number: {plan_data.get('sprint_number')}")
            print(f"   - Team Name: {plan_data.get('team_name')}")
            print(f"   - Team Members Count: {len(plan_data.get('team_members', []))}")
            print(f"   - Backlog Items Count: {len(plan_data.get('backlog_items', []))}")
            print(f"   - Generated Plan Length: {len(plan_data.get('generated_plan', ''))}")
            print(f"   - Generated Plan Preview: {plan_data.get('generated_plan', '')[:200]}...")
            
            generated_plan_content = plan_data.get('generated_plan', '')
            if generated_plan_content:
                print(f"🔍 [GENERATED PLAN ANALYSIS] Content analysis:")
                print(f"   - Contains HTML tags: {'<h2>' in generated_plan_content or '<section>' in generated_plan_content}")
                print(f"   - Contains 'Sprint Overview': {'Sprint Overview' in generated_plan_content}")
                print(f"   - Contains 'Confirmed Sprint Goal': {'Confirmed Sprint Goal' in generated_plan_content}")
                print(f"   - Contains 'Team Capacity': {'Team Capacity' in generated_plan_content}")
                print(f"   - First 500 characters: {generated_plan_content[:500]}...")
            else:
                print("❌ [GENERATED PLAN ANALYSIS] Generated plan is empty!")
            
            
            # Add delay to ensure generated_plan is fully populated
            import time
            print("⏳ [PDF GENERATION] Waiting 3 seconds for data to settle...")
            time.sleep(3)
            
            # Double-check the generated_plan content before PDF generation
            final_generated_plan = plan_data.get('generated_plan', '')
            print(f"🔍 [PDF GENERATION] Final check - Generated plan length: {len(final_generated_plan)}")
            print(f"🔍 [PDF GENERATION] Final check - Generated plan preview: {final_generated_plan[:300]}...")
            
            if not final_generated_plan or len(final_generated_plan.strip()) < 100:
                print("❌ [PDF GENERATION] Generated plan is still empty or too short!")
                print("❌ [PDF GENERATION] Waiting additional 5 seconds...")
                time.sleep(5)
                final_generated_plan = plan_data.get('generated_plan', '')
                print(f"🔍 [PDF GENERATION] After additional wait - Length: {len(final_generated_plan)}")
            
            # Update plan_data with final content
            plan_data['generated_plan'] = final_generated_plan
            
            # CRITICAL: Ensure generated_plan is not lost
            print(f"🔍 [PDF GENERATION] CRITICAL CHECK - Before PDF generation:")
            print(f"   - Generated plan in plan_data: {len(plan_data.get('generated_plan', ''))} characters")
            print(f"   - Generated plan content preview: {plan_data.get('generated_plan', '')[:300]}...")
            print(f"   - Generated plan contains HTML: {'<h2>' in plan_data.get('generated_plan', '')}")
            
            # FINAL SAFEGUARD: If generated_plan is still empty, use backup
            if not plan_data.get('generated_plan') or len(plan_data.get('generated_plan', '').strip()) < 100:
                print("🚨 [SAFEGUARD] Generated plan is empty, using backup!")
                plan_data['generated_plan'] = generated_plan_backup
                print(f"🚨 [SAFEGUARD] Restored from backup: {len(plan_data['generated_plan'])} characters")
            
            # Generate Word document content (replica of HTML rendered output)
            print("📝 [WORD GENERATION] Generating Word document content...")
            word_document_content = generate_word_document_content(
                plan_data.get('generated_plan', ''),
                user_inputs
            )
            plan_data['word_document'] = word_document_content
            print(f"📝 [WORD GENERATION] Word document generated: {len(word_document_content)} characters")
            print(f"📝 [WORD GENERATION] Word document preview: {word_document_content[:200]}...")
            
            # Generate PDF with SprintResultsPage styling
            # pdf_result = pdf_service.generate_sprint_plan_pdf(plan_data) # This line is removed
            
            # Save to database with new structure including PDF and optional SOW content
            db_result = db_service.save_sprint_plan(db, {
                **plan_data,
                "sow_content": request.sow_content,
                "workspace_id": request.workspace_id,
                # PDF generation is now handled by frontend html2pdf.js
            }, request.user_email)
            
            if db_result.get("success"):
                return {
                    "success": True,
                    "message": "Sprint plan generated and saved successfully",
                    "response": gemini_response.get("response", ""),
                    "plan_id": db_result.get("plan_id")
                }
            else:
                print(f"❌ [GENERATE PLAN] Database save failed: {db_result.get('message')}")
                return {
                    "success": False,
                                    "message": f"Failed to save sprint plan to database: {db_result.get('message')}"
            }
        else:
            return {
                "success": False,
                "message": "Failed to generate sprint plan",
                "error": gemini_response.get("error", "Unknown error")
            }
            
    except Exception as e:
        print(f"❌ [GENERATE PLAN] Error: {str(e)}")
        return {"success": False, "message": f"Error generating sprint plan: {str(e)}"}

# LLM and Gemini endpoints
@app.post("/api/llm/chat", response_model=LLMChatResponse)
async def llm_chat(request: LLMChatRequest):
    """Mock LLM chat endpoint"""
    return llm_service.chat(request)



@app.post("/api/gemini/chat")
async def gemini_chat(request: dict):
    """Gemini chat endpoint"""
    from services.gemini_service import gemini_service
    messages = request.get("messages", []) or []

    # Prepend structured HTML system prompt if not already included
    system_prompt = _get_structured_html_system_prompt()
    if not (messages and messages[0].get("role") == "system"):
        messages = [{"role": "system", "content": system_prompt}] + messages
    else:
        # If a system message exists, combine ours ahead of it
        messages = [{"role": "system", "content": system_prompt}] + messages

    return gemini_service.chat(messages, request.get("max_tokens", 3000))

# Background task for indexing files in Pinecone
def index_file_background(file_id: int, text: str, source_filename: str, file_type: str, uploaded_by: str, uploaded_at):
    """
    Background task to index a file in Pinecone.
    This runs asynchronously and doesn't block the upload response.
    """
    try:
        from services.pinecone_service import pinecone_service
        from services.chunking_service import chunking_service
        from services.embedding_service import embedding_service
        from database import SessionLocal
        from datetime import datetime
        import traceback
        
        # Create a new database session for the background task
        db = SessionLocal()
        try:
            # Get the uploaded file record
            from models import UploadedFile
            uploaded_file = db.query(UploadedFile).filter(UploadedFile.id == file_id).first()
            
            if not uploaded_file:
                print(f"⚠️ [BACKGROUND INDEX] File {file_id} not found in database")
                return
            
            # Use extracted_text from database if passed text is empty or None
            text_to_index = text
            if not text_to_index or not text_to_index.strip():
                text_to_index = uploaded_file.extracted_text or ""
                print(f"📋 [BACKGROUND INDEX] File {file_id}: Using extracted_text from database (length: {len(text_to_index) if text_to_index else 0})")
            
            # Check if we have text to index
            if not text_to_index or not text_to_index.strip():
                print(f"⚠️ [BACKGROUND INDEX] File {file_id}: No text content to index (empty extracted_text)")
                uploaded_file.indexing_status = "error"
                db.commit()
                return
            
            # Handle datetime conversion if needed
            upload_datetime = uploaded_at
            if isinstance(upload_datetime, str):
                try:
                    # Try parsing ISO format datetime string
                    upload_datetime = datetime.fromisoformat(upload_datetime.replace('Z', '+00:00'))
                except:
                    upload_datetime = uploaded_file.upload_time or datetime.utcnow()
            elif upload_datetime is None:
                upload_datetime = uploaded_file.upload_time or datetime.utcnow()
            
            # Use values from database if not provided
            source_filename = source_filename or uploaded_file.file_name or "unknown"
            file_type = file_type or uploaded_file.file_type or "unknown"
            uploaded_by = uploaded_by or uploaded_file.uploaded_by or "anonymous"
            
            print(f"📋 [BACKGROUND INDEX] Starting indexing for file {file_id} ({source_filename}), text length: {len(text_to_index)}")
            
            # Create Pinecone index for this file
            pinecone_index_result = pinecone_service.create_index_for_file(
                file_id=file_id,
                file_name=source_filename
            )
            
            if not pinecone_index_result.get("success"):
                error_msg = pinecone_index_result.get("error", "Unknown error creating Pinecone index")
                uploaded_file.indexing_status = "error"
                db.commit()
                print(f"⚠️ [BACKGROUND INDEX] File {file_id} indexing failed: {error_msg}")
                return
            
            # Chunk text for Pinecone (400 chars, 100 overlap)
            chunks = chunking_service.chunk_text_by_characters(
                text=text_to_index,
                chunk_size=400,
                chunk_overlap=100,
                metadata={
                    "file_id": file_id,
                    "file_name": source_filename,
                    "file_type": file_type,
                    "uploaded_by": uploaded_by,
                    "uploaded_at": upload_datetime.isoformat() if upload_datetime else None
                }
            )
            
            if not chunks:
                uploaded_file.indexing_status = "error"
                db.commit()
                print(f"⚠️ [BACKGROUND INDEX] File {file_id}: No chunks generated for Pinecone indexing")
                return
            
            # Generate embeddings for chunks
            chunk_texts = [chunk["text"] for chunk in chunks]
            embeddings = embedding_service.embed(chunk_texts)
            
            # Index chunks in Pinecone
            index_chunks_result = pinecone_service.index_file_chunks(
                file_id=file_id,
                file_name=source_filename,
                chunks=chunks,
                embeddings=embeddings
            )
            
            if index_chunks_result.get("success"):
                uploaded_file.indexing_status = "indexed"
                db.commit()
                print(f"✅ [BACKGROUND INDEX] File {file_id} indexed successfully in Pinecone: {index_chunks_result.get('chunks_indexed', 0)} chunks")
            else:
                error_msg = index_chunks_result.get('error', 'Unknown error during Pinecone indexing')
                uploaded_file.indexing_status = "error"
                db.commit()
                print(f"⚠️ [BACKGROUND INDEX] File {file_id} indexing failed: {error_msg}")
                print(f"⚠️ [BACKGROUND INDEX] File {file_id} error details: {traceback.format_exc()}")
            
        except Exception as e:
            # Update status to error if indexing fails
            error_msg = str(e)
            error_trace = traceback.format_exc()
            print(f"⚠️ [BACKGROUND INDEX] File {file_id} indexing exception: {error_msg}")
            print(f"⚠️ [BACKGROUND INDEX] File {file_id} exception traceback:\n{error_trace}")
            
            try:
                uploaded_file = db.query(UploadedFile).filter(UploadedFile.id == file_id).first()
                if uploaded_file:
                    uploaded_file.indexing_status = "error"
                    db.commit()
            except Exception as db_error:
                print(f"⚠️ [BACKGROUND INDEX] Failed to update error status: {str(db_error)}")
        finally:
            db.close()
    except Exception as e:
        error_msg = str(e)
        error_trace = traceback.format_exc()
        print(f"⚠️ [BACKGROUND INDEX] Critical error indexing file {file_id}: {error_msg}")
        print(f"⚠️ [BACKGROUND INDEX] Critical error traceback:\n{error_trace}")

# Helper function to process a single file
async def process_single_file(
    file: UploadFile,
    uploaded_by: str,
    db: Session,
    background_tasks: BackgroundTasks
) -> dict:
    """
    Process a single file: validate, save, extract text, save to DB, and queue indexing.
    Returns a dict with success status and file information or error.
    """
    import uuid
    from pathlib import Path
    
    try:
        # Validate file type
        file_extension = file.filename.lower().split('.')[-1] if '.' in file.filename else ''
        
        if file_extension not in ['pdf', 'docx', 'txt', 'doc', 'xlsx']:
            return {
                "success": False,
                "error": f"Invalid file type. Supported formats: PDF, DOCX, TXT, XLSX. Got: {file_extension}",
                "file_name": file.filename
            }
        
        # Create uploads directory if it doesn't exist
        upload_dir = Path("uploads")
        upload_dir.mkdir(parents=True, exist_ok=True)
        
        # Generate unique filename
        unique_filename = f"{uuid.uuid4()}_{file.filename}"
        file_path = upload_dir / unique_filename
        
        # Save file to disk
        file_content = await file.read()
        with open(file_path, "wb") as f:
            f.write(file_content)
        
        # Extract text based on file type
        extracted_text = ""
        if file_extension == 'pdf':
            result = pdf_service.extract_text_from_pdf(file_content)
            if result['success']:
                extracted_text = result['text']
            else:
                return {
                    "success": False,
                    "error": f"Failed to extract text from PDF: {result.get('error', 'Unknown error')}",
                    "file_name": file.filename
                }
        elif file_extension in ['docx', 'doc']:
            # Use enhanced docx extraction to preserve hyperlinks
            try:
                from services.docx_extraction_helper import extract_text_with_hyperlinks_from_docx
                extracted_text = extract_text_with_hyperlinks_from_docx(file_content)
            except Exception as e:
                # Fallback to simple extraction if enhanced extraction fails
                print(f"⚠️ [UPLOAD] Enhanced DOCX extraction failed for {file.filename}, using fallback: {str(e)}")
                from docx import Document
                import io
                doc = Document(io.BytesIO(file_content))
                text_parts = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text.strip())
                extracted_text = '\n'.join(text_parts)
        elif file_extension == 'txt':
            extracted_text = file_content.decode('utf-8', errors='ignore')
        elif file_extension == 'xlsx':
            # Extract text from Excel using openpyxl (cells by rows)
            try:
                import io
                from openpyxl import load_workbook
                workbook = load_workbook(filename=io.BytesIO(file_content), data_only=True)
                lines = []
                
                for sheet in workbook.worksheets:
                    sheet_lines = [f"Sheet: {sheet.title}"]
                    row_count = 0
                    for row in sheet.iter_rows(values_only=True):
                        # Filter out None, empty strings, and convert to strings
                        cells = []
                        for cell in row:
                            if cell is not None:
                                cell_str = str(cell).strip()
                                if cell_str:  # Only add non-empty cells
                                    cells.append(cell_str)
                        
                        if cells:  # Only add rows with content
                            sheet_lines.append("\t".join(cells))
                            row_count += 1
                    
                    # Only add sheet if it has content
                    if row_count > 0:
                        lines.extend(sheet_lines)
                
                extracted_text = "\n".join(lines) if lines else ""
                
                # Ensure we have some text content
                if not extracted_text or not extracted_text.strip():
                    extracted_text = f"[XLSX file: {file.filename}]\nThis file contains spreadsheet data but no extractable text content was found. The file may contain only formulas, images, or empty cells."
                    print(f"⚠️ [UPLOAD] XLSX file {file.filename} has no extractable text, using placeholder")
                else:
                    print(f"✅ [UPLOAD] XLSX file {file.filename} extracted {len(extracted_text)} characters")
                    
            except Exception as e:
                print(f"⚠️ [UPLOAD] XLSX extraction error for {file.filename}: {str(e)}")
                import traceback
                print(f"⚠️ [UPLOAD] XLSX extraction traceback:\n{traceback.format_exc()}")
                # Don't fail the upload, use placeholder text
                extracted_text = f"[XLSX file: {file.filename}]\nError extracting text: {str(e)}\nThis file may contain only formulas, images, or be in an unsupported format."
        
        # Determine uploaded_by (use provided value or default)
        user_email = uploaded_by or "anonymous"
        
        # Save to database
        uploaded_file = UploadedFile(
            file_name=file.filename,
            file_type=file_extension,
            file_path=str(file_path),
            uploaded_by=user_email,
            status="Processed",
            extracted_text=extracted_text,
            indexing_status="pending_index"
        )
        
        db.add(uploaded_file)
        db.commit()
        db.refresh(uploaded_file)
        
        # Schedule indexing in Pinecone (always background to prevent blocking)
        if background_tasks is None:
            background_tasks = BackgroundTasks()
        
        background_tasks.add_task(
            index_file_background,
            file_id=uploaded_file.id,
            text=extracted_text,
            source_filename=uploaded_file.file_name,
            file_type=uploaded_file.file_type,
            uploaded_by=uploaded_file.uploaded_by,
            uploaded_at=uploaded_file.upload_time
        )
        print(f"📋 [UPLOAD] File {uploaded_file.id} ({uploaded_file.file_name}) queued for Pinecone indexing")
        
        return {
            "success": True,
            "file_id": uploaded_file.id,
            "file_name": uploaded_file.file_name,
            "file_type": uploaded_file.file_type,
            "indexing_status": uploaded_file.indexing_status,
            "extracted_length": len(extracted_text),
            "message": f"File uploaded and processed successfully. Extracted {len(extracted_text)} characters. Indexing in progress..."
        }
        
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        print(f"⚠️ [UPLOAD] Error processing file {file.filename}: {str(e)}")
        print(f"⚠️ [UPLOAD] Traceback:\n{error_trace}")
        return {
            "success": False,
            "error": f"Error processing file: {str(e)}",
            "file_name": file.filename
        }

# Chatbot file upload and question endpoints
@app.post("/api/upload-file")
async def upload_chatbot_file(
    request: Request,
    uploaded_by: str = Form(None),
    background_tasks: BackgroundTasks = None,
    db: Session = Depends(get_db)
):
    """
    Upload one or multiple files for chatbot analysis
    Supports PDF, DOCX, TXT, and XLSX files
    Accepts up to 10 files in a single request
    Saves files to /uploads/ and extracts text content
    Each file is processed independently and indexed separately
    
    Backward compatible: accepts both 'file' (single) and 'files' (multiple) parameters
    """
    import uuid
    from datetime import datetime
    from pathlib import Path
    
    try:
        # Parse form data manually to handle both 'file' and 'files' parameters
        form = await request.form()
        normalized_files = []
        
        # Get uploaded_by from form if not provided as parameter
        if uploaded_by is None and 'uploaded_by' in form:
            uploaded_by = form['uploaded_by']
        
        # Check for 'files' (multiple) parameter first (new format)
        if 'files' in form:
            files_list = form.getlist('files')
            if files_list and len(files_list) > 0:
                # Filter out None values
                files_list = [f for f in files_list if f is not None]
                if files_list:
                    normalized_files = files_list
            # If 'files' exists but is empty, check 'file' parameter
            if len(normalized_files) == 0 and 'file' in form:
                file_obj = form['file']
                if file_obj:
                    normalized_files = [file_obj]
        # Fallback: check 'file' (single) parameter for backward compatibility
        elif 'file' in form:
            file_obj = form['file']
            if file_obj:
                normalized_files = [file_obj]
        
        if len(normalized_files) == 0:
            return {
                "success": False,
                "error": "No files provided. Please upload at least one file."
            }
        
        files = normalized_files
        
        # Validate number of files (max 10)
        if len(files) > 10:
            return {
                "success": False,
                "error": f"Too many files. Maximum 10 files allowed. Got {len(files)} files."
            }
        
        # Process each file independently
        results = []
        successful_files = []
        failed_files = []
        
        print(f"📤 [UPLOAD] Processing {len(files)} file(s)...")
        
        for idx, file in enumerate(files, 1):
            print(f"📄 [UPLOAD] Processing file {idx}/{len(files)}: {file.filename}")
            result = await process_single_file(file, uploaded_by, db, background_tasks)
            results.append(result)
            
            if result["success"]:
                successful_files.append(result)
            else:
                failed_files.append(result)
        
        # Prepare response
        total_files = len(files)
        success_count = len(successful_files)
        failed_count = len(failed_files)
        
        response = {
            "success": True,
            "total_files": total_files,
            "successful_uploads": success_count,
            "failed_uploads": failed_count,
            "files": results,
            "message": f"Processed {total_files} file(s): {success_count} successful, {failed_count} failed"
        }
        
        # If all files failed, mark overall success as False
        if success_count == 0:
            response["success"] = False
            response["error"] = "All files failed to upload"
        # If some files failed, still return success but with warning
        elif failed_count > 0:
            response["warning"] = f"{failed_count} file(s) failed to upload. Check individual file results for details."
        
        print(f"✅ [UPLOAD] Batch upload complete: {success_count}/{total_files} successful")
        return response
        
    except Exception as e:
        db.rollback()
        import traceback
        error_trace = traceback.format_exc()
        print(f"⚠️ [UPLOAD] Batch upload error: {str(e)}")
        print(f"⚠️ [UPLOAD] Traceback:\n{error_trace}")
        return {
            "success": False,
            "error": f"Error processing batch upload: {str(e)}"
        }

def _router_answerer_system_prompt() -> str:
    """Returns the ROUTER + ANSWERER system prompt"""
    return """You are a reliable two-step assistant (ROUTER + ANSWERER) that receives:

  - a user question,
  - a list of candidates from the vector DB (top-K chunk hits across all files) summarized as `file_scores`, and
  - the top context chunks (documents) retrieved for each file candidate.

Your job:

  1) ROUTE: Decide which file(s) are most appropriate to answer the question using the `file_scores`.

  2) RETRIEVE/ANSWER: Use only the provided `context_chunks` (the chunks attached from Pinecone) from the selected file(s) and produce a concise, factual answer. Do NOT use external knowledge beyond light, safe inference. Always cite sources inline and in the `sources` array.

  3) OUTPUT: Return a single valid JSON object (no extra text) matching the specified schema exactly.

Routing rules (apply before answering):

  - `file_scores` gives, per file: file_name, top_chunk_score (similarity 0..1), and 1-2 line summary.
  - CONFIDENCE thresholds:
      HIGH_CONFIDENCE if top score >= 0.72
      MEDIUM_CONFIDENCE if 0.55 <= top score < 0.72
      LOW_CONFIDENCE if top score < 0.55
  - If top two files have scores within 0.03 of each other, include both files as `selected_files` (order by score).
  - If LOW_CONFIDENCE: do not hallucinate. Instead, set `"status": "LOW_CONFIDENCE"` and attempt a conservative best-effort answer using top available chunks from all candidate files, and include a `clarifying_question` suggestion for the user.

Answering rules:

  - Use ONLY the `context_chunks` provided. Do not invent facts. You may make short, reasonable inferences, but label them as "inference" if not explicitly present in chunks.
  - Keep answers concise by default (<= 200 words). If the user explicitly asks for more depth, extend.
  - Cite each factual statement that comes from a chunk using inline bracketed citations like: [Project Management Playbook.docx — chunk_12].
  - After the answer include a `Sources` section listing each chunk used with `file_name`, `chunk_id`, and `score`.
  - If no chunk contains an exact answer, explicitly say: "I couldn't find an exact answer in the provided documents; here's the best guidance based on them." Then answer and cite.

Output JSON schema (MUST match exactly):

{
  "status": "OK" | "MEDIUM_CONFIDENCE" | "LOW_CONFIDENCE",
  "selected_files": ["file1.docx", "file2.xlsx"],
  "routing_detail": {
    "top_file": "file1.docx",
    "top_score": 0.82,
    "file_scores": [ { "file_name": "...", "top_chunk_score": 0.82, "summary":"..." }, ... ]
  },
  "answer": "<the assistant's concise answer (string)>",
  "sources": [
    { "file_name": "Project Management Playbook.docx", "chunk_id": "chunk_72", "score": 0.83, "excerpt": "first 200 chars of the chunk..." },
    ...
  ],
  "confidence_explanation": "<one-line explanation of confidence and any limitations>",
  "clarifying_question": "<optional string - present only when status is LOW_CONFIDENCE or ambiguous>",
  "raw_used_chunks": [ { "file_name":"...", "chunk_id":"...", "text":"full chunk text (or truncated to X chars)", "score": 0.83 } ]
}

Formatting constraints:
  - Return only the JSON object (no surrounding commentary or markdown).
  - Truncate long chunk excerpts in `sources` to 200 chars. `raw_used_chunks` may be truncated to 2000 chars each if necessary to keep output practical."""


def _resolve_user_email(db: Session, chat_id: str, provided_email: str = None) -> str:
    """Resolve which user_email to store for a chat."""
    if provided_email:
        return provided_email
    
    try:
        existing = (
            db.query(ChatMessage.user_email)
            .filter(ChatMessage.chat_id == chat_id, ChatMessage.user_email.isnot(None))
            .order_by(ChatMessage.created_at.asc())
            .first()
        )
        if existing and existing[0]:
            return existing[0]
    except Exception:
        pass
    return None


def _save_chat_message(
    db: Session,
    chat_id: str,
    role: str,
    message: str,
    user_email: str = None
) -> None:
    """Persist chat messages by updating conversation JSON instead of creating individual records."""
    if not chat_id or not role or message is None:
        return
    
    cleaned_message = message.strip() if isinstance(message, str) else ""
    if not cleaned_message:
        return
    
    try:
        import json
        resolved_email = _resolve_user_email(db, chat_id, user_email)
        truncated_message = cleaned_message[:8000]  # Prevent excessively long entries
        
        # Find or create conversation record
        conversation = db.query(Conversation).filter(
            Conversation.chat_id == chat_id
        ).first()
        
        if conversation:
            # Update existing conversation - preserve project_id if it exists
            conversation_data = conversation.conversation_json
            if not isinstance(conversation_data, dict):
                conversation_data = json.loads(conversation_data) if isinstance(conversation_data, str) else {}
            
            messages = conversation_data.get("messages", [])
            
            # Find the last message entry or create new one
            if messages and role == "assistant" and len(messages) > 0:
                # If assistant message, update the last message entry
                last_msg = messages[-1]
                if "assistant" not in last_msg or not last_msg.get("assistant"):
                    last_msg["assistant"] = truncated_message
                else:
                    # Create new message entry
                    new_message_id = len(messages) + 1
                    messages.append({
                        "message_id": new_message_id,
                        "user": "",
                        "assistant": truncated_message
                    })
            elif messages and role == "user":
                # If user message, check if last message has user or create new
                last_msg = messages[-1] if messages else None
                if last_msg and not last_msg.get("user"):
                    last_msg["user"] = truncated_message
                else:
                    # Create new message entry
                    new_message_id = len(messages) + 1
                    messages.append({
                        "message_id": new_message_id,
                        "user": truncated_message,
                        "assistant": ""
                    })
            else:
                # First message
                messages.append({
                    "message_id": 1,
                    "user": truncated_message if role == "user" else "",
                    "assistant": truncated_message if role == "assistant" else ""
                })
            
            conversation_data["messages"] = messages
            conversation.conversation_json = conversation_data
            if resolved_email:
                conversation.user_email = resolved_email
            # Note: project_id is already set on the conversation, we don't need to update it
            db.commit()
        else:
            # Create new conversation
            conversation_id = str(uuid.uuid4())
            conversation_data = {
                "conversation_id": 1,  # Will be updated to actual id after save
                "messages": [{
                    "message_id": 1,
                    "user": truncated_message if role == "user" else "",
                    "assistant": truncated_message if role == "assistant" else ""
                }]
            }
            new_conversation = Conversation(
                conversation_id=conversation_id,
                chat_id=chat_id,
                user_email=resolved_email,
                conversation_json=conversation_data
            )
            db.add(new_conversation)
            db.commit()
            db.refresh(new_conversation)
            # Update conversation_id in JSON to actual database id
            conversation_data["conversation_id"] = new_conversation.id
            new_conversation.conversation_json = conversation_data
            db.commit()
    except Exception as e:
        db.rollback()
        print(f"⚠️ [CHAT] Failed to save chat message: {str(e)}")


def _search_across_all_files_and_route(
    question: str,
    top_k: int = 10,
    db: Session = None
) -> Dict[str, Any]:
    """
    Search across all indexed files, group by file, calculate file scores,
    and prepare data for ROUTER + ANSWERER system.
    
    Returns:
        Dict with file_scores and context_chunks ready for LLM routing
    """
    from services.pinecone_service import pinecone_service
    from services.embedding_service import embedding_service
    from models import UploadedFile, MandatoryFile, ProjectKnowledgeBaseFile
    
    print(f"🔍 [ROUTER] Searching across all Pinecone indexes with top_k={top_k}")
    
    if db is None:
        print(f"⚠️ [ROUTER] Database session not provided")
        return {"file_scores": [], "context_chunks": []}
    
    # Gather indexed uploaded files
    uploaded_files = db.query(UploadedFile).filter(UploadedFile.indexing_status == "indexed").all()
    
    # Also include knowledge base mandatory files that are indexed
    kb_file_ids = [
        kb_file.mandatory_file_id
        for kb_file in db.query(ProjectKnowledgeBaseFile).all()
    ]
    mandatory_files = []
    if kb_file_ids:
        mandatory_files = db.query(MandatoryFile).filter(MandatoryFile.id.in_(kb_file_ids)).all()
    
    if not uploaded_files and not mandatory_files:
        print(f"⚠️ [ROUTER] No indexed files available for routing")
        return {"file_scores": [], "context_chunks": []}
    
    existing_indexes = pinecone_service.list_indexes()
    index_names = []
    file_info_map = {}
    
    # Add uploaded file indexes
    for file in uploaded_files:
        index_name = pinecone_service.get_index_name_for_file(file.id, file.file_name)
        if index_name in existing_indexes:
            index_names.append(index_name)
            file_info_map[index_name] = {
                "file_id": file.id,
                "file_name": file.file_name
            }
    
    # Add mandatory file indexes
    for file in mandatory_files:
        index_name = pinecone_service.get_index_name_for_file(file.id, file.file_name)
        if index_name in existing_indexes:
            index_names.append(index_name)
            file_info_map[index_name] = {
                "file_id": file.id,
                "file_name": file.file_name
            }
    
    # Deduplicate index names
    index_names = list(dict.fromkeys(index_names))
    
    if not index_names:
        print(f"⚠️ [ROUTER] No Pinecone indexes available for routing")
        return {"file_scores": [], "context_chunks": []}
    
    print(f"📄 [ROUTER] Searching across indexes: {[file_info_map[idx]['file_name'] for idx in index_names]}")
    
    query_embedding = embedding_service.embed_query(question)
    search_result = pinecone_service.search_across_indexes(
        query_embedding=query_embedding,
        index_names=index_names,
        top_k=max(3, top_k)  # ensure at least 3 per index
    )
    
    if not search_result.get("success"):
        print(f"⚠️ [ROUTER] Pinecone search failed: {search_result.get('error')}")
        return {"file_scores": [], "context_chunks": []}
    
    results = search_result.get("results", [])
    if not results:
        print(f"⚠️ [ROUTER] No Pinecone results found")
        return {"file_scores": [], "context_chunks": []}
    
    print(f"✅ [ROUTER] Found {len(results)} total Pinecone results")
    
    files_dict = {}
    for result in results:
        index_name = result.get("index_name")
        metadata = result.get("metadata", {})
        file_info = file_info_map.get(index_name, {})
        file_name = file_info.get("file_name", metadata.get("file_name", "unknown"))
        file_id = file_info.get("file_id", metadata.get("file_id"))
        score = result.get("score", 0.0)
        text = metadata.get("text", "")
        
        if file_name not in files_dict:
            files_dict[file_name] = {
                "file_id": file_id,
                "chunks": [],
                "top_score": 0.0
            }
        
        files_dict[file_name]["chunks"].append({
            "score": score,
            "text": text,
            "metadata": metadata,
            "chunk_id": result.get("id", f"chunk_{metadata.get('chunk_index', '?')}"),
            "index_name": index_name
        })
        
        if score > files_dict[file_name]["top_score"]:
            files_dict[file_name]["top_score"] = score
    
    file_scores = []
    context_chunks = []
    
    for file_name, file_data in sorted(files_dict.items(), key=lambda x: x[1]["top_score"], reverse=True):
        chunks = sorted(file_data["chunks"], key=lambda x: x["score"], reverse=True)
        top_chunk = chunks[0] if chunks else None
        summary = ""
        if top_chunk:
            chunk_text = top_chunk.get("text", "")
            summary = chunk_text[:200].replace("\n", " ").strip()
            if len(chunk_text) > 200:
                summary += "..."
        
        file_scores.append({
            "file_name": file_name,
            "top_chunk_score": file_data["top_score"],
            "summary": summary
        })
        
        for chunk in chunks[:3]:
            context_chunks.append({
                "file_name": file_name,
                "chunk_id": chunk.get("chunk_id"),
                "text": chunk.get("text", ""),
                "score": chunk.get("score", 0.0)
            })
    
    print(f"📊 [ROUTER] Prepared {len(file_scores)} file scores and {len(context_chunks)} context chunks")
    
    return {
        "file_scores": file_scores,
        "context_chunks": context_chunks
    }


@app.post("/api/ask-question")
async def ask_chatbot_question(
    question: str = Form(...),
    file_id: int = Form(None),
    file_context: str = Form(None),
    mandatory_file_ids: str = Form(None),
    chat_id: str = Form(None),
    user_email: str = Form(None),
    db: Session = Depends(get_db)
):
    """
    Ask a question to the chatbot using vector search.
    - If file_context is provided, uses it directly (for mandatory files - backward compatibility)
    - If file_id is provided, uses Pinecone vector search to retrieve relevant chunks
    - If neither is provided, searches across all files using Pinecone indexes and ROUTER + ANSWERER system
    """
    question = (question or "").strip()
    chat_id = chat_id or str(uuid.uuid4())
    
    print(f"\n{'='*80}")
    print(f"🔵 [ASK-QUESTION] New question received")
    print(f"🔵 [ASK-QUESTION] Question: {question[:200]}...")
    print(f"🔵 [ASK-QUESTION] Parameters - file_id: {file_id}, file_context: {'Yes' if file_context else 'No'}, mandatory_file_ids: {mandatory_file_ids}")
    print(f"{'='*80}\n")
    
    try:
        from services.gemini_service import gemini_service
        from services.embedding_service import embedding_service
        from services.pinecone_service import pinecone_service
        import json
        
        context_text = ""
        use_router_answerer = False
        use_pinecone_search = False  # Initialize Pinecone search flag
        
        _save_chat_message(db, chat_id, "user", question, user_email)
        
        if file_context:
            # Use provided file_context directly (e.g., from multiple mandatory files)
            # This maintains backward compatibility for mandatory files
            context_text = file_context
            print(f"📝 [ASK-QUESTION] ✅ Using provided file_context (PLAYBOOK/MANDATORY FILES)")
            print(f"📝 [ASK-QUESTION] Context length: {len(context_text)} characters")
            if mandatory_file_ids:
                try:
                    import json
                    ids = json.loads(mandatory_file_ids)
                    print(f"📝 [ASK-QUESTION] Mandatory file IDs used: {ids}")
                    # Fetch file names for these IDs
                    if ids:
                        mandatory_files = db.query(MandatoryFile).filter(MandatoryFile.id.in_(ids)).all()
                        file_names = [f.file_name for f in mandatory_files]
                        print(f"📝 [ASK-QUESTION] 📄 Documents being used: {file_names}")
                except:
                    print(f"📝 [ASK-QUESTION] Mandatory file IDs (raw): {mandatory_file_ids}")
            print(f"📝 [ASK-QUESTION] Question: {question[:200]}...")
        elif file_id:
            # Use Pinecone vector search to retrieve relevant chunks
            uploaded_file = db.query(UploadedFile).filter(UploadedFile.id == file_id).first()
            if not uploaded_file:
                return {
                    "success": False,
                    "error": f"File with ID {file_id} not found.",
                    "chat_id": chat_id
                }
            
            print(f"📝 [ASK-QUESTION] 📄 Document being used: {uploaded_file.file_name} (ID: {file_id})")
            
            # Check if file is indexed
            if uploaded_file.indexing_status != "indexed":
                # Fallback to full text if not indexed yet
                if uploaded_file.extracted_text:
                    context_text = uploaded_file.extracted_text
                    print(f"⚠️ [ASK-QUESTION] File {file_id} not indexed, using full text (length: {len(context_text)} characters)")
                else:
                    return {
                        "success": False,
                        "error": f"File {file_id} is not indexed and has no extracted text available.",
                        "chat_id": chat_id
                    }
            else:
                # Use Pinecone vector search
                print(f"🔍 [ASK-QUESTION] Using Pinecone search for file_id {file_id}")
                
                index_name = pinecone_service.get_index_name_for_file(file_id, uploaded_file.file_name)
                if not pinecone_service.index_exists(index_name):
                    print(f"⚠️ [ASK-QUESTION] Pinecone index {index_name} not found. Falling back to full text.")
                    if uploaded_file.extracted_text:
                        context_text = uploaded_file.extracted_text
                    else:
                        context_text = ""
                else:
                    query_embedding = embedding_service.embed_query(question)
                    search_result = pinecone_service.search_across_indexes(
                        query_embedding=query_embedding,
                        index_names=[index_name],
                        top_k=5
                    )
                    
                    results = search_result.get("results") if search_result.get("success") else []
                    
                    if results:
                        chunk_texts = []
                        for result in results:
                            metadata = result.get("metadata", {})
                            chunk_text = metadata.get("text", "")
                            chunk_index = metadata.get("chunk_index", "?")
                            file_name = metadata.get("file_name", uploaded_file.file_name)
                            score = result.get("score", 0.0)
                            chunk_texts.append(
                                f"[Chunk {chunk_index} from {file_name} (score: {score:.3f})]\n{chunk_text}"
                            )
                        
                        context_text = "\n\n---\n\n".join(chunk_texts)
                        print(f"✅ [ASK-QUESTION] Retrieved {len(results)} relevant Pinecone chunks (total length: {len(context_text)} characters)")
                    else:
                        if uploaded_file.extracted_text:
                            context_text = uploaded_file.extracted_text
                            print(f"⚠️ [ASK-QUESTION] No Pinecone results, using full text (length: {len(context_text)} characters)")
                        else:
                            context_text = ""
                            print(f"⚠️ [ASK-QUESTION] No Pinecone results and no extracted text available")
        else:
            # No file_id or file_context provided - search Pinecone indexes for knowledge base files
            use_pinecone_search = True
            print(f"🌲 [ASK-QUESTION] No specific file provided, searching Pinecone knowledge base indexes...")
        
        # Build prompt with file context if available
        if use_pinecone_search:
            # Search across all Pinecone indexes for knowledge base files
            kb_context_found = False
            try:
                knowledge_base_files = db.query(ProjectKnowledgeBaseFile).all()
                
                if knowledge_base_files:
                    kb_file_ids = [kb_file.mandatory_file_id for kb_file in knowledge_base_files]
                    all_mandatory_files = db.query(MandatoryFile).filter(
                        MandatoryFile.id.in_(kb_file_ids),
                        MandatoryFile.is_active == True,
                        MandatoryFile.extracted_text.isnot(None),
                        MandatoryFile.extracted_text != ""
                    ).all()
                    
                    if all_mandatory_files:
                        index_names = []
                        file_info_map = {}
                        existing_indexes = pinecone_service.list_indexes()
                        
                        for file in all_mandatory_files:
                            index_name = pinecone_service.get_index_name_for_file(file.id, file.file_name)
                            if index_name in existing_indexes:
                                index_names.append(index_name)
                                file_info_map[index_name] = {
                                    "file_id": file.id,
                                    "file_name": file.file_name
                                }
                        
                        if index_names:
                            print(f"🌲 [PINECONE] Searching across {len(index_names)} indexes: {[file_info_map[idx]['file_name'] for idx in index_names]}")
                            
                            query_embedding = embedding_service.embed_query(question)
                            search_result = pinecone_service.search_across_indexes(
                                query_embedding=query_embedding,
                                index_names=index_names,
                                top_k=3
                            )
                            
                            if search_result.get("success") and search_result.get("results"):
                                top_results = search_result["results"][:5]
                                
                                best_index = None
                                best_score = 0.0
                                for result in top_results:
                                    if result["score"] > best_score:
                                        best_score = result["score"]
                                        best_index = result["index_name"]
                                
                                best_file_info = file_info_map.get(best_index, {})
                                print(f"🌲 [PINECONE] Best match: {best_file_info.get('file_name', 'Unknown')} (score: {best_score:.3f})")
                                
                                chunk_texts = []
                                for result in top_results:
                                    metadata = result.get("metadata", {})
                                    chunk_text = metadata.get("text", "")
                                    file_name = metadata.get("file_name", "Unknown")
                                    chunk_id = result.get("chunk_id", result.get("id", "?"))
                                    score = result.get("score", 0.0)
                                    
                                    chunk_texts.append(f"[Chunk {chunk_id} from {file_name} (score: {score:.3f})]\n{chunk_text}")
                                
                                context_text = "\n\n---\n\n".join(chunk_texts)
                                print(f"✅ [PINECONE] Retrieved {len(top_results)} relevant chunks (best score: {best_score:.3f})")
                                kb_context_found = True
                            else:
                                print(f"⚠️ [PINECONE] No relevant results found in knowledge base.")
                        else:
                            print(f"⚠️ [PINECONE] No Pinecone indexes found for knowledge base files.")
                    else:
                        print(f"⚠️ [PINECONE] Mandatory files referenced by knowledge base are missing or empty.")
                else:
                    print(f"⚠️ [PINECONE] No files selected in knowledge base.")
            
            except Exception as e:
                print(f"❌ [PINECONE] Error searching Pinecone: {str(e)}")
                import traceback
                print(traceback.format_exc())
                kb_context_found = False
            
            if not kb_context_found:
                use_router_answerer = True
                context_text = ""
        elif use_router_answerer:
            # Use ROUTER + ANSWERER system (Pinecone fallback across all indexes)
            router_data = _search_across_all_files_and_route(question, top_k=10, db=db)
            
            if not router_data['file_scores'] or not router_data['context_chunks']:
                # No results found, fallback to simple response
                error_text = "No relevant documents found in the knowledge base. Please try rephrasing your question or upload relevant files."
                _save_chat_message(db, chat_id, "assistant", error_text, user_email)
                return {
                    "success": False,
                    "error": error_text,
                    "chat_id": chat_id
                }
            
            # Build the prompt for ROUTER + ANSWERER
            user_prompt = f"""Process the following question using the provided file scores and context chunks.

{{
  "user_question": "{question}",
  "file_scores": {json.dumps(router_data['file_scores'], indent=2)},
  "context_chunks": {json.dumps(router_data['context_chunks'], indent=2)}
}}

Return ONLY the JSON object matching the schema specified in your instructions."""
            
            # Send to Gemini with ROUTER + ANSWERER system prompt
            messages = [
                {"role": "system", "content": _router_answerer_system_prompt()},
                {"role": "user", "content": user_prompt}
            ]
            
            print(f"📤 [ROUTER] Sending to LLM with {len(router_data['file_scores'])} files, {len(router_data['context_chunks'])} chunks")
            result = gemini_service.chat(messages, max_tokens=4000)
            
            if result['success']:
                llm_response = result.get('response', '')
                
                # Try to extract JSON from response (may be wrapped in markdown code blocks)
                import re
                json_match = re.search(r'\{[\s\S]*\}', llm_response)
                if json_match:
                    json_str = json_match.group(0)
                else:
                    json_str = llm_response
                
                try:
                    router_result = json.loads(json_str)
                    
                    # Extract answer and format response
                    answer_text = router_result.get('answer', '')
                    status = router_result.get('status', 'OK')
                    
                    # Add sources section if available
                    sources = router_result.get('sources', [])
                    if sources:
                        answer_text += "\n\n**Sources:**\n"
                        for source in sources:
                            answer_text += f"- {source.get('file_name', 'Unknown')} (chunk: {source.get('chunk_id', '?')}, score: {source.get('score', 0):.3f})\n"
                    
                    # Format as HTML for frontend
                    formatted_answer = answer_text.replace('\n', '<br/>')
                    
                    print(f"✅ [ROUTER] LLM response received, status: {status}")
                    print(f"📊 [ROUTER] Selected files: {router_result.get('selected_files', [])}")
                    
                    formatted_answer = formatted_answer or ""
                    _save_chat_message(db, chat_id, "assistant", formatted_answer, user_email)
                    return {
                        "success": True,
                        "response": formatted_answer,
                        "router_result": router_result,  # Include full router result for debugging
                        "status": status,
                        "chat_id": chat_id
                    }
                except json.JSONDecodeError as e:
                    print(f"⚠️ [ROUTER] Failed to parse JSON response: {e}")
                    print(f"⚠️ [ROUTER] Raw response: {llm_response[:500]}")
                    # Fallback: return the raw response
                    cleaned_response = llm_response.replace('\n', '<br/>')
                    _save_chat_message(db, chat_id, "assistant", cleaned_response, user_email)
                    return {
                        "success": True,
                        "response": cleaned_response,
                        "status": "OK",
                        "chat_id": chat_id
                    }
            else:
                error_msg = result.get('error', 'Unknown error')
                print(f"❌ [ROUTER] Gemini API error: {error_msg}")
                assistant_error = f"Gemini API error: {error_msg}"
                _save_chat_message(db, chat_id, "assistant", assistant_error, user_email)
                return {
                    "success": False,
                    "error": assistant_error,
                    "chat_id": chat_id
                }
        elif context_text:
            prompt = f"""You are a helpful project management assistant. Based on the following document content, please answer the user's question in a clear, structured, and concise manner.

DOCUMENT CONTENT:
{context_text}

USER QUESTION:
{question}

INSTRUCTIONS:
- Provide a well-structured answer based on the document content
- Use headings, bullet points, and clear formatting
- Focus on answering the user's specific question
- If the question cannot be answered using the document, clearly state that
- Do NOT simply repeat the document content - synthesize and summarize the relevant information
- When multiple documents are provided, consider information from all of them
- **CRITICAL: The document contains links in the format "link_text (url)" or "[Link: url]". You MUST preserve ALL external links from the source document in your response.**
- **When you mention any item that has a link in the source (like "Link to standard documentation/templates", "MOM Template", "RAID Log", "Project Plan", etc.), you MUST include the actual clickable HTML link in the format: <a href="url" target="_blank">link_text</a>**
- **If a link appears in the format "link_text (url)" in the document content where url starts with http:// or https://, convert it to: <a href="url" target="_blank">link_text</a> in your response**
- **If you see "[Link: url]" format in the document, convert it to: <a href="url" target="_blank">View Document</a> or <a href="url" target="_blank">Link</a>**
- **IMPORTANT: If a section mentions a link text (like "Link to sample design document") and you can find a URL in the document content (even if not directly next to the text), include that URL as a clickable link.**
- **Always scan the document content for any URLs (especially Google Sheets links like https://docs.google.com/spreadsheets/...) and include them as clickable links when they relate to the content being discussed.**
- **For links marked as internal (format: "link_text (#internal:...)"), you can skip including the URL but still mention the link text if relevant.**
- **Do NOT skip external links. If a section mentions a link with a valid URL, include that link in your response.**
- **Look for patterns like "Link to...", "...Template", "...Log", "...Plan" - these are likely link references that need to be included with their URLs if they have valid external URLs.**
- **For any Google Sheets or document links (URLs starting with http:// or https://), always include them as clickable links.**
- **Example 1: If you see "Link to standard documentation/templates (https://example.com/templates)" in the document, your response should include: <a href="https://example.com/templates" target="_blank">Link to standard documentation/templates</a>**
- **Example 2: If you see "Link to sample design document" and later find "https://docs.google.com/spreadsheets/d/1Gg4W2tmwaWqFQHpTqFxk3EVdWVuLKFrz/edit..." in the document, include: <a href="https://docs.google.com/spreadsheets/d/1Gg4W2tmwaWqFQHpTqFxk3EVdWVuLKFrz/edit..." target="_blank">Link to sample design document</a>**

Please provide your answer:"""
        else:
            # No context available, answer without document reference
            prompt = question
        
        # Debug: Log prompt length (but not the full content to avoid cluttering logs)
        print(f"📝 [ASK-QUESTION] Prompt length: {len(prompt)} characters")
        print(f"📝 [ASK-QUESTION] Question: {question[:100]}...")
        
        # Send to Gemini
        messages = [
            {"role": "system", "content": _get_structured_html_system_prompt()},
            {"role": "user", "content": prompt}
        ]
        
        result = gemini_service.chat(messages, max_tokens=3000)
        
        if result['success']:
            # Ensure we return the LLM response, not the file context
            llm_response = result.get('response', '')
            
            # Clean up markdown code blocks if present (Gemini sometimes wraps HTML in ```html blocks)
            if llm_response and '```' in llm_response:
                import re
                # Remove opening code fences (```html, ```, etc.)
                llm_response = re.sub(r'```[a-zA-Z]*\s*\n?', '', llm_response)
                # Remove closing code fences
                llm_response = re.sub(r'```\s*\n?', '', llm_response)
                llm_response = llm_response.strip()
                print(f"📝 [ASK-QUESTION] Cleaned markdown code blocks from response")
            
            # Debug: Log response length to ensure we're getting LLM output
            print(f"📝 [ASK-QUESTION] LLM Response length: {len(llm_response)} characters")
            print(f"📝 [ASK-QUESTION] LLM Response preview: {llm_response[:200]}...")
            
            # Ensure response is not empty
            if not llm_response or not llm_response.strip():
                llm_response = "<p>I apologize, but I couldn't generate a response. Please try again.</p>"
                print(f"⚠️ [ASK-QUESTION] Empty response detected, using fallback message")
            
            _save_chat_message(db, chat_id, "assistant", llm_response, user_email)
            return {
                "success": True,
                "response": llm_response,
                "file_id": file_id,
                "chat_id": chat_id
            }
        else:
            error_msg = result.get('error', 'Unknown error')
            print(f"❌ [ASK-QUESTION] Gemini API error: {error_msg}")
            assistant_error = f"Gemini API error: {error_msg}"
            _save_chat_message(db, chat_id, "assistant", assistant_error, user_email)
            return {
                "success": False,
                "error": assistant_error,
                "response": result.get('response', 'Failed to get response'),
                "chat_id": chat_id
            }
            
    except Exception as e:
        error_message = f"Error processing question: {str(e)}"
        _save_chat_message(db, chat_id, "assistant", error_message, user_email)
        return {
            "success": False,
            "error": error_message,
            "chat_id": chat_id
        }


@app.get("/api/chat/sessions")
async def get_chat_sessions(user_email: str = None, db: Session = Depends(get_db)):
    """Return a list of distinct chat sessions for the current user from conversations table.
    Only returns conversations that don't belong to a project (project_id IS NULL).
    Project conversations are shown under their respective projects."""
    try:
        import json
        
        # Query conversations table - only get conversations WITHOUT project_id
        query = db.query(Conversation).filter(Conversation.project_id.is_(None))
        
        if user_email:
            query = query.filter(
                or_(
                    Conversation.user_email == user_email,
                    Conversation.user_email.is_(None)
                )
            )
        
        conversations = query.order_by(Conversation.updated_at.desc()).all()
        
        sessions = []
        for conv in conversations:
            # Extract preview from first message in conversation JSON
            preview = ""
            try:
                conv_data = conv.conversation_json
                if not isinstance(conv_data, dict):
                    conv_data = json.loads(conv_data) if isinstance(conv_data, str) else {}
                
                messages = conv_data.get("messages", [])
                if messages and len(messages) > 0:
                    first_msg = messages[0]
                    # Get first non-empty user or assistant message
                    preview = first_msg.get("user", "") or first_msg.get("assistant", "")
                    if preview:
                        preview = preview.strip()[:120]
            except Exception as e:
                print(f"⚠️ [CHAT] Error extracting preview: {str(e)}")
            
            sessions.append({
                "chat_id": conv.chat_id,
                "conversation_id": conv.id,
                "first_message_preview": preview,
                "last_message_at": conv.updated_at.isoformat() if conv.updated_at else None,
                "updated_at": conv.updated_at.isoformat() if conv.updated_at else None
            })
        
        return {
            "success": True,
            "chats": sessions
        }
    except Exception as e:
        return {
            "success": False,
            "error": f"Error fetching chat sessions: {str(e)}"
        }


@app.get("/api/chat/messages")
async def get_chat_messages(chat_id: str, user_email: str = None, db: Session = Depends(get_db)):
    """Return the ordered messages for a specific chat session from conversations table."""
    if not chat_id:
        return {
            "success": False,
            "error": "chat_id is required"
        }
    
    try:
        import json
        
        # Query conversation by chat_id
        query = db.query(Conversation).filter(Conversation.chat_id == chat_id)
        if user_email:
            query = query.filter(
                or_(
                    Conversation.user_email == user_email,
                    Conversation.user_email.is_(None)
                )
            )
        
        conversation = query.first()
        
        if not conversation:
            return {
                "success": True,
                "messages": [],
                "project": None
            }
        
        # Get project info if this conversation belongs to a project
        project_info = None
        if conversation.project_id:
            project = db.query(Project).filter(Project.id == conversation.project_id).first()
            if project:
                project_info = {
                    "id": project.id,
                    "name": project.name
                }
        
        # Extract messages from conversation JSON
        conv_data = conversation.conversation_json
        if not isinstance(conv_data, dict):
            conv_data = json.loads(conv_data) if isinstance(conv_data, str) else {}
        
        messages_json = conv_data.get("messages", [])
        
        # Convert JSON format to message list format
        messages = []
        for msg_obj in messages_json:
            message_id = msg_obj.get("message_id", 0)
            user_msg = msg_obj.get("user", "").strip()
            assistant_msg = msg_obj.get("assistant", "").strip()
            
            # Add user message if exists
            if user_msg:
                messages.append({
                    "id": f"{conversation.id}-{message_id}-user",
                    "chat_id": conversation.chat_id,
                    "user_email": conversation.user_email,
                    "role": "user",
                    "message": user_msg,
                    "created_at": conversation.created_at.isoformat() if conversation.created_at else None
                })
            
            # Add assistant message if exists
            if assistant_msg:
                messages.append({
                    "id": f"{conversation.id}-{message_id}-assistant",
                    "chat_id": conversation.chat_id,
                    "user_email": conversation.user_email,
                    "role": "assistant",
                    "message": assistant_msg,
                    "created_at": conversation.updated_at.isoformat() if conversation.updated_at else None
                })
        
        return {
            "success": True,
            "messages": messages,
            "project": project_info
        }
    except Exception as e:
        return {
            "success": False,
            "error": f"Error fetching chat messages: {str(e)}"
        }


@app.post("/api/chat/save-message")
async def save_chat_message(
    chat_id: str = Form(...),
    role: str = Form(...),
    message: str = Form(...),
    user_email: str = Form(None),
    db: Session = Depends(get_db)
):
    """Save a single chat message in real-time by updating conversation JSON."""
    try:
        import json
        
        # Find or create conversation record
        conversation = db.query(Conversation).filter(
            Conversation.chat_id == chat_id
        ).first()
        
        if conversation:
            # Update existing conversation
            conversation_data = conversation.conversation_json
            if not isinstance(conversation_data, dict):
                conversation_data = json.loads(conversation_data) if isinstance(conversation_data, str) else {}
            
            messages = conversation_data.get("messages", [])
            
            # Find the last message entry or create new one
            if messages and role == "assistant" and len(messages) > 0:
                # If assistant message, update the last message entry
                last_msg = messages[-1]
                if "assistant" not in last_msg or not last_msg.get("assistant"):
                    last_msg["assistant"] = message
                else:
                    # Create new message entry
                    new_message_id = len(messages) + 1
                    messages.append({
                        "message_id": new_message_id,
                        "user": "",
                        "assistant": message
                    })
            elif messages and role == "user":
                # If user message, check if last message has user or create new
                last_msg = messages[-1] if messages else None
                if last_msg and not last_msg.get("user"):
                    last_msg["user"] = message
                else:
                    # Create new message entry
                    new_message_id = len(messages) + 1
                    messages.append({
                        "message_id": new_message_id,
                        "user": message,
                        "assistant": ""
                    })
            else:
                # First message
                messages.append({
                    "message_id": 1,
                    "user": message if role == "user" else "",
                    "assistant": message if role == "assistant" else ""
                })
            
            conversation_data["messages"] = messages
            conversation.conversation_json = conversation_data
            if user_email:
                conversation.user_email = user_email
            db.commit()
        else:
            # Create new conversation
            conversation_id = str(uuid.uuid4())
            conversation_data = {
                "conversation_id": 1,  # Will be updated to actual id after save
                "messages": [{
                    "message_id": 1,
                    "user": message if role == "user" else "",
                    "assistant": message if role == "assistant" else ""
                }]
            }
            new_conversation = Conversation(
                conversation_id=conversation_id,
                chat_id=chat_id,
                user_email=user_email,
                conversation_json=conversation_data
            )
            db.add(new_conversation)
            db.commit()
            db.refresh(new_conversation)
            # Update conversation_id in JSON to actual database id
            conversation_data["conversation_id"] = new_conversation.id
            new_conversation.conversation_json = conversation_data
            db.commit()
        
        return {
            "success": True,
            "message": "Message saved successfully"
        }
    except Exception as e:
        db.rollback()
        return {
            "success": False,
            "error": f"Error saving message: {str(e)}"
        }


@app.post("/api/chat/save-conversation")
async def save_conversation_json(
    chat_id: str = Form(...),
    conversation_json: str = Form(...),
    user_email: str = Form(None),
    db: Session = Depends(get_db)
):
    """Save full conversation as JSON for easy retrieval."""
    try:
        import json
        conversation_data = json.loads(conversation_json)
        
        # Find or create conversation record
        conversation = db.query(Conversation).filter(
            Conversation.chat_id == chat_id
        ).first()
        
        if conversation:
            # Update existing conversation
            # Ensure conversation_id is set correctly
            if "conversation_id" not in conversation_data or conversation_data["conversation_id"] != conversation.id:
                conversation_data["conversation_id"] = conversation.id
            
            conversation.conversation_json = conversation_data
            if user_email:
                conversation.user_email = user_email
            db.commit()
        else:
            # Create new conversation
            conversation_id = str(uuid.uuid4())
            # Ensure conversation_id is set in the JSON
            if "conversation_id" not in conversation_data:
                conversation_data["conversation_id"] = 1  # Will be updated after save
            
            new_conversation = Conversation(
                conversation_id=conversation_id,
                chat_id=chat_id,
                user_email=user_email,
                conversation_json=conversation_data
            )
            db.add(new_conversation)
            db.commit()
            db.refresh(new_conversation)
            # Update conversation_id in JSON to actual database id
            conversation_data["conversation_id"] = new_conversation.id
            new_conversation.conversation_json = conversation_data
            db.commit()
        
        return {
            "success": True,
            "message": "Conversation saved successfully"
        }
    except Exception as e:
        db.rollback()
        return {
            "success": False,
            "error": f"Error saving conversation: {str(e)}"
        }


@app.post("/api/chat/create")
async def create_chat_record(
    chat_id: str = Form(...),
    user_email: str = Form(None),
    db: Session = Depends(get_db)
):
    """Create a new chat record when a chat session starts."""
    try:
        # Create an initial message to establish the chat record
        # The chat_id will be used to group messages
        print(f"💾 [CHAT] Creating chat record for chat_id: {chat_id}, user: {user_email}")
        
        # The chat record is automatically created when the first message is saved
        # This endpoint just ensures the chat_id exists in the system
        return {
            "success": True,
            "message": "Chat record created successfully",
            "chat_id": chat_id
        }
    except Exception as e:
        return {
            "success": False,
            "error": f"Error creating chat record: {str(e)}"
        }


# CSV Sprint Plans endpoints
@app.get("/api/sprint-plans")
async def get_sprint_plans(user_email: str = None, workspace: str = None, filter_type: str = None, db: Session = Depends(get_db)):
    """Get sprint plans from database with optional filtering"""
    try:
        from services.db_service import db_service
        from models import SprintPlan
        
        # If filter_type is "Created by Me", show only user's plans
        if filter_type == "Created by Me":
            return db_service.get_sprint_plans_by_user(db, user_email)
        
        # If workspace is provided, filter by workspace
        if workspace and workspace != "Created by Me":
            # Query workspace by name to get ID
            from models import Workspace
            workspace_obj = db.query(Workspace).filter(Workspace.name == workspace).first()
            
            if not workspace_obj:
                return {"success": True, "plans": []}
            
            plans = db.query(SprintPlan).filter(
                SprintPlan.workspace_id == workspace_obj.id
            ).order_by(SprintPlan.created_at.desc()).all()
            
            # Convert to dictionary format
            plans_data = []
            for plan in plans:
                plan_dict = {
                    "id": plan.id,
                    "sprint_number": plan.sprint_number,
                    "sprint_dates": plan.sprint_dates,
                    "sprint_duration": plan.sprint_duration,
                    "team_name": plan.team_name,
                    "sprint_goal": plan.sprint_goal,
                    "total_hours_per_person": plan.total_hours_per_person,
                    "number_of_members": plan.number_of_members,
                    "team_members": plan.team_members,
                    "historical_story_points": plan.historical_story_points,
                    "backlog_items": plan.backlog_items,
                    "definition_of_done": plan.definition_of_done,
                    "risks_and_impediments": plan.risks_and_impediments,
                    "generated_plan": plan.generated_plan,
                    "word_document": plan.word_document,
                    "created_by": plan.created_by,
                    "workspace": workspace,
                    "workspace_id": plan.workspace_id,
                    "created_at": plan.created_at.isoformat() if plan.created_at else None
                }
                plans_data.append(plan_dict)
            
            return {"success": True, "plans": plans_data}
        
        # If user_email is provided, filter by user
        if user_email:
            return db_service.get_sprint_plans_by_user(db, user_email)
        else:
            # Return all plans (for admin purposes)
            return db_service.get_all_sprint_plans(db)
    except Exception as e:
        return {"success": False, "message": str(e)}

@app.get("/api/sprint-plans/{plan_id}")
async def get_sprint_plan(plan_id: int, db: Session = Depends(get_db)):
    """Get a specific sprint plan by ID"""
    return db_service.get_sprint_plan_by_id(db, plan_id)

@app.delete("/api/sprint-plans/{plan_id}")
async def delete_sprint_plan(plan_id: int, user_email: str, db: Session = Depends(get_db)):
    """Delete a sprint plan (only if user owns it)"""
    return db_service.delete_sprint_plan(db, plan_id, user_email)

@app.put("/api/sprint-plans/{plan_id}/edit")
async def edit_sprint_plan(plan_id: int, request: EditSprintPlanRequest, db: Session = Depends(get_db)):
    """Edit an existing sprint plan and create a new version"""
    # Format the data for the database service
    updated_data = {
        'sprint_number': request.sprint_overview.get('SprintNumber', ''),
        'sprint_dates': request.sprint_overview.get('SprintDates', ''),
        'sprint_duration': request.sprint_overview.get('SprintDuration', ''),
        'team_name': request.sprint_overview.get('TeamName', ''),
        'sprint_goal': request.sprint_overview.get('SprintGoal', ''),
        
        'total_hours_per_person': request.team_capacity.get('TotalHoursPerPerson', ''),
        'number_of_members': request.team_capacity.get('NumberofMembers', ''),
        'team_members': request.team_capacity.get('TeamMembers', []),
        'historical_story_points': request.team_capacity.get('HistoricalStoryPoints', ''),
        
        'backlog_items': request.product_backlog.get('BacklogItems', []),
        'definition_of_done': request.definition_of_done.get('DefinitionOfDone', ''),
        'risks_and_impediments': request.risks_and_impediments.get('RisksAndImpediments', '')
    }
    
    return db_service.edit_sprint_plan(db, plan_id, updated_data, request.edit_comments, request.edited_by)

@app.get("/api/sprint-plans/{plan_id}/versions")
async def get_plan_versions(plan_id: int, db: Session = Depends(get_db)):
    """Get all versions of a specific sprint plan"""
    return db_service.get_plan_versions(db, plan_id)

# PDF download endpoint removed - PDF generation now handled by frontend html2pdf.js

@app.post("/api/sprint-plans/{plan_id}/download-version")
async def download_plan_version(plan_id: int, request: DownloadVersionRequest, db: Session = Depends(get_db)):
    """Get a specific version of a sprint plan for download"""
    versions_result = db_service.get_plan_versions(db, plan_id)
    if not versions_result.get("success"):
        return versions_result
    
    # Find the requested version
    requested_version = None
    for version in versions_result.get("versions", []):
        if version.get("version_number") == request.version_number:
            requested_version = version
            break
    
    if not requested_version:
        return {"success": False, "message": f"Version {request.version_number} not found"}
    
    return {"success": True, "version": requested_version}

@app.post("/api/sprint/edit-plan")
async def edit_sprint_plan(request: EditSprintPlanRequest, db: Session = Depends(get_db)):
    """Edit and regenerate a sprint plan based on new requirements"""
    try:
        print("=" * 80)
        print("✏️ [EDIT SPRINT PLAN] REQUEST RECEIVED")
        print("=" * 80)
        print(f"📝 [EDIT PLAN] Edit Comments: {request.edit_comments}")
        print(f"👤 [EDIT PLAN] Edited By: {request.edited_by}")
        print("=" * 80)
        
        from services.gemini_service import gemini_service
        from services.db_service import db_service
        
        # Get stored prompt from global variable or fetch from DB
        stored_prompt = get_global_prompt_data()
        if not stored_prompt:
            # Fallback: fetch from database
            print("🔍 [EDIT PLAN] Global prompt is None, fetching from database...")
            result = db.execute(text("SELECT id, prompt, feature FROM documents WHERE feature = 'SprintPlan' LIMIT 1"))
            row = result.fetchone()
            if row:
                stored_prompt = row[1]
                # Update global variable for future use
                global GLOBAL_PROMPT_DATA
                GLOBAL_PROMPT_DATA = stored_prompt
                print(f"🔍 [EDIT PLAN] Loaded prompt from database: {len(stored_prompt)} characters")
            else:
                print("❌ [EDIT PLAN] No prompt found in database!")
                return {"success": False, "message": "No prompt found in database"}
        
        # Create enhanced prompt with edit requirements
        enhanced_prompt = f"""
{stored_prompt}

IMPORTANT EDIT REQUIREMENTS:
The user has requested the following changes to the previously generated sprint plan:

"{request.edit_comments}"

Please regenerate the sprint plan incorporating these specific requirements. 
Make sure the new plan addresses all the user's edit requests while maintaining 
the professional structure and format of the original plan.

User Inputs for Sprint Plan:
"""
        
        # Format user inputs into prompt template
        user_inputs = {
            "sprint_overview": request.sprint_overview,
            "team_capacity": request.team_capacity,
            "product_backlog": request.product_backlog,
            "definition_of_done": request.definition_of_done,
            "risks_and_impediments": request.risks_and_impediments,
            "additional_comments": request.additional_comments
        }
        
        # Format team members and backlog items
        team_members_text = ""
        if user_inputs['team_capacity'].get('TeamMembers'):
            team_members_text = "\n".join([
                f"- {member.get('roleName', 'N/A')}: {member.get('workingHours', 'N/A')}"
                for member in user_inputs['team_capacity'].get('TeamMembers', [])
            ])

        backlog_items_text = ""
        if user_inputs['product_backlog'].get('BacklogItems'):
            backlog_items_text = "\n".join([
                f"- {item.get('userStorySummary', 'N/A')} (Priority: {item.get('priority', 'N/A')}, Effort: {item.get('effortEstimate', 0)} hours)"
                for item in user_inputs['product_backlog'].get('BacklogItems', [])
            ])

        # Build the complete prompt
        complete_prompt = f"""
{enhanced_prompt}

SPRINT OVERVIEW:
- Sprint Number: {user_inputs['sprint_overview'].get('SprintNumber', 'N/A')}
- Sprint Dates: {user_inputs['sprint_overview'].get('SprintDates', 'N/A')}
- Team Name: {user_inputs['sprint_overview'].get('TeamName', 'N/A')}
- Sprint Goal: {user_inputs['sprint_overview'].get('SprintGoal', 'N/A')}

TEAM CAPACITY:
- Total Hours per Person: {user_inputs['team_capacity'].get('TotalHoursPerPerson', 'N/A')}
- Number of Members: {user_inputs['team_capacity'].get('NumberOfMembers', 'N/A')}
- Historical Story Points: {user_inputs['team_capacity'].get('HistoricalStoryPoints', 'N/A')}
- Team Members:
{team_members_text if team_members_text else '  - No team members specified'}

PRODUCT BACKLOG ITEMS:
{backlog_items_text if backlog_items_text else '- No backlog items specified'}

DEFINITION OF DONE:
{user_inputs['definition_of_done'].get('DoDContent', 'N/A')}

RISKS & IMPEDIMENTS:
{user_inputs['risks_and_impediments'].get('RisksContent', 'N/A')}

ADDITIONAL COMMENTS:
{user_inputs['additional_comments'].get('CommentsContent', 'N/A')}

Remember to incorporate the edit requirements: "{request.edit_comments}"
"""
        
        print(f"🔍 [EDIT PLAN] Enhanced prompt length: {len(complete_prompt)} characters")
        print(f"🔍 [EDIT PLAN] Prompt preview: {complete_prompt[:300]}...")
        
        # Generate new plan using LLM service - use the EXACT same format as main plan generation
        messages = [
            {"role": "system", "content": stored_prompt},           # Base prompt from documents table
            {"role": "user", "content": complete_prompt}            # User inputs + edit requirements
        ]
        
        print("🚀 [EDIT PLAN] Calling LLM service to regenerate plan...")
        llm_response = gemini_service.chat(messages, max_tokens=4000)
        
        # Check if LLM service actually succeeded
        if not llm_response:
            print("❌ [EDIT PLAN] LLM service returned no response")
            return {"success": False, "message": "Failed to generate new plan from LLM"}
        
        if not llm_response.get('success', False):
            error_msg = llm_response.get('response', 'Unknown error from LLM service')
            print(f"❌ [EDIT PLAN] LLM service failed: {error_msg}")
            return {"success": False, "message": f"LLM service failed: {error_msg}"}
        
        if not llm_response.get('response'):
            print("❌ [EDIT PLAN] LLM service returned empty response")
            return {"success": False, "message": "Failed to generate new plan from LLM"}
        
        new_generated_plan = llm_response['response']
        print(f"✅ [EDIT PLAN] New plan generated successfully, length: {len(new_generated_plan)} characters")
        
        # Count expected Product Backlog Items from user input
        expected_pb_count = len(user_inputs['product_backlog'].get('BacklogItems', []))
        print(f"🔍 [EDIT PB COUNT] Expected Product Backlog Items: {expected_pb_count}")
        
        # STEP 2: VALIDATION AND FINE-TUNING FOR EDITED PLAN PB COMPLETENESS
        print("🔍 [EDIT VALIDATION] Starting PB completeness validation for edited plan...")
        
        # Pre-validation check for missing sections in edited plan
        missing_sections = []
        
        if "Detailed Task Breakdown" not in new_generated_plan:
            missing_sections.append("Detailed Task Breakdown")
        if "Committed Sprint Backlog" not in new_generated_plan:
            missing_sections.append("Committed Sprint Backlog")
        if "Sprint Overview" not in new_generated_plan:
            missing_sections.append("Sprint Overview")
        if "Team Capacity" not in new_generated_plan:
            missing_sections.append("Team Capacity")
        
        if missing_sections:
            print(f"❌ [EDIT PRE-VALIDATION] Missing critical sections: {missing_sections}")
            print("🔄 [EDIT PRE-VALIDATION] Edited plan will be regenerated to include missing sections")
        else:
            print("✅ [EDIT PRE-VALIDATION] All critical sections present")
        
        # Create user inputs text for validation
        user_inputs_text = f"""
I. Sprint Overview & Proposed Goal:
- Sprint Number: {user_inputs['sprint_overview'].get('SprintNumber', 'N/A')}
- Sprint Dates: {user_inputs['sprint_overview'].get('SprintDates', 'N/A')}
- Sprint Duration: {user_inputs['sprint_overview'].get('SprintDuration', 'N/A')}
- Team Name: {user_inputs['sprint_overview'].get('TeamName', 'N/A')}
- Sprint Goal: {user_inputs['sprint_overview'].get('SprintGoal', 'N/A')}

II. Team Capacity & Availability:
- Total Hours per Person: {user_inputs['team_capacity'].get('TotalHoursPerPerson', 'N/A')}
- Number of Members: {user_inputs['team_capacity'].get('NumberOfMembers', 'N/A')}
- Team Members: {team_members_text}
- Historical Story Points: {user_inputs['team_capacity'].get('HistoricalStoryPoints', 'N/A')}

III. Prioritized Product Backlog Items:
{backlog_items_text}

IV. Definition of Done (DoD):
{user_inputs['definition_of_done'].get('DoDContent', 'N/A')}

V. Known Impediments, Dependencies & Risks:
{user_inputs['risks_and_impediments'].get('RisksContent', 'N/A')}

VI. Additional Comments:
{user_inputs['additional_comments'].get('CommentsContent', 'N/A')}
"""
        
        validation_result = gemini_service.validate_and_finetune_sprint_plan(
            original_plan=new_generated_plan,
            user_inputs=user_inputs_text,
            stored_prompt=stored_prompt,
            expected_pb_count=expected_pb_count
        )
        
        if validation_result.get("success"):
            if validation_result.get("improved"):
                print(f"🔄 [EDIT VALIDATION] Edited plan was regenerated to include all {expected_pb_count} PBs")
                new_generated_plan = validation_result.get("response", "")
            else:
                print(f"✅ [EDIT VALIDATION] Edited plan passed validation with all {expected_pb_count} PBs")
            
            print(f"🔍 [EDIT VALIDATION] Final edited plan length: {len(new_generated_plan)} characters")
        else:
            print(f"⚠️ [EDIT VALIDATION] Validation failed, using original edited plan: {validation_result.get('validation_error', 'Unknown error')}")
        
        print("🔍 [EDIT VALIDATION] PB completeness validation completed")
        
        # Generate Word document content for the new plan
        print("📝 [EDIT PLAN] Generating Word document content for new plan...")
        word_document_content = generate_word_document_content(new_generated_plan, user_inputs)
        print(f"📝 [EDIT PLAN] Word document generated: {len(word_document_content)} characters")
        
        # Prepare data for saving new plan
        plan_data = {
            'sprint_number': user_inputs['sprint_overview'].get('SprintNumber'),
            'sprint_dates': user_inputs['sprint_overview'].get('SprintDates'),
            'sprint_duration': user_inputs['sprint_overview'].get('SprintDuration'),
            'team_name': user_inputs['sprint_overview'].get('TeamName'),
            'sprint_goal': user_inputs['sprint_overview'].get('SprintGoal'),
            'total_hours_per_person': user_inputs['team_capacity'].get('TotalHoursPerPerson'),
            'number_of_members': user_inputs['team_capacity'].get('NumberOfMembers'),
            'team_members': user_inputs['team_capacity'].get('TeamMembers'),
            'historical_story_points': user_inputs['team_capacity'].get('HistoricalStoryPoints'),
            'backlog_items': user_inputs['product_backlog'].get('BacklogItems'),
            'definition_of_done': user_inputs['definition_of_done'].get('DoDContent'),
            'risks_and_impediments': user_inputs['risks_and_impediments'].get('RisksContent'),
            'generated_plan': new_generated_plan,
            'word_document': word_document_content,
            'user_email': request.edited_by
        }
        
        # First, delete old plans with same sprint number and team name
        print("🗑️ [EDIT PLAN] Replacing old plans...")
        replace_result = db_service.find_and_delete_old_plan(
            db, 
            request.edited_by, 
            plan_data['sprint_number'], 
            plan_data['team_name']
        )
        
        if not replace_result['success']:
            print(f"⚠️ [EDIT PLAN] Warning: Could not replace old plans: {replace_result['message']}")
        else:
            print(f"✅ [EDIT PLAN] {replace_result['message']}")
        
        # Save new plan to database
        print("💾 [EDIT PLAN] Saving new plan to database...")
        save_result = db_service.save_sprint_plan(db, plan_data, request.edited_by)
        
        if not save_result['success']:
            print(f"❌ [EDIT PLAN] Failed to save new plan: {save_result['message']}")
            return {"success": False, "message": f"Failed to save new plan: {save_result['message']}"}
        
        print(f"✅ [EDIT PLAN] New plan saved successfully with ID: {save_result['plan_id']}")
        
        # Return the new plan data in the format expected by frontend
        return {
            "success": True,
            "message": "Sprint plan edited and regenerated successfully",
            "plan": {
                "id": save_result['plan_id'],
                "response": new_generated_plan,
                "summary": new_generated_plan,
                "sprint_number": plan_data['sprint_number'],
                "sprint_dates": plan_data['sprint_dates'],
                "team_name": plan_data['team_name']
            }
        }
        
    except Exception as e:
        print(f"❌ [EDIT PLAN] Error: {str(e)}")
        import traceback
        traceback.print_exc()
        return {"success": False, "message": f"Error editing sprint plan: {str(e)}"}

# Risk Assessment endpoints
@app.post("/api/risk-assessment/start", response_model=RiskAssessmentStartResponse)
async def start_risk_assessment(
    request: RiskAssessmentStartRequest,
    db: Session = Depends(get_db)
):
    """Start a new risk assessment session"""
    return risk_service.start_risk_assessment(request, db)

@app.get("/api/risk-assessments")
async def get_risk_assessments(user_email: str = None, workspace: str = None, filter_type: str = None, db: Session = Depends(get_db)):
    """Get risk assessments from database with optional filtering"""
    try:
        from services.db_service import db_service
        from models import RiskAssessment, Workspace
        
        # If filter_type is "Created by Me", show only user's assessments
        if filter_type == "Created by Me":
            return db_service.get_risk_assessments_by_user(db, user_email)
        
        # If workspace is provided, filter by workspace
        if workspace and workspace != "Created by Me":
            # Query workspace by name to get ID
            workspace_obj = db.query(Workspace).filter(Workspace.name == workspace).first()
            
            if not workspace_obj:
                return {"success": True, "assessments": []}
            
            assessments = db.query(RiskAssessment).filter(
                RiskAssessment.workspace_id == workspace_obj.id
            ).order_by(RiskAssessment.created_at.desc()).all()
            
            # Convert to dictionary format
            assessments_data = []
            for assessment in assessments:
                assessment_dict = {
                    "id": assessment.id,
                    "project_name": assessment.project_name,
                    "project_dates": assessment.project_dates,
                    "project_duration": assessment.project_duration,
                    "team_name": assessment.team_name,
                    "project_scope": assessment.project_scope,
                    "risk_categories": assessment.risk_categories,
                    "risk_mitigation": assessment.risk_mitigation,
                    "risk_monitoring": assessment.risk_monitoring,
                    "stakeholders": assessment.stakeholders,
                    "risk_matrix": assessment.risk_matrix,
                    "risk_register": assessment.risk_register,
                    "generated_assessment": assessment.generated_assessment,
                    "word_document": assessment.word_document,
                    "created_by": assessment.created_by,
                    "workspace": workspace,
                    "workspace_id": assessment.workspace_id,
                    "created_at": assessment.created_at.isoformat() if assessment.created_at else None
                }
                assessments_data.append(assessment_dict)
            
            return {"success": True, "assessments": assessments_data}
        
        # If user_email is provided, filter by user
        if user_email:
            return db_service.get_risk_assessments_by_user(db, user_email)
        else:
            # Return all assessments (for admin purposes)
            return db_service.get_all_risk_assessments(db)
    except Exception as e:
        return {"success": False, "message": str(e)}

@app.get("/api/risk-assessments/{assessment_id}")
async def get_risk_assessment(assessment_id: int, db: Session = Depends(get_db)):
    """Get a specific risk assessment by ID"""
    return db_service.get_risk_assessment_by_id(db, assessment_id)

@app.delete("/api/risk-assessments/{assessment_id}")
async def delete_risk_assessment(assessment_id: int, user_email: str, db: Session = Depends(get_db)):
    """Delete a risk assessment (only if user owns it)"""
    return db_service.delete_risk_assessment(db, assessment_id, user_email)

@app.post("/api/risk-assessment/chat", response_model=RiskAssessmentChatResponse)
async def chat_with_llm_risk(
    request: RiskAssessmentChatRequest,
    db: Session = Depends(get_db)
):
    """Send message to LLM and get response for risk assessment"""
    # Fetch RiskAssessment prompt from database
    result = db.execute(text("SELECT id, prompt, feature FROM documents WHERE feature = 'RiskAssessment' LIMIT 1"))
    row = result.fetchone()
    if row:
        risk_prompt = row[1]
        print(f"🔍 [RISK CHAT] Loaded RiskAssessment prompt from database: {len(risk_prompt)} characters")
    else:
        print("❌ [RISK CHAT] No RiskAssessment prompt found in database!")
        return {"success": False, "message": "No RiskAssessment prompt found in database"}
    
    return risk_service.chat_with_llm(request, db, risk_prompt)

@app.post("/api/risk-assessment/finish", response_model=RiskAssessmentFinishResponse)
async def finish_risk_assessment(
    request: RiskAssessmentFinishRequest,
    db: Session = Depends(get_db)
):
    """Complete risk assessment and get summary"""
    # Fetch RiskAssessment prompt from database
    result = db.execute(text("SELECT id, prompt, feature FROM documents WHERE feature = 'RiskAssessment' LIMIT 1"))
    row = result.fetchone()
    if row:
        risk_prompt = row[1]
        print(f"🔍 [RISK FINISH] Loaded RiskAssessment prompt from database: {len(risk_prompt)} characters")
    else:
        print("❌ [RISK FINISH] No RiskAssessment prompt found in database!")
        return {"success": False, "message": "No RiskAssessment prompt found in database"}
    
    return risk_service.finish_risk_assessment(request, db, risk_prompt)

@app.post("/api/risk-assessment/generate-assessment")
async def generate_risk_assessment(request: GenerateRiskAssessmentRequest, db: Session = Depends(get_db)):
    """Generate risk assessment with new structure"""
    try:
        import json
        
        # Log the complete JSON request data
        print("=" * 80)
        print("🚀 [GENERATE RISK ASSESSMENT] REQUEST RECEIVED")
        print("=" * 80)
        print("📋 [GENERATE RISK ASSESSMENT] Complete JSON Request Data:")
        print("=" * 80)
        
        # Convert request to dict for logging
        request_dict = {
            "project_overview": request.project_overview,
            "risk_categories": request.risk_categories,
            "stakeholders": request.stakeholders,
            "risk_matrix": request.risk_matrix,
            "risk_register": request.risk_register,
            "additional_comments": request.additional_comments,
            "all_risks_data": request.all_risks_data
        }
        
        # Pretty print the JSON
        print(json.dumps(request_dict, indent=2, ensure_ascii=False))
        print("=" * 80)
        
        from services.gemini_service import gemini_service
        from services.db_service import db_service
        
        # Get stored prompt from global variable or fetch from DB
        stored_prompt = get_global_prompt_data()
        if not stored_prompt:
            # Fallback: fetch from database
            print("🔍 [PROMPT LOADING] Global prompt is None, fetching from database...")
            result = db.execute(text("SELECT id, prompt, feature FROM documents WHERE feature = 'RiskAssessment' LIMIT 1"))
            row = result.fetchone()
            if row:
                stored_prompt = row[1]
                # Update global variable for future use
                global GLOBAL_PROMPT_DATA
                GLOBAL_PROMPT_DATA = stored_prompt
                print(f"🔍 [PROMPT LOADING] Loaded prompt from database: {len(stored_prompt)} characters")
            else:
                print("❌ [PROMPT LOADING] No prompt found in database!")
                return {"success": False, "message": "No prompt found in database"}
        else:
            print(f"🔍 [PROMPT LOADING] Using global prompt: {len(stored_prompt)} characters")
        
        # Fetch RiskAssessment prompt from database
        print("🔍 [PROMPT LOADING] Fetching RiskAssessment prompt from database...")
        result = db.execute(text("SELECT id, prompt, feature FROM documents WHERE feature = 'RiskAssessment' LIMIT 1"))
        row = result.fetchone()
        if row:
            stored_prompt = row[1]
            print(f"🔍 [PROMPT LOADING] Loaded RiskAssessment prompt from database: {len(stored_prompt)} characters")
        else:
            print("❌ [PROMPT LOADING] No RiskAssessment prompt found in database!")
            return {"success": False, "message": "No RiskAssessment prompt found in database"}
        
        print(f"🔍 [GENERATE ASSESSMENT] Final prompt length: {len(stored_prompt)} characters")
        print(f"🔍 [GENERATE ASSESSMENT] Prompt preview: {stored_prompt[:200]}...")
        
        # Validate that we have a proper prompt
        if not stored_prompt or len(stored_prompt.strip()) < 50:
            print("❌ [PROMPT VALIDATION] Prompt is too short or empty!")
            return {"success": False, "message": "Invalid prompt data. Please ensure prompt is properly loaded."}
        
        # Prepare user inputs as JSON for the database prompt to handle
        user_inputs = {
            "project_overview": request.project_overview,
            "risk_categories": request.risk_categories,
            "stakeholders": request.stakeholders,
            "risk_matrix": request.risk_matrix,
            "risk_register": request.risk_register,
            "additional_comments": request.additional_comments,
            "all_risks_data": request.all_risks_data  # Include the all_risks_data field
        }
        
        # Convert user inputs to JSON string for the prompt
        import json
        user_inputs_json = json.dumps(user_inputs, indent=2, ensure_ascii=False)
        
        print(f"🔍 [GENERATE ASSESSMENT] User inputs JSON length: {len(user_inputs_json)}")
        print(f"🔍 [GENERATE ASSESSMENT] User inputs JSON: {user_inputs_json}")
        
        # Call Gemini service with database prompt handling the formatting
        messages = [
            {"role": "system", "content": stored_prompt},
            {"role": "user", "content": f"Please generate a comprehensive risk assessment based on the following project data:\n\n{user_inputs_json}"}
        ]
        
        print(f"🔍 [GENERATE ASSESSMENT] Messages being sent to Gemini:")
        print(f"   - System message length: {len(stored_prompt)}")
        print(f"   - User message length: {len(messages[1]['content'])}")
        print(f"   - User message content: {messages[1]['content'][:200]}...")
        
        print("🔍 [GEMINI CALL] Calling Gemini service with messages:")
        print(f"   - System message length: {len(stored_prompt)}")
        print(f"   - User message length: {len(messages[1]['content'])}")
        
        gemini_response = gemini_service.chat(messages, max_tokens=4000)
        
        print("🔍 [GEMINI CALL] Gemini service response received:")
        print(f"   - Response object: {gemini_response}")
        print(f"   - Response type: {type(gemini_response)}")
        print(f"   - Response keys: {list(gemini_response.keys()) if isinstance(gemini_response, dict) else 'Not a dict'}")
        
        # Check if Gemini service actually succeeded
        if not gemini_response:
            print("❌ [GEMINI CALL] Gemini service returned no response")
            return {"success": False, "message": "Failed to generate risk assessment from Gemini"}
        
        if not gemini_response.get('success', False):
            error_msg = gemini_response.get('response', 'Unknown error from Gemini service')
            print(f"❌ [GEMINI CALL] Gemini service failed: {error_msg}")
            return {"success": False, "message": f"Gemini service failed: {error_msg}"}
        
        if not gemini_response.get('response'):
            print("❌ [GEMINI CALL] Gemini service returned empty response")
            return {"success": False, "message": "Failed to generate risk assessment from Gemini"}
        
        if gemini_response.get("success"):
            print("🔍 [STEP 1] Initial risk assessment generation completed successfully")
            
            # Count expected risks from user input
            expected_risk_count = len(user_inputs.get('all_risks_data', []))
            print(f"🔍 [RISK COUNT] Expected risks: {expected_risk_count}")
            
            # STEP 2: VALIDATION AND FINE-TUNING FOR RISK COMPLETENESS
            print("🔍 [STEP 2] Starting risk completeness validation and fine-tuning...")
            
            # Pre-validation check for risk count
            original_assessment = gemini_response.get("response", "")
            
            # Count risks in the original assessment
            risk_count_in_output = original_assessment.count("**Risk ID:**")
            print(f"🔍 [PRE-VALIDATION] Risks found in output: {risk_count_in_output}")
            print(f"🔍 [PRE-VALIDATION] Expected risks: {expected_risk_count}")
            
            if risk_count_in_output < expected_risk_count:
                print(f"❌ [PRE-VALIDATION] Missing risks: {expected_risk_count - risk_count_in_output}")
                print("🔄 [PRE-VALIDATION] Assessment will be regenerated to include all risks")
            else:
                print("✅ [PRE-VALIDATION] All risks present")
            
            # Create user inputs text for validation
            user_inputs_text = f"""
PROJECT OVERVIEW:
- Project Name: {user_inputs['project_overview'].get('ProjectName', 'N/A')}
- Project Dates: {user_inputs['project_overview'].get('ProjectDates', 'N/A')}
- Project Duration: {user_inputs['project_overview'].get('ProjectDuration', 'N/A')}
- Team Name: {user_inputs['project_overview'].get('TeamName', 'N/A')}
- Project Scope: {user_inputs['project_overview'].get('ProjectScope', 'N/A')}

RISK CATEGORIES:
{user_inputs['risk_categories'].get('RiskCategories', [])}

STAKEHOLDERS:
{user_inputs['stakeholders'].get('Stakeholders', [])}

RISK MATRIX:
{user_inputs['risk_matrix'].get('RiskMatrixContent', 'N/A')}

RISK REGISTER:
{user_inputs['risk_register'].get('RiskRegisterContent', 'N/A')}

ALL RISKS DATA:
{user_inputs.get('all_risks_data', [])}

ADDITIONAL COMMENTS:
{user_inputs['additional_comments'].get('CommentsContent', 'N/A')}
"""
            
            validation_result = gemini_service.validate_and_finetune_risk_assessment(
                original_assessment=original_assessment,
                user_inputs=user_inputs_text,
                stored_prompt=stored_prompt,
                expected_risk_count=expected_risk_count
            )
            
            if validation_result.get("success"):
                print("✅ [STEP 2] Risk assessment validation completed successfully")
                final_assessment = validation_result.get("response", original_assessment)
            else:
                print(f"⚠️ [STEP 2] Validation failed, using original assessment: {validation_result.get('error', 'Unknown error')}")
                final_assessment = original_assessment
            
            # CRITICAL: Store final validated assessment content
            generated_assessment_backup = final_assessment
            
            # Clean up any unwanted markdown tags
            import re
            generated_assessment_backup = re.sub(r'```html\s*', '', generated_assessment_backup, flags=re.IGNORECASE)
            generated_assessment_backup = re.sub(r'```\s*$', '', generated_assessment_backup, flags=re.IGNORECASE)
            generated_assessment_backup = generated_assessment_backup.strip()
            
            print(f"🔍 [FINAL ASSESSMENT] Final validated assessment created: {len(generated_assessment_backup)} characters")
            print(f"🔍 [FINAL ASSESSMENT] Final assessment preview: {generated_assessment_backup[:200]}...")
            
            # Prepare assessment data for saving
            assessment_data = {
                "project_name": user_inputs['project_overview'].get('ProjectName', ''),
                "project_dates": user_inputs['project_overview'].get('ProjectDates', ''),
                "project_duration": user_inputs['project_overview'].get('ProjectDuration', ''),
                "team_name": user_inputs['project_overview'].get('TeamName', ''),
                "project_scope": user_inputs['project_overview'].get('ProjectScope', ''),
                "risk_categories": user_inputs['risk_categories'].get('RiskCategories', []),
                "risk_mitigation": user_inputs['risk_categories'].get('RiskMitigation', ''),
                "risk_monitoring": user_inputs['risk_categories'].get('RiskMonitoring', ''),
                "stakeholders": user_inputs['stakeholders'].get('Stakeholders', []),
                "risk_matrix": user_inputs['risk_matrix'].get('RiskMatrixData', {}),
                "risk_register": user_inputs['risk_register'].get('RiskRegisterData', {}),
                "generated_assessment": generated_assessment_backup,
                "word_document": generated_assessment_backup,  # For now, use the same content
                "workspace_id": request.workspace_id  # Add workspace_id from request
            }
            
            # Generate Word document content (replica of HTML rendered output)
            print("📝 [WORD GENERATION] Generating Word document content...")
            word_document_content = generate_risk_assessment_word_document(
                assessment_data.get('generated_assessment', ''),
                user_inputs
            )
            assessment_data['word_document'] = word_document_content
            print(f"📝 [WORD GENERATION] Word document generated: {len(word_document_content)} characters")
            
            # Save to database with new structure
            db_result = db_service.save_risk_assessment(db, assessment_data, request.user_email)
            
            if db_result.get("success"):
                return {
                    "success": True,
                    "message": "Risk assessment generated, validated, and saved successfully",
                    "response": final_assessment,  # Use the validated assessment
                    "assessment_id": db_result.get("assessment_id")
                }
            else:
                print(f"❌ [GENERATE ASSESSMENT] Database save failed: {db_result.get('message')}")
                return {
                    "success": False,
                    "message": f"Failed to save risk assessment to database: {db_result.get('message')}"
                }
        else:
            return {
                "success": False,
                "message": "Failed to generate risk assessment",
                "error": gemini_response.get("error", "Unknown error")
            }
            
    except Exception as e:
        print(f"❌ [GENERATE ASSESSMENT] Error: {str(e)}")
        return {"success": False, "message": f"Error generating risk assessment: {str(e)}"}

def generate_risk_assessment_word_document(generated_assessment: str, user_inputs: dict) -> str:
    """Generate Word document content that is a replica of HTML rendered output for risk assessment"""
    try:
        # Remove markdown code blocks but keep HTML formatting
        html_content = generated_assessment or ''
        import re
        html_content = re.sub(r'```html\s*', '', html_content, flags=re.IGNORECASE)
        html_content = re.sub(r'```\s*$', '', html_content, flags=re.IGNORECASE)
        html_content = html_content.strip()
        
        # If the content doesn't start with <!DOCTYPE or <html, wrap it properly
        if not html_content.startswith('<!DOCTYPE') and not html_content.startswith('<html'):
            html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Risk Assessment - {user_inputs.get('project_overview', {}).get('ProjectName', 'N/A')}</title>
</head>
<body>
{html_content}
</body>
</html>"""
        
        if not html_content:
            return '<p>No content available for Word document.</p>'
        
        # Create Word-compatible HTML with all formatting preserved
        word_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>Risk Assessment - {user_inputs.get('project_overview', {}).get('ProjectName', 'N/A')}</title>
            <style>
                body {{
                    font-family: 'Calibri', 'Arial', sans-serif;
                    margin: 20px;
                    padding: 0;
                    background: white;
                    color: #2c3e50;
                    line-height: 1.6;
                    font-size: 11pt;
                }}
                h1, h2, h3, h4 {{
                    color: #2d3748;
                    margin: 20px 0 15px 0;
                    font-weight: 600;
                }}
                h1 {{ font-size: 18pt; }}
                h2 {{ font-size: 16pt; }}
                h3 {{ font-size: 14pt; }}
                h4 {{ font-size: 12pt; }}
                p {{ margin: 10px 0; }}
                table {{
                    width: 100%;
                    border-collapse: collapse;
                    margin: 20px 0;
                    border: 1px solid #ddd;
                }}
                th, td {{
                    border: 1px solid #ddd;
                    padding: 8px 12px;
                    text-align: left;
                    vertical-align: top;
                }}
                th {{
                    background-color: #f8f9fa;
                    font-weight: bold;
                    color: #2d3748;
                }}
                ul, ol {{ 
                    margin: 15px 0; 
                    padding-left: 30px; 
                }}
                li {{ 
                    margin: 8px 0; 
                    line-height: 1.6;
                }}
                strong {{
                    color: #2d3748;
                    font-weight: 600;
                }}
                em {{
                    color: #718096;
                    font-style: italic;
                }}
                .assessment-header {{
                    background: #f7fafc;
                    border: 2px solid #e2e8f0;
                    border-radius: 8px;
                    padding: 20px;
                    margin-bottom: 20px;
                    text-align: center;
                }}
                .assessment-header h1 {{
                    margin: 0 0 10px 0;
                    color: #1a202c;
                }}
                .assessment-header p {{
                    margin: 0;
                    color: #4a5568;
                    font-size: 12pt;
                }}
            </style>
        </head>
        <body>
            <div class="assessment-header">
                <h1>Risk Assessment - {user_inputs.get('project_overview', {}).get('ProjectName', 'N/A')}</h1>
                <p>Generated on {__import__('datetime').datetime.now().strftime('%Y-%m-%d')}</p>
            </div>
            
            {html_content}
            
            <div style="margin-top: 40px; padding-top: 20px; border-top: 1px solid #e2e8f0; text-align: center; color: #718096; font-size: 10pt;">
                <p>This risk assessment was automatically generated by the Risk Assessment System</p>
                <p>Document ID: RA-{user_inputs.get('project_overview', {}).get('ProjectName', 'N/A')}-{__import__('datetime').datetime.now().strftime('%Y-%m-%d')}</p>
            </div>
        </body>
        </html>
        """
        
        return word_html.strip()
        
    except Exception as e:
        print(f"❌ [WORD GENERATION] Error generating Word document: {str(e)}")
        return f'<p>Error generating Word document: {str(e)}</p>'

# Risk Assessment CRUD endpoints
@app.get("/api/risk-assessments")
async def get_risk_assessments(user_email: str = None, db: Session = Depends(get_db)):
    """Get risk assessments from database"""
    try:
        from services.db_service import db_service
        
        # If user_email is provided, filter by user
        if user_email:
            return db_service.get_risk_assessments_by_user(db, user_email)
        else:
            # Return all assessments (for admin purposes)
            return db_service.get_all_risk_assessments(db)
    except Exception as e:
        return {"success": False, "message": str(e)}

@app.get("/api/risk-assessments/{assessment_id}")
async def get_risk_assessment(assessment_id: int, db: Session = Depends(get_db)):
    """Get a specific risk assessment by ID"""
    return db_service.get_risk_assessment_by_id(db, assessment_id)

@app.delete("/api/risk-assessments/{assessment_id}")
async def delete_risk_assessment(assessment_id: int, user_email: str, db: Session = Depends(get_db)):
    """Delete a risk assessment (only if user owns it)"""
    return db_service.delete_risk_assessment(db, assessment_id, user_email)

# Feedback endpoints
@app.post("/api/feedback", response_model=FeedbackResponse)
async def submit_feedback(feedback_data: FeedbackRequest, db: Session = Depends(get_db)):
    """Submit feedback form data"""
    try:
        print(f"📝 [FEEDBACK] Received feedback submission")
        print(f"📝 [FEEDBACK] Data: {feedback_data.dict()}")
        
        # Create new feedback record
        feedback = Feedback(
            name=feedback_data.name,
            email=feedback_data.email,
            clarity_of_sprint_goals=feedback_data.clarity_of_sprint_goals,
            workload_distribution=feedback_data.workload_distribution,
            plan_alignment_sow=feedback_data.plan_alignment_sow,
            suggestions_sprint_planning=feedback_data.suggestions_sprint_planning,
            risks_clear=feedback_data.risks_clear,
            mitigation_practical=feedback_data.mitigation_practical,
            suggestions_risk_assessment=feedback_data.suggestions_risk_assessment,
            overall_sprint_planning_rating=feedback_data.overall_sprint_planning_rating,
            overall_risk_assessment_rating=feedback_data.overall_risk_assessment_rating,
            additional_comments=feedback_data.additional_comments,
            created_by=feedback_data.user_email or feedback_data.email
        )
        
        # Add to database
        db.add(feedback)
        db.commit()
        db.refresh(feedback)
        
        print(f"✅ [FEEDBACK] Successfully saved feedback with ID: {feedback.id}")
        
        return FeedbackResponse(
            success=True,
            message="Feedback submitted successfully!",
            feedback_id=feedback.id
        )
        
    except Exception as e:
        print(f"❌ [FEEDBACK] Error submitting feedback: {str(e)}")
        db.rollback()
        return FeedbackResponse(
            success=False,
            message=f"Error submitting feedback: {str(e)}"
        )

@app.get("/api/feedback")
async def get_feedback(user_email: str = None, db: Session = Depends(get_db)):
    """Get feedback submissions (admin or user-specific)"""
    try:
        if user_email:
            # Get feedback for specific user
            feedback = db.query(Feedback).filter(Feedback.created_by == user_email).all()
        else:
            # Get all feedback (admin)
            feedback = db.query(Feedback).all()
        
        return {
            "success": True,
            "feedback": [
                {
                    "id": f.id,
                    "name": f.name,
                    "email": f.email,
                    "clarity_of_sprint_goals": f.clarity_of_sprint_goals,
                    "workload_distribution": f.workload_distribution,
                    "plan_alignment_sow": f.plan_alignment_sow,
                    "suggestions_sprint_planning": f.suggestions_sprint_planning,
                    "risks_clear": f.risks_clear,
                    "mitigation_practical": f.mitigation_practical,
                    "suggestions_risk_assessment": f.suggestions_risk_assessment,
                    "overall_sprint_planning_rating": f.overall_sprint_planning_rating,
                    "overall_risk_assessment_rating": f.overall_risk_assessment_rating,
                    "additional_comments": f.additional_comments,
                    "created_at": f.created_at,
                    "created_by": f.created_by
                }
                for f in feedback
            ]
        }
    except Exception as e:
        return {"success": False, "message": str(e)}

# Workspace endpoints
@app.get("/api/workspaces")
async def get_workspaces(db: Session = Depends(get_db)):
    """Get all workspaces"""
    try:
        from models import Workspace
        
        workspaces = db.query(Workspace).order_by(Workspace.is_default.desc(), Workspace.name).all()
        
        return {
            "success": True,
            "workspaces": [
                {
                    "id": ws.id,
                    "name": ws.name,
                    "description": ws.description,
                    "is_default": ws.is_default,
                    "created_at": ws.created_at.isoformat() if ws.created_at else None
                }
                for ws in workspaces
            ]
        }
    except Exception as e:
        return {"success": False, "message": str(e)}

@app.post("/api/workspaces")
async def create_workspace(
    name: str = Form(...),
    description: str = Form(None),
    db: Session = Depends(get_db)
):
    """Create a new workspace"""
    try:
        from models import Workspace
        
        # Check if workspace with same name already exists
        existing_workspace = db.query(Workspace).filter(Workspace.name == name).first()
        if existing_workspace:
            return {"success": False, "message": f"Workspace '{name}' already exists"}
        
        # Create new workspace
        workspace = Workspace(
            name=name,
            description=description,
            is_default=False
        )
        
        db.add(workspace)
        db.commit()
        db.refresh(workspace)
        
        return {
            "success": True,
            "workspace": {
                "id": workspace.id,
                "name": workspace.name,
                "description": workspace.description,
                "is_default": workspace.is_default,
                "created_at": workspace.created_at.isoformat() if workspace.created_at else None
            }
        }
    except Exception as e:
        db.rollback()
        return {"success": False, "message": str(e)}

@app.delete("/api/workspaces/{workspace_id}")
async def delete_workspace(
    workspace_id: int, 
    user_email: str = None,
    db: Session = Depends(get_db)
):
    """Delete a workspace - Admin only"""
    try:
        from models import Workspace
        
        # Define admin emails
        ADMIN_EMAILS = [
            "shaik.sharuk@forsysinc.com"
            # Add more admin emails here as needed
        ]
        
        # Check if user is admin
        if not user_email or user_email not in ADMIN_EMAILS:
            return {"success": False, "message": "Only admins can delete workspaces"}
        
        # Find the workspace
        workspace = db.query(Workspace).filter(Workspace.id == workspace_id).first()
        
        if not workspace:
            return {"success": False, "message": "Workspace not found"}
        
        # Prevent deletion of default workspace
        if workspace.is_default:
            return {"success": False, "message": "Cannot delete the default workspace"}
        
        # Delete the workspace
        db.delete(workspace)
        db.commit()
        
        return {"success": True, "message": f"Workspace '{workspace.name}' deleted successfully"}
    except Exception as e:
        db.rollback()
        return {"success": False, "message": str(e)}

# Projects endpoints
@app.post("/api/projects")
async def create_project(
    name: str = Form(...),
    user_email: str = Form(...),
    db: Session = Depends(get_db)
):
    """Create a new project with a default conversation"""
    try:
        # Validate project name
        if not name or not name.strip():
            return {"success": False, "message": "Project name cannot be empty"}
        
        # Begin transaction - create project and conversation atomically
        # Create new project with unique ID
        project = Project(
            name=name.strip(),
            user_email=user_email
        )
        
        db.add(project)
        db.flush()  # Flush to get project.id without committing
        
        # Create default conversation linked to the project
        conversation_id = str(uuid.uuid4())
        chat_id = str(uuid.uuid4())
        conversation_data = {
            "conversation_id": 0,  # Will be updated to actual id after save
            "messages": []
        }
        
        conversation = Conversation(
            conversation_id=conversation_id,
            chat_id=chat_id,
            user_email=user_email,
            project_id=project.id,
            conversation_json=conversation_data
        )
        
        db.add(conversation)
        db.commit()  # Commit both project and conversation together
        db.refresh(project)
        db.refresh(conversation)
        
        # Update conversation_id in JSON to actual database id
        conversation_data["conversation_id"] = conversation.id
        conversation.conversation_json = conversation_data
        db.commit()
        db.refresh(conversation)
        
        return {
            "success": True,
            "project": {
                "id": project.id,
                "name": project.name,
                "user_email": project.user_email,
                "created_at": project.created_at.isoformat() if project.created_at else None
            },
            "conversation": {
                "id": conversation.id,
                "conversation_id": conversation.conversation_id,
                "chat_id": conversation.chat_id,
                "title": "Default chat",
                "project_id": conversation.project_id,
                "user_email": conversation.user_email,
                "created_at": conversation.created_at.isoformat() if conversation.created_at else None
            }
        }
    except Exception as e:
        db.rollback()
        error_msg = str(e)
        import traceback
        error_trace = traceback.format_exc()
        print(f"❌ [PROJECT CREATE ERROR] {error_msg}")
        print(f"❌ [PROJECT CREATE TRACEBACK]\n{error_trace}")
        
        # Check if it's the UUID/integer type error
        if "invalid input syntax for type integer" in error_msg or "InvalidTextRepresentation" in error_msg:
            error_msg = "Database schema error: Projects table has wrong ID type. Please restart the backend server to fix this automatically."
        elif "foreign key" in error_msg.lower() or "constraint" in error_msg.lower():
            error_msg = f"Database constraint error: {error_msg}. Please check foreign key relationships."
        elif "conversations" in error_msg.lower() and "project_id" in error_msg.lower():
            error_msg = f"Conversation creation error: {error_msg}. Please check conversations table schema."
        
        return {"success": False, "message": error_msg}

@app.get("/api/projects")
async def get_projects(
    user_email: str,
    db: Session = Depends(get_db)
):
    """Get all projects for a user with their conversations"""
    try:
        projects = db.query(Project).filter(Project.user_email == user_email).order_by(Project.created_at.desc()).all()
        
        projects_list = []
        for project in projects:
            # Get conversations for this project
            conversations = db.query(Conversation).filter(
                Conversation.project_id == project.id
            ).order_by(Conversation.created_at.asc()).all()
            
            projects_list.append({
                "id": project.id,
                "name": project.name,
                "user_email": project.user_email,
                "created_at": project.created_at.isoformat() if project.created_at else None,
                "conversations": [
                    {
                        "id": conv.id,
                        "conversation_id": conv.conversation_id,
                        "chat_id": conv.chat_id,
                        "title": "Default chat",  # Could be extracted from conversation_json if stored
                        "project_id": conv.project_id,
                        "created_at": conv.created_at.isoformat() if conv.created_at else None
                    }
                    for conv in conversations
                ]
            })
        
        return {
            "success": True,
            "projects": projects_list
        }
    except Exception as e:
        return {"success": False, "message": str(e)}

@app.delete("/api/projects/{project_id}")
async def delete_project(
    project_id: str,
    user_email: str,
    db: Session = Depends(get_db)
):
    """Delete a project"""
    try:
        project = db.query(Project).filter(Project.id == project_id, Project.user_email == user_email).first()
        
        if not project:
            return {"success": False, "message": "Project not found or you don't have permission to delete it"}
        
        db.delete(project)
        db.commit()
        
        return {"success": True, "message": "Project deleted successfully"}
    except Exception as e:
        db.rollback()
        return {"success": False, "message": str(e)}

@app.post("/api/send-sprint-plan-email")
async def send_sprint_plan_email(email_data: dict):
    """Send sprint plan via email with PDF attachment"""
    try:
        import requests
        import os
        import base64
        from datetime import datetime
        
        # Extract email data
        to_email = email_data.get('to')
        subject = email_data.get('subject', 'Sprint Plan Shared')
        body = email_data.get('body', 'Please find the attached sprint plan.')
        sprint_plan_content = email_data.get('sprintPlanContent', '')
        sprint_plan_name = email_data.get('sprintPlanName', 'Sprint Plan')
        
        if not to_email:
            return {"success": False, "message": "Recipient email is required"}
        
        # Generate PDF from sprint plan content
        print(f"Generating PDF for: {sprint_plan_name}")
        pdf_content = await generate_sprint_plan_pdf(sprint_plan_content, sprint_plan_name)
        
        # Generate HTML content for email body
        html_email_body = await generate_html_email_content(sprint_plan_content, sprint_plan_name, body)
        
        if not pdf_content:
            print("ERROR: PDF generation failed!")
            return {"success": False, "message": "Failed to generate PDF"}
        
        print(f"PDF generated successfully, size: {len(pdf_content)} bytes")
        
        # SendGrid API configuration
        SENDGRID_API_KEY = os.getenv('SENDGRID_API_KEY')
        FROM_EMAIL = os.getenv('FROM_EMAIL', 'noreply@yourdomain.com')
        
        if not SENDGRID_API_KEY:
            # Fallback to Gmail SMTP if SendGrid not configured
            print("Using Gmail SMTP...")
            return await send_email_with_pdf_via_smtp(to_email, subject, html_email_body, FROM_EMAIL, pdf_content, sprint_plan_name)
        
        # Encode PDF for SendGrid
        pdf_base64 = base64.b64encode(pdf_content).decode('utf-8')
        pdf_filename = f"{sprint_plan_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        
        # SendGrid API call with PDF attachment
        url = "https://api.sendgrid.com/v3/mail/send"
        
        payload = {
            "personalizations": [
                {
                    "to": [{"email": to_email}],
                    "subject": subject
                }
            ],
            "from": {"email": FROM_EMAIL},
            "content": [
                {
                    "type": "text/html",
                    "value": html_email_body
                },
                {
                    "type": "text/plain",
                    "value": body
                }
            ],
            "attachments": [
                {
                    "content": pdf_base64,
                    "filename": pdf_filename,
                    "type": "application/pdf",
                    "disposition": "attachment"
                }
            ]
        }
        
        headers = {
            "Authorization": f"Bearer {SENDGRID_API_KEY}",
            "Content-Type": "application/json"
        }
        
        response = requests.post(url, json=payload, headers=headers)
        
        if response.status_code == 202:
            return {
                "success": True, 
                "message": f"Sprint plan sent successfully to {to_email} with PDF attachment"
            }
        else:
            return {
                "success": False, 
                "message": f"Failed to send email: {response.text}"
            }
        
    except Exception as e:
        return {"success": False, "message": f"Error sending email: {str(e)}"}

@app.post("/api/send-risk-assessment-email")
async def send_risk_assessment_email(email_data: dict):
    """Send risk assessment via email with PDF attachment"""
    try:
        import requests
        import os
        import base64
        from datetime import datetime
        
        # Extract email data
        to_email = email_data.get('to')
        subject = email_data.get('subject', 'Risk Assessment Shared')
        body = email_data.get('body', 'Please find the attached risk assessment.')
        risk_assessment_content = email_data.get('riskAssessmentContent', '')
        assessment_name = email_data.get('assessmentName', 'Risk Assessment')
        
        if not to_email:
            return {"success": False, "message": "Recipient email is required"}
        
        # Generate PDF from risk assessment content
        print(f"Generating PDF for: {assessment_name}")
        pdf_content = await generate_risk_assessment_pdf(risk_assessment_content, assessment_name)
        
        # Generate HTML content for email body
        html_email_body = await generate_risk_assessment_html_email_content(risk_assessment_content, assessment_name, body)
        
        if not pdf_content:
            print("ERROR: PDF generation failed!")
            return {"success": False, "message": "Failed to generate PDF"}
        
        print(f"PDF generated successfully, size: {len(pdf_content)} bytes")
        
        # SendGrid API configuration
        SENDGRID_API_KEY = os.getenv('SENDGRID_API_KEY')
        FROM_EMAIL = os.getenv('FROM_EMAIL', 'noreply@yourdomain.com')
        
        if not SENDGRID_API_KEY:
            # Fallback to Gmail SMTP if SendGrid not configured
            print("Using Gmail SMTP...")
            return await send_email_with_pdf_via_smtp(to_email, subject, html_email_body, FROM_EMAIL, pdf_content, assessment_name)
        
        # Encode PDF for SendGrid
        pdf_base64 = base64.b64encode(pdf_content).decode('utf-8')
        pdf_filename = f"{assessment_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        
        # SendGrid API call with PDF attachment
        url = "https://api.sendgrid.com/v3/mail/send"
        
        payload = {
            "personalizations": [
                {
                    "to": [{"email": to_email}],
                    "subject": subject
                }
            ],
            "from": {"email": FROM_EMAIL},
            "content": [
                {
                    "type": "text/html",
                    "value": html_email_body
                },
                {
                    "type": "text/plain",
                    "value": body
                }
            ],
            "attachments": [
                {
                    "content": pdf_base64,
                    "filename": pdf_filename,
                    "type": "application/pdf",
                    "disposition": "attachment"
                }
            ]
        }
        
        headers = {
            "Authorization": f"Bearer {SENDGRID_API_KEY}",
            "Content-Type": "application/json"
        }
        
        response = requests.post(url, json=payload, headers=headers)
        
        if response.status_code == 202:
            return {
                "success": True, 
                "message": f"Risk assessment sent successfully to {to_email} with PDF attachment"
            }
        else:
            return {
                "success": False, 
                "message": f"Failed to send email: {response.text}"
            }
        
    except Exception as e:
        return {"success": False, "message": f"Error sending email: {str(e)}"}

async def generate_risk_assessment_html_email_content(risk_assessment_content: str, assessment_name: str, custom_message: str = ""):
    """Generate HTML email content for risk assessment"""
    try:
        import re
        from datetime import datetime
        
        print("Generating HTML email content for risk assessment...")
        
        # Clean up the HTML content
        cleaned_content = risk_assessment_content.strip()
        
        # Remove markdown code blocks
        if cleaned_content.startswith('```html'):
            cleaned_content = cleaned_content.replace('```html', '', 1)
            cleaned_content = cleaned_content.rsplit('```', 1)[0]
        elif cleaned_content.startswith('```'):
            cleaned_content = cleaned_content.replace('```', '', 1)
            cleaned_content = cleaned_content.rsplit('```', 1)[0]
        
        # Extract body content if full HTML
        body_match = re.search(r'<body[^>]*>(.*?)</body>', cleaned_content, re.DOTALL | re.IGNORECASE)
        if body_match:
            cleaned_content = body_match.group(1)
        
        # Wrap all tables in table-wrapper divs for better responsive handling
        cleaned_content = re.sub(
            r'<table([^>]*)>(.*?)</table>',
            r'<div class="table-wrapper"><table\1>\2</table></div>',
            cleaned_content,
            flags=re.DOTALL | re.IGNORECASE
        )
        
        # Create HTML email template
        html_email = f"""
        <!DOCTYPE html>
        <html lang="en">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>Risk Assessment - {assessment_name}</title>
            <style>
                body {{
                    font-family: 'Segoe UI', -apple-system, BlinkMacSystemFont, sans-serif;
                    margin: 0;
                    padding: 10px;
                    background: #f8fafc;
                    color: #2c3e50;
                    line-height: 1.6;
                    font-size: 14px;
                    overflow-x: auto;
                }}
                
                .email-container {{
                    max-width: 100%;
                    margin: 0 auto;
                    background: white;
                    border-radius: 12px;
                    box-shadow: 0 4px 6px rgba(0,0,0,0.1);
                    overflow-x: auto;
                }}
                
                .email-header {{
                    background: linear-gradient(135deg, #e53e3e 0%, #d97706 100%);
                    color: white;
                    padding: 30px;
                    text-align: center;
                }}
                
                .email-header h1 {{
                    margin: 0 0 10px 0;
                    font-size: 2em;
                    font-weight: 700;
                }}
                
                .email-header p {{
                    margin: 0;
                    opacity: 0.9;
                    font-size: 1.1em;
                }}
                
                .email-body {{
                    padding: 30px;
                }}
                
                .custom-message {{
                    background: #fff5f5;
                    border: 1px solid #fc8181;
                    border-radius: 8px;
                    padding: 20px;
                    margin-bottom: 30px;
                    color: #742a2a;
                }}
                
                .table-wrapper {{
                    overflow-x: auto;
                    margin: 20px 0;
                }}
                
                table {{
                    width: 100%;
                    border-collapse: collapse;
                    font-size: 14px;
                }}
                
                table th, table td {{
                    padding: 12px;
                    text-align: left;
                    border: 1px solid #e2e8f0;
                }}
                
                table th {{
                    background: linear-gradient(135deg, #e53e3e 0%, #d97706 100%);
                    color: white;
                    font-weight: 600;
                }}
                
                table tr:nth-child(even) {{
                    background: #f7fafc;
                }}
                
                .attachment-note {{
                    background: #e6fffa;
                    border: 1px solid #38b2ac;
                    border-radius: 8px;
                    padding: 15px;
                    margin-top: 20px;
                    color: #234e52;
                }}
            </style>
        </head>
        <body>
            <div class="email-container">
                <div class="email-header">
                    <h1>Risk Assessment</h1>
                    <p>{assessment_name}</p>
                </div>
                <div class="email-body">
                    {f'<div class="custom-message">{custom_message}</div>' if custom_message else ''}
                    <div>
                        {cleaned_content}
                    </div>
                    <div class="attachment-note">
                        <strong>📎 Attachment:</strong> Please find the complete risk assessment as a PDF attachment.
                    </div>
                </div>
            </div>
        </body>
        </html>
        """
        
        print("HTML email content generated successfully")
        return html_email
        
    except Exception as e:
        print(f"Error generating HTML email content: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

async def generate_risk_assessment_pdf(risk_assessment_content: str, assessment_name: str):
    """Generate PDF from risk assessment content using ReportLab"""
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from io import BytesIO
        from datetime import datetime
        import re
        
        print("Starting PDF generation for risk assessment with ReportLab...")
        
        # Clean up the HTML content
        cleaned_content = risk_assessment_content.strip()
        
        # Remove markdown code blocks
        if cleaned_content.startswith('```html'):
            cleaned_content = cleaned_content.replace('```html', '', 1)
            cleaned_content = cleaned_content.rsplit('```', 1)[0]
        elif cleaned_content.startswith('```'):
            cleaned_content = cleaned_content.replace('```', '', 1)
            cleaned_content = cleaned_content.rsplit('```', 1)[0]
        
        # Extract body content if full HTML
        body_match = re.search(r'<body[^>]*>(.*?)</body>', cleaned_content, re.DOTALL | re.IGNORECASE)
        if body_match:
            cleaned_content = body_match.group(1)
        
        # Create PDF buffer
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, 
                              rightMargin=72, leftMargin=72,
                              topMargin=72, bottomMargin=72)
        
        # Container for PDF elements
        story = []
        
        # Styles
        styles = getSampleStyleSheet()
        
        # Title style with risk assessment colors
        title_style = ParagraphStyle(
            'RiskAssessmentTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#e53e3e'),
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        # Add title
        story.append(Paragraph(assessment_name, title_style))
        story.append(Paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}", 
                              styles['Normal']))
        story.append(Spacer(1, 0.3*inch))
        
        # Parse HTML content properly to handle tables
        story.extend(await parse_html_content_for_pdf(cleaned_content, styles))
        
        # Add footer
        story.append(Spacer(1, 0.5*inch))
        footer_style = ParagraphStyle(
            'Footer',
            parent=styles['Normal'],
            fontSize=9,
            textColor=colors.grey,
            alignment=TA_CENTER
        )
        story.append(Paragraph("This risk assessment was generated and shared via PM Portal", footer_style))
        
        # Build PDF
        doc.build(story)
        
        # Get PDF bytes
        pdf_bytes = buffer.getvalue()
        buffer.close()
        
        print(f"PDF generated successfully using ReportLab, size: {len(pdf_bytes)} bytes")
        return pdf_bytes
        
    except Exception as e:
        print(f"ERROR in PDF generation for risk assessment: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

async def generate_html_email_content(sprint_plan_content: str, sprint_plan_name: str, custom_message: str = ""):
    """Generate HTML email content with proper table structure and formatting"""
    try:
        import re
        from datetime import datetime
        
        print("Generating HTML email content...")
        
        # Clean up the HTML content
        cleaned_content = sprint_plan_content.strip()
        
        # Remove markdown code blocks
        if cleaned_content.startswith('```html'):
            cleaned_content = cleaned_content.replace('```html', '', 1)
            cleaned_content = cleaned_content.rsplit('```', 1)[0]
        elif cleaned_content.startswith('```'):
            cleaned_content = cleaned_content.replace('```', '', 1)
            cleaned_content = cleaned_content.rsplit('```', 1)[0]
        
        # Extract body content if full HTML
        body_match = re.search(r'<body[^>]*>(.*?)</body>', cleaned_content, re.DOTALL | re.IGNORECASE)
        if body_match:
            cleaned_content = body_match.group(1)
        
        # Wrap all tables in table-wrapper divs for better responsive handling
        cleaned_content = re.sub(
            r'<table([^>]*)>(.*?)</table>',
            r'<div class="table-wrapper"><table\1>\2</table></div>',
            cleaned_content,
            flags=re.DOTALL | re.IGNORECASE
        )
        
        # Create HTML email template
        html_email = f"""
        <!DOCTYPE html>
        <html lang="en">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>Sprint Plan - {sprint_plan_name}</title>
            <style>
                body {{
                    font-family: 'Segoe UI', -apple-system, BlinkMacSystemFont, sans-serif;
                    margin: 0;
                    padding: 10px;
                    background: #f8fafc;
                    color: #2c3e50;
                    line-height: 1.6;
                    font-size: 14px;
                    overflow-x: auto;
                }}
                
                .email-container {{
                    max-width: 100%;
                    margin: 0 auto;
                    background: white;
                    border-radius: 12px;
                    box-shadow: 0 4px 6px rgba(0,0,0,0.1);
                    overflow-x: auto;
                }}
                
                .email-header {{
                    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                    color: white;
                    padding: 30px;
                    text-align: center;
                }}
                
                .email-header h1 {{
                    margin: 0 0 10px 0;
                    font-size: 2em;
                    font-weight: 700;
                }}
                
                .email-header p {{
                    margin: 0;
                    opacity: 0.9;
                    font-size: 1.1em;
                }}
                
                .email-body {{
                    padding: 30px;
                }}
                
                .custom-message {{
                    background: #e6fffa;
                    border: 1px solid #81e6d9;
                    border-radius: 8px;
                    padding: 20px;
                    margin-bottom: 30px;
                    color: #234e52;
                }}
                
                .custom-message h3 {{
                    margin: 0 0 10px 0;
                    color: #2d3748;
                    font-size: 1.2em;
                }}
                
                .sprint-content {{
                    background: #f8fafc;
                    border: 1px solid #e2e8f0;
                    border-radius: 8px;
                    padding: 25px;
                }}
                
                .sprint-content h2, .sprint-content h3, .sprint-content h4 {{
                    color: #2d3748;
                    margin: 20px 0 15px 0;
                    font-weight: 600;
                }}
                
                .sprint-content h2 {{
                    font-size: 1.4em;
                    color: #1a202c;
                    border-bottom: 2px solid #e2e8f0;
                    padding-bottom: 10px;
                }}
                
                .sprint-content h3 {{
                    font-size: 1.2em;
                    color: #2d3748;
                }}
                
                .sprint-content h4 {{
                    font-size: 1.1em;
                    color: #4a5568;
                }}
                
                .sprint-content p {{
                    margin: 12px 0;
                    line-height: 1.6;
                    color: #4a5568;
                }}
                
                .sprint-content ul, .sprint-content ol {{
                    margin: 15px 0;
                    padding-left: 25px;
                }}
                
                .sprint-content li {{
                    margin: 8px 0;
                    line-height: 1.6;
                    color: #4a5568;
                }}
                
                .sprint-content strong {{
                    color: #2d3748;
                    font-weight: 600;
                }}
                
                .sprint-content em {{
                    color: #718096;
                    font-style: italic;
                }}
                
                .sprint-content table {{
                    width: 100%;
                    border-collapse: collapse;
                    margin: 20px 0;
                    background: white;
                    border-radius: 8px;
                    box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                    table-layout: auto;
                    min-width: 100%;
                }}
                
                .sprint-content th {{
                    background: #667eea;
                    color: white;
                    padding: 15px 12px;
                    text-align: left;
                    font-weight: 600;
                    font-size: 0.9em;
                    text-transform: uppercase;
                    letter-spacing: 0.5px;
                    white-space: normal;
                    word-wrap: break-word;
                    overflow-wrap: break-word;
                    vertical-align: top;
                    min-width: 120px;
                }}
                
                .sprint-content td {{
                    padding: 15px 12px;
                    border-bottom: 1px solid #e2e8f0;
                    color: #4a5568;
                    white-space: normal;
                    word-wrap: break-word;
                    overflow-wrap: break-word;
                    font-size: 0.9em;
                    line-height: 1.5;
                    vertical-align: top;
                    min-width: 120px;
                }}
                
                .sprint-content tr:nth-child(even) {{
                    background: #f8fafc;
                }}
                
                .sprint-content tr:hover {{
                    background: #edf2f7;
                }}
                
                /* Responsive table wrapper */
                .table-wrapper {{
                    overflow-x: auto;
                    margin: 20px 0;
                    border-radius: 8px;
                    box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                    width: 100%;
                }}
                
                .table-wrapper table {{
                    margin: 0;
                    min-width: 800px;
                    width: 100%;
                }}
                
                /* Specific column width improvements */
                .sprint-content table td:nth-child(1),
                .sprint-content table th:nth-child(1) {{
                    min-width: 100px;
                    max-width: 150px;
                }}
                
                .sprint-content table td:nth-child(2),
                .sprint-content table th:nth-child(2) {{
                    min-width: 200px;
                    max-width: 300px;
                }}
                
                .sprint-content table td:nth-child(3),
                .sprint-content table th:nth-child(3) {{
                    min-width: 120px;
                    max-width: 180px;
                }}
                
                .sprint-content table td:nth-child(4),
                .sprint-content table th:nth-child(4) {{
                    min-width: 250px;
                    max-width: 400px;
                }}
                
                .sprint-content table td:nth-child(5),
                .sprint-content table th:nth-child(5) {{
                    min-width: 120px;
                    max-width: 180px;
                }}
                
                .email-footer {{
                    background: #f7fafc;
                    padding: 20px 30px;
                    text-align: center;
                    color: #718096;
                    font-size: 0.9em;
                    border-top: 1px solid #e2e8f0;
                }}
                
                .attachment-notice {{
                    background: #fff5f5;
                    border: 1px solid #feb2b2;
                    border-radius: 8px;
                    padding: 15px;
                    margin-top: 20px;
                    color: #742a2a;
                }}
                
                .attachment-notice strong {{
                    color: #c53030;
                }}
                
                
                /* Scroll hint */
                .scroll-hint {{
                    background: #e6fffa;
                    border: 1px solid #81e6d9;
                    border-radius: 6px;
                    padding: 10px;
                    margin: 10px 0;
                    color: #234e52;
                    font-size: 0.9em;
                    text-align: center;
                }}
            </style>
        </head>
        <body>
            <div class="email-container">
                <div class="email-header">
                    <h1>Sprint Plan - {sprint_plan_name}</h1>
                    <p>Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}</p>
                </div>
                
                <div class="email-body">
                    {f'<div class="custom-message"><h3>Message from sender:</h3><p>{custom_message}</p></div>' if custom_message else ''}
                    
                    <div class="scroll-hint">
                        💡 <strong>Tip:</strong> If tables appear cut off, scroll horizontally to see all columns
                    </div>
                    
                    <div class="sprint-content">
                        {cleaned_content}
                    </div>
                    
                    <div class="attachment-notice">
                        <strong>📎 PDF Attachment:</strong> A formatted PDF version of this sprint plan is attached to this email for your convenience.
                    </div>
                </div>
                
                <div class="email-footer">
                    <p>This sprint plan was generated and shared via PM Portal</p>
                    <p>For any questions or concerns, please contact the sender.</p>
                </div>
            </div>
        </body>
        </html>
        """
        
        print("HTML email content generated successfully")
        return html_email
        
    except Exception as e:
        print(f"ERROR generating HTML email content: {str(e)}")
        import traceback
        traceback.print_exc()
        # Fallback to simple HTML
        return f"""
        <html>
        <body>
            <h2>Sprint Plan - {sprint_plan_name}</h2>
            <p>{custom_message}</p>
            <div>{sprint_plan_content}</div>
            <p><strong>Note:</strong> A PDF attachment is also included with this email.</p>
        </body>
        </html>
        """

async def parse_html_content_for_pdf(html_content: str, styles):
    """Parse HTML content and convert to ReportLab elements, properly handling tables"""
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from io import BytesIO
        from datetime import datetime
        import re
        
        elements = []
        
        # Clean up HTML
        cleaned_content = re.sub(r'<style[^>]*>.*?</style>', '', html_content, flags=re.DOTALL)
        cleaned_content = re.sub(r'<script[^>]*>.*?</script>', '', cleaned_content, flags=re.DOTALL)
        
        # Split content by tables and other elements
        parts = re.split(r'(<table[^>]*>.*?</table>)', cleaned_content, flags=re.DOTALL | re.IGNORECASE)
        
        for part in parts:
            part = part.strip()
            if not part:
                continue
                
            # Check if this is a table
            if re.match(r'<table[^>]*>', part, re.IGNORECASE):
                # Parse table
                table_elements = parse_html_table_for_pdf(part, styles)
                elements.extend(table_elements)
            else:
                # Parse other content (headings, paragraphs, etc.)
                other_elements = parse_html_text_for_pdf(part, styles)
                elements.extend(other_elements)
        
        return elements
        
    except Exception as e:
        print(f"ERROR parsing HTML content for PDF: {str(e)}")
        import traceback
        traceback.print_exc()
        return []

def parse_html_table_for_pdf(table_html: str, styles):
    """Parse HTML table and convert to ReportLab Table"""
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        import re
        
        elements = []
        
        # Extract table rows
        rows = re.findall(r'<tr[^>]*>(.*?)</tr>', table_html, re.DOTALL | re.IGNORECASE)
        
        if not rows:
            return elements
        
        table_data = []
        
        for row in rows:
            # Extract cells (both th and td)
            cells = re.findall(r'<(?:th|td)[^>]*>(.*?)</(?:th|td)>', row, re.DOTALL | re.IGNORECASE)
            row_data = []
            
            for cell in cells:
                # Clean cell content
                cell_text = re.sub(r'<[^>]+>', '', cell)  # Remove HTML tags
                cell_text = cell_text.strip()
                if not cell_text:
                    cell_text = " "
                
                # Handle long text better - split into multiple lines if needed
                if len(cell_text) > 80:
                    # Split long text into multiple lines for better readability
                    words = cell_text.split()
                    lines = []
                    current_line = ""
                    
                    for word in words:
                        if len(current_line + " " + word) <= 80:
                            current_line += (" " + word) if current_line else word
                        else:
                            if current_line:
                                lines.append(current_line)
                            current_line = word
                    
                    if current_line:
                        lines.append(current_line)
                    
                    # Join lines with line breaks for PDF
                    cell_text = "\n".join(lines)
                
                row_data.append(cell_text)
            
            if row_data:
                table_data.append(row_data)
        
        if table_data:
            # Calculate column widths based on content and page width
            page_width = A4[0] - 144  # A4 width minus margins (72*2)
            num_cols = len(table_data[0]) if table_data else 1
            
            # Calculate column widths based on number of columns
            if num_cols <= 2:
                # For 1-2 columns, use wider columns
                col_widths = [page_width / num_cols] * num_cols
            elif num_cols <= 4:
                # For 3-4 columns, use moderate width
                col_widths = [page_width / num_cols] * num_cols
            else:
                # For 5+ columns, use smaller width but ensure minimum readability
                max_col_width = page_width / num_cols
                min_col_width = 60  # Minimum column width for readability
                if max_col_width < min_col_width:
                    # If columns would be too narrow, use minimum width and allow horizontal overflow
                    col_widths = [min_col_width] * num_cols
                else:
                    col_widths = [max_col_width] * num_cols
            
            # Create ReportLab Table with calculated column widths
            pdf_table = Table(table_data, colWidths=col_widths)
            
            # Style the table
            table_style = TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#667eea')),  # Header background
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),  # Header text color
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),  # Left align all cells
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),  # Bold header
                ('FONTSIZE', (0, 0), (-1, 0), 9),  # Header font size
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),  # Regular font for data
                ('FONTSIZE', (0, 1), (-1, -1), 8),  # Data font size
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),  # Header padding
                ('TOPPADDING', (0, 0), (-1, 0), 12),  # Header padding
                ('BOTTOMPADDING', (0, 1), (-1, -1), 10),  # Data padding
                ('TOPPADDING', (0, 1), (-1, -1), 10),  # Data padding
                ('LEFTPADDING', (0, 0), (-1, -1), 8),  # Left padding
                ('RIGHTPADDING', (0, 0), (-1, -1), 8),  # Right padding
                ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e2e8f0')),  # Grid lines
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),  # Top align all cells
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8fafc')]),  # Alternating row colors
                ('LINEBELOW', (0, 0), (-1, 0), 1, colors.HexColor('#667eea')),  # Header bottom line
                ('LINEBELOW', (0, 1), (-1, -1), 0.5, colors.HexColor('#e2e8f0')),  # Row separator lines
                ('MINROWHEIGHT', (0, 0), (-1, -1), 20),  # Minimum row height
            ])
            
            pdf_table.setStyle(table_style)
            
            # Add spacing before and after table
            elements.append(Spacer(1, 0.15*inch))
            elements.append(pdf_table)
            elements.append(Spacer(1, 0.15*inch))
        
        return elements
        
    except Exception as e:
        print(f"ERROR parsing HTML table for PDF: {str(e)}")
        import traceback
        traceback.print_exc()
        return []

def parse_html_text_for_pdf(text_content: str, styles):
    """Parse HTML text content (headings, paragraphs) and convert to ReportLab elements"""
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        import re
        
        elements = []
        
        # Split by headings
        sections = re.split(r'(<h[1-6][^>]*>.*?</h[1-6]>)', text_content, flags=re.DOTALL | re.IGNORECASE)
        
        for section in sections:
            section = section.strip()
            if not section:
                continue
            
            # Check if this is a heading
            heading_match = re.match(r'<h([1-6])[^>]*>(.*?)</h[1-6]>', section, re.DOTALL | re.IGNORECASE)
            if heading_match:
                level = int(heading_match.group(1))
                heading_text = re.sub(r'<[^>]+>', '', heading_match.group(2)).strip()
                
                if heading_text:
                    # Create heading style based on level
                    if level <= 2:
                        heading_style = ParagraphStyle(
                            'CustomHeading',
                            parent=styles['Heading2'],
                            fontSize=16,
                            textColor=colors.HexColor('#34495e'),
                            spaceAfter=12,
                            spaceBefore=20
                        )
                    else:
                        heading_style = ParagraphStyle(
                            'CustomSubHeading',
                            parent=styles['Heading3'],
                            fontSize=14,
                            textColor=colors.HexColor('#4a5568'),
                            spaceAfter=10,
                            spaceBefore=15
                        )
                    
                    elements.append(Paragraph(heading_text, heading_style))
            else:
                # Parse paragraphs and other content
                paragraphs = re.split(r'</p>', section)
                for para in paragraphs:
                    para = para.replace('<p>', '').replace('<p ', '<p ').strip()
                    if para:
                        # Clean HTML tags but preserve basic formatting
                        para = para.replace('<strong>', '<b>').replace('</strong>', '</b>')
                        para = para.replace('<em>', '<i>').replace('</em>', '</i>')
                        para = para.replace('<br>', '<br/>')
                        para = para.replace('<br/><br/>', '<br/><br/>')
                        
                        # Remove any remaining HTML tags except basic formatting
                        para = re.sub(r'<(?!/?(?:b|i|br/?)>)[^>]+>', '', para)
                        
                        if para.strip():
                            try:
                                elements.append(Paragraph(para, styles['Normal']))
                                elements.append(Spacer(1, 0.1*inch))
                            except:
                                # If paragraph fails, add as plain text
                                plain_text = re.sub(r'<[^>]+>', '', para)
                                if plain_text.strip():
                                    elements.append(Paragraph(plain_text, styles['Normal']))
                                    elements.append(Spacer(1, 0.1*inch))
        
        return elements
        
    except Exception as e:
        print(f"ERROR parsing HTML text for PDF: {str(e)}")
        import traceback
        traceback.print_exc()
        return []

async def generate_sprint_plan_pdf(sprint_plan_content: str, sprint_plan_name: str):
    """Generate PDF from sprint plan HTML content using ReportLab"""
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from io import BytesIO
        from datetime import datetime
        from html.parser import HTMLParser
        import re
        
        print("Starting PDF generation with ReportLab...")
        
        # Clean up the HTML content
        cleaned_content = sprint_plan_content.strip()
        
        # Remove markdown code blocks
        if cleaned_content.startswith('```html'):
            cleaned_content = cleaned_content.replace('```html', '', 1)
            cleaned_content = cleaned_content.rsplit('```', 1)[0]
        elif cleaned_content.startswith('```'):
            cleaned_content = cleaned_content.replace('```', '', 1)
            cleaned_content = cleaned_content.rsplit('```', 1)[0]
        
        # Extract body content if full HTML
        body_match = re.search(r'<body[^>]*>(.*?)</body>', cleaned_content, re.DOTALL | re.IGNORECASE)
        if body_match:
            cleaned_content = body_match.group(1)
        
        # Create PDF buffer
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, 
                              rightMargin=72, leftMargin=72,
                              topMargin=72, bottomMargin=72)
        
        # Container for PDF elements
        story = []
        
        # Styles
        styles = getSampleStyleSheet()
        
        # Title style
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#2c3e50'),
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        # Add title
        story.append(Paragraph(sprint_plan_name, title_style))
        story.append(Paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}", 
                              styles['Normal']))
        story.append(Spacer(1, 0.3*inch))
        
        # Parse HTML content properly to handle tables
        story.extend(await parse_html_content_for_pdf(cleaned_content, styles))
        
        # Add footer
        story.append(Spacer(1, 0.5*inch))
        footer_style = ParagraphStyle(
            'Footer',
            parent=styles['Normal'],
            fontSize=9,
            textColor=colors.grey,
            alignment=TA_CENTER
        )
        story.append(Paragraph("This sprint plan was generated and shared via PM Portal", footer_style))
        
        # Build PDF
        doc.build(story)
        
        # Get PDF bytes
        pdf_bytes = buffer.getvalue()
        buffer.close()
        
        print(f"PDF generated successfully using ReportLab, size: {len(pdf_bytes)} bytes")
        return pdf_bytes
        
    except Exception as e:
        print(f"ERROR in PDF generation: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

async def send_email_with_pdf_via_smtp(to_email: str, subject: str, body: str, from_email: str, pdf_content: bytes, sprint_plan_name: str):
    """Fallback SMTP email sending with PDF attachment"""
    try:
        import smtplib
        from email.mime.text import MIMEText
        from email.mime.multipart import MIMEMultipart
        from email.mime.application import MIMEApplication
        import os
        from datetime import datetime
        
        # SMTP configuration
        SMTP_SERVER = os.getenv('SMTP_SERVER', 'smtp.gmail.com')
        SMTP_PORT = int(os.getenv('SMTP_PORT', '587'))
        SMTP_USERNAME = os.getenv('SMTP_USERNAME')
        SMTP_PASSWORD = os.getenv('SMTP_PASSWORD')
        
        if not SMTP_USERNAME or not SMTP_PASSWORD:
            return {
                "success": False, 
                "message": "Email configuration not set. Please configure SMTP_USERNAME and SMTP_PASSWORD in environment variables."
            }
        
        # Create message with attachment
        msg = MIMEMultipart()
        msg['Subject'] = subject
        msg['From'] = from_email
        msg['To'] = to_email
        
        # Add body (HTML content)
        msg.attach(MIMEText(body, 'html'))
        
        # Add PDF attachment
        if pdf_content:
            pdf_filename = f"{sprint_plan_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            print(f"Attaching PDF: {pdf_filename}, size: {len(pdf_content)} bytes")
            pdf_attachment = MIMEApplication(pdf_content, _subtype='pdf')
            pdf_attachment.add_header('Content-Disposition', 'attachment', filename=pdf_filename)
            msg.attach(pdf_attachment)
            print("PDF attachment added successfully")
        else:
            print("WARNING: No PDF content to attach!")
        
        # Send email
        print(f"Connecting to SMTP server: {SMTP_SERVER}:{SMTP_PORT}")
        server = smtplib.SMTP(SMTP_SERVER, SMTP_PORT)
        server.starttls()
        print("Logging in to SMTP...")
        server.login(SMTP_USERNAME, SMTP_PASSWORD)
        print("Sending email...")
        server.send_message(msg)
        server.quit()
        print("Email sent successfully!")
        
        return {
            "success": True, 
            "message": f"Sprint plan sent successfully to {to_email} with PDF attachment"
        }
        
    except Exception as e:
        return {"success": False, "message": f"SMTP Error: {str(e)}"}

@app.get("/api/workspaces/default")
async def get_default_workspace(db: Session = Depends(get_db)):
    """Get the default EJM workspace"""
    try:
        from models import Workspace
        
        # Check if EJM workspace exists
        workspace = db.query(Workspace).filter(Workspace.name == "EJM").first()
        
        if not workspace:
            # Create EJM workspace if it doesn't exist
            workspace = Workspace(
                name="EJM",
                description="Default EJM workspace",
                is_default=True
            )
            db.add(workspace)
            db.commit()
            db.refresh(workspace)
        
        return {
            "success": True,
            "workspace": {
                "id": workspace.id,
                "name": workspace.name,
                "description": workspace.description,
                "is_default": workspace.is_default
            }
        }
    except Exception as e:
        return {"success": False, "message": str(e)}


if __name__ == "__main__":
    import uvicorn
    # Use 0.0.0.0 to allow external connections on the same network
    uvicorn.run(app, host="0.0.0.0", port=8000)
