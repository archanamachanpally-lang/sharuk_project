import io
import re
import time
import json
from typing import Dict, List, Any, Optional
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from .gemini_service import gemini_service

class DocxService:
    """Service for parsing DOCX files and extracting structured data using LLM"""
    
    def __init__(self):
        self.section_mappings = {
            'sprint_overview': [
                'Sprint Overview', 'Sprint Overview & Proposed Goal', 'I. Sprint Overview',
                'Sprint Number', 'Sprint Dates', 'Sprint Duration', 'Team Name', 'Sprint Goal'
            ],
            'team_capacity': [
                'Team Capacity', 'Team Capacity & Availability', 'II. Team Capacity',
                'Available Working Hours', 'Team Member Details', 'Individual Team Member'
            ],
            'product_backlog': [
                'Product Backlog', 'Product Backlog Items', 'III. Product Backlog',
                'Backlog Items', 'User Stories', 'Epics', 'Features'
            ],
            'definition_of_done': [
                'Definition of Done', 'IV. Definition of Done', 'DoD',
                'Acceptance Criteria', 'Quality Gates'
            ],
            'risks_and_impediments': [
                'Risks and Impediments', 'V. Risks and Impediments', 'Risk Assessment',
                'Potential Issues', 'Blockers'
            ]
        }
    
    def parse_docx_file(self, file_content: bytes) -> Dict[str, Any]:
        """Parse DOCX file and extract structured data using LLM"""
        try:
            # Extract raw text from DOCX
            raw_text = self._extract_text_from_document(file_content)
            
            if not raw_text.strip():
                return {
                    'success': False,
                    'message': 'No text content found in the document'
                }
            
            # Use LLM to convert unstructured text to structured format
            structured_data = self._convert_to_structured_format(raw_text)
            
            if structured_data['success']:
                return {
                    'success': True,
                    'data': structured_data['data'],
                    'message': 'Document parsed successfully using LLM'
                }
            else:
                return {
                    'success': False,
                    'message': f'LLM parsing failed: {structured_data["error"]}'
                }
                
        except Exception as e:
            return {
                'success': False,
                'message': f'Error processing document: {str(e)}'
            }
    
    def _extract_text_from_document(self, file_content: bytes) -> str:
        """Extract raw text from DOCX document"""
        try:
            doc = Document(io.BytesIO(file_content))
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            return '\n'.join(text_parts)
            
        except Exception as e:
            raise Exception(f"Error extracting text from document: {str(e)}")
    
    def _convert_to_structured_format(self, raw_text: str) -> Dict[str, Any]:
        """Use LLM to convert unstructured text to structured format"""
        try:
            # Create a comprehensive prompt for the LLM
            prompt = self._create_llm_prompt(raw_text)
            
            # Call Gemini service
            messages = [
                {
                    "role": "system", 
                    "content": "You are an expert at analyzing sprint planning documents and converting unstructured content into structured JSON format."
                },
                {
                    "role": "user", 
                    "content": prompt
                }
            ]
            
            result = gemini_service.chat(messages, max_tokens=4000)
            
            if not result['success']:
                return {
                    'success': False,
                    'error': f"LLM service error: {result.get('error', 'Unknown error')}"
                }
            
            # Parse the LLM response
            structured_data = self._parse_llm_response(result['response'])
            
            return {
                'success': True,
                'data': structured_data
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f"Error in LLM conversion: {str(e)}"
            }
    
    def _create_llm_prompt(self, raw_text: str) -> str:
        """Create a comprehensive prompt for the LLM"""
        return f"""
Please analyze the following sprint planning document and convert it into a structured JSON format.

DOCUMENT CONTENT:
{raw_text}

INSTRUCTIONS:
1. Extract all relevant information from the document
2. Convert it into the following JSON structure
3. Use "NA" for any missing information
4. Ensure all arrays are properly formatted
 5. Extract team members with their FULL names (including identifiers like A, B, C), roles, and working hours
 6. Extract backlog items with PBI numbers, user stories, acceptance criteria, priority, and effort estimates
 7. Convert effort estimates to hours where possible

REQUIRED JSON STRUCTURE:
{{
  "sprintPlan": {{
    "sprintOverview": {{
      "sprintNumber": "string",
      "sprintDates": "string", 
      "sprintDuration": "string",
      "teamName": "string",
      "sprintGoal": "string"
    }},
    "teamCapacity": {{
      "totalHoursPerPerson": "number",
      "dailyProductiveHours": "number",
      "teamMembers": [
        {{
          "name": "string (full name with identifier, e.g., 'Developer A', 'Product Owner G')",
          "role": "string (role only, e.g., 'Developer', 'Product Owner')", 
          "workingHours": "number"
        }}
      ],
      "historicalStoryPoints": "string"
    }},
    "productBacklog": {{
      "backlogItems": [
        {{
          "pbiNumber": "number",
          "userStorySummary": "string",
          "acceptanceCriteria": ["string"],
          "priority": "string",
          "effortEstimateHours": "number"
        }}
      ]
    }},
    "definitionOfDone": ["string"],
    "risksAndImpediments": ["string"]
  }}
}}

IMPORTANT:
- Return ONLY valid JSON, no additional text
- Ensure all JSON is properly formatted and valid
- Use arrays for multiple items (team members, backlog items, etc.)
- Convert any effort estimates to hours
 - Extract PBI numbers from user stories if present
 - Split acceptance criteria into separate array items
 - Use "High", "Medium", "Low" for priorities
 - Use "NA" for missing values, not null or empty strings
 - For team members, extract the FULL name including any identifiers (e.g., "Developer A", "Product Owner G")
"""
    
    def _parse_llm_response(self, llm_response: str) -> Dict[str, Any]:
        """Parse the LLM response and extract JSON"""
        try:
            # Clean the response to extract JSON
            response_text = llm_response.strip()
            
            # Try to find JSON in the response
            json_start = response_text.find('{')
            json_end = response_text.rfind('}') + 1
            
            if json_start == -1 or json_end == 0:
                raise Exception("No JSON found in LLM response")
            
            json_str = response_text[json_start:json_end]
            
            # Parse the JSON
            data = json.loads(json_str)
            
            # Validate the structure
            if 'sprintPlan' not in data:
                raise Exception("Invalid JSON structure: missing 'sprintPlan' key")
            
            print(f"🔍 [DOCX SERVICE] Returning data: {json.dumps(data, indent=2)}")
            return data
            
        except json.JSONDecodeError as e:
            raise Exception(f"Invalid JSON in LLM response: {str(e)}")
        except Exception as e:
            raise Exception(f"Error parsing LLM response: {str(e)}")

# Create service instance
docx_service = DocxService()
